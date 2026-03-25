"""
LG Factory — DOCX Exporter
Parametrized DOCX generation from markdown outputs.
"""

import os
import re
from pathlib import Path
from typing import Optional

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def add_formatted_text(paragraph, text: str):
    """Handle basic markdown bold and inline code."""
    parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(100, 50, 50)
        else:
            paragraph.add_run(part)


def flush_table(doc, table_rows: list):
    """Create a docx table from collected rows."""
    if not table_rows:
        return
    max_cols = max(len(r) for r in table_rows)
    table = doc.add_table(rows=len(table_rows), cols=max_cols)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(table_rows):
        for j, cell_text in enumerate(row_data):
            if j < max_cols:
                cell = table.cell(i, j)
                cell.text = ''
                p = cell.paragraphs[0]
                add_formatted_text(p, cell_text)
                for run in p.runs:
                    run.font.size = Pt(9)
                    if i == 0:
                        run.bold = True
    doc.add_paragraph()


def markdown_to_docx(doc, md_content: str):
    """Convert markdown content to docx paragraphs."""
    lines = md_content.split('\n')
    table_rows = []
    in_code_block = False

    for line in lines:
        stripped = line.strip()

        # Code block toggles
        if stripped.startswith('```'):
            if in_code_block:
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(line.rstrip())
            run.font.name = 'Courier New'
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(60, 60, 60)
            continue

        # Table rows
        if stripped.startswith('|') and '|' in stripped[1:]:
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            if not all(set(c) <= set('- :') for c in cells):
                table_rows.append(cells)
                continue

        # Flush pending table
        if table_rows:
            flush_table(doc, table_rows)
            table_rows = []

        # Horizontal rules
        if stripped == '---':
            p = doc.add_paragraph()
            run = p.add_run('─' * 60)
            run.font.size = Pt(6)
            run.font.color.rgb = RGBColor(180, 180, 180)
            continue

        if not stripped:
            continue

        # Headings
        if stripped.startswith('# '):
            h = stripped[2:].strip()
            heading = doc.add_heading(h, level=1)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(31, 58, 95)
            continue
        if stripped.startswith('## '):
            h = stripped[3:].strip()
            heading = doc.add_heading(h, level=2)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(31, 58, 95)
            continue
        if stripped.startswith('### '):
            h = stripped[4:].strip()
            heading = doc.add_heading(h, level=3)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(68, 114, 196)
            continue
        if stripped.startswith('#### '):
            h = stripped[5:].strip()
            doc.add_heading(h, level=4)
            continue

        # Blockquotes
        if stripped.startswith('> '):
            text = stripped[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(text)
            run.font.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(80, 80, 80)
            continue

        # Bullet points
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            add_formatted_text(p, text)
            continue

        # Numbered items
        num_match = re.match(r'^(\d+)\.\s', stripped)
        if num_match:
            text = stripped[num_match.end():].strip()
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            add_formatted_text(p, text)
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        add_formatted_text(p, stripped)

    # Flush remaining table
    if table_rows:
        flush_table(doc, table_rows)


def generate_unit_docx(unit_outputs: dict, program_name: str, unit_name: str,
                        unit_number: str, cefr: str, grammar_targets: list,
                        key_vocabulary: list, output_path: Path) -> Path:
    """Generate a complete unit DOCX from all PM outputs."""

    if not HAS_DOCX:
        raise ImportError("pip install python-docx")

    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Cover page
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(program_name.upper())
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(31, 58, 95)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Unidad {unit_number}: {unit_name}')
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(68, 114, 196)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Guía de Aprendizaje Completa')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(100, 100, 100)

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('FPI SENA — Bilingüismo')
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Nivel CEFR: {cefr}')
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Grammar: {", ".join(grammar_targets)}')
    run.font.size = Pt(11)
    run.font.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Key Terms: {", ".join(key_vocabulary[:10])}...')
    run.font.size = Pt(9)
    run.font.italic = True

    # Table of contents
    doc.add_page_break()
    doc.add_heading('CONTENIDO', level=1)
    pm_order = [
        "PM-1.2", "PM-2.1", "PM-2.2", "PM-2.3", "PM-2.4",
        "PM-2.5", "PM-2.6", "PM-2.7", "PM-2.8", "PM-2.9", "PM-2.10"
    ]
    for pm in pm_order:
        if pm in unit_outputs and unit_outputs[pm]:
            p = doc.add_paragraph(pm, style='List Number')
            p.paragraph_format.space_after = Pt(4)

    # Content sections
    for pm in pm_order:
        content = unit_outputs.get(pm)
        if not content:
            continue
        doc.add_page_break()
        markdown_to_docx(doc, content)

    # Footer
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('— FIN DE LA GUÍA —')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(31, 58, 95)

    # Save
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"  DOCX saved: {output_path} ({output_path.stat().st_size:,} bytes)")
    return output_path


def generate_gfpi_docx(gfpi_content: str, program_name: str, unit_name: str,
                         unit_number: str, output_path: Path) -> Path:
    """Generate DOCX for GFPI-F-135 document."""
    if not HAS_DOCX:
        raise ImportError("pip install python-docx")

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    markdown_to_docx(doc, gfpi_content)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"  GFPI DOCX saved: {output_path}")
    return output_path
