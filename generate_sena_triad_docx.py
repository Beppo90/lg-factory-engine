#!/usr/bin/env python3
"""Generate SENA Evaluation Triad as a Word document."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def sh(cell,color):
    e=OxmlElement('w:shd');e.set(qn('w:fill'),color);e.set(qn('w:val'),'clear')
    cell._tc.get_or_add_tcPr().append(e)

def tbl(doc,headers,rows,sc='1F3A5F'):
    t=doc.add_table(rows=1,cols=len(headers));t.style='Table Grid';t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i];c.text='';r=c.paragraphs[0].add_run(str(h))
        r.font.size=Pt(10);r.font.name='Calibri';r.bold=True;r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF);sh(c,sc)
    for rd in rows:
        row=t.add_row()
        for i,v in enumerate(rd):
            c=row.cells[i];c.text='';r=c.paragraphs[0].add_run(str(v));r.font.size=Pt(9);r.font.name='Calibri'
    return t

def h(d,t,l=1):
    hd=d.add_heading(t,level=l)
    for r in hd.runs:r.font.name='Calibri'

def bp(d,l,v=''):
    p=d.add_paragraph();r=p.add_run(l);r.bold=True;r.font.name='Calibri';r.font.size=Pt(11)
    if v:r2=p.add_run(v);r2.font.name='Calibri';r2.font.size=Pt(11)

def ip(d,t):
    p=d.add_paragraph();r=p.add_run(t);r.italic=True;r.font.name='Calibri';r.font.size=Pt(11)

def st(d,t):
    p=d.add_paragraph(t)
    for r in p.runs:r.font.name='Calibri';r.font.size=Pt(11)

def line(d,n=1):
    for _ in range(n):
        p=d.add_paragraph('____________________________________________________________')
        for r in p.runs:r.font.name='Calibri';r.font.size=Pt(11);r.font.color.rgb=RGBColor(0xBB,0xBB,0xBB)

doc = Document()
for s in doc.sections:
    s.top_margin=Cm(2.5);s.bottom_margin=Cm(2.5);s.left_margin=Cm(2.5);s.right_margin=Cm(2.5)
style=doc.styles['Normal'];style.font.name='Calibri';style.font.size=Pt(11)

# TITLE
for _ in range(3):doc.add_paragraph('')
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('ADSO — GUÍA 1: THE HARDWARE SPECIALIST');r.bold=True;r.font.size=Pt(20);r.font.name='Calibri';r.font.color.rgb=RGBColor(0x1F,0x3A,0x5F)
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Instrumentos de Evaluación SENA');r.font.size=Pt(16);r.font.name='Calibri';r.font.color.rgb=RGBColor(0x44,0x72,0xC4)
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('IE-01 (Conocimiento) + IE-02 (Producto) + IE-03 (Desempeño)');r.font.size=Pt(14);r.font.name='Calibri';r.italic=True
doc.add_paragraph('')
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('⚠️ SOLO PARA EL INSTRUCTOR. No distribuir a los aprendices.');r.bold=True;r.font.size=Pt(12);r.font.name='Calibri';r.font.color.rgb=RGBColor(0xC0,0x39,0x2B)
doc.add_paragraph('')
for txt in ['Sistema de Prompts Maestros — LG Factory — FPI SENA — Bilingüismo','Instructor Sergio Cortés Perdomo · Marzo 2026']:
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(txt);r.font.size=Pt(10);r.font.name='Calibri';r.font.color.rgb=RGBColor(0x7F,0x8C,0x8D)
doc.add_page_break()

# ═══ IE-01 SUMMARY ═══
h(doc,'IE-01: CUESTIONARIO TÉCNICO — CASOS DE HARDWARE Y SELECCIÓN DE COMPONENTES',1)
bp(doc,'Evidencia SENA: ','Conocimiento')
bp(doc,'Técnica: ','Formulación de preguntas')
bp(doc,'Puntuación: ','50 puntos')
bp(doc,'Ubicación: ','Session 7 (Prove What You Know)')
bp(doc,'Documento: ','ADSO — GUÍA 1 — The Hardware Specialist — Cuestionario Técnico (IE-01).docx')
doc.add_paragraph('')
tbl(doc,['Sección','Contenido','Pts'],[
    ['1. Reading','Email de Laura Méndez — workstations para testing lab','10'],
    ['2. Writing','Miguel Torres — Tech Request para estación de trabajo','10'],
    ['3. Listening','Voicemail IT Department — plan de upgrades','10'],
    ['4. Vocabulary HOTS','Diagnóstico + Odd One Out + Presupuesto','10'],
    ['5. Grammar HOTS','Inventario + Error Log + Recomendación','10'],
    ['TOTAL','','50'],
],sc='1F3A5F')
doc.add_page_break()

# ═══ IE-02 ═══
h(doc,'IE-02: TECH REQUEST EMAIL — RÚBRICA ANALÍTICA DE ESCRITURA TÉCNICA',1)
bp(doc,'Evidencia SENA: ','Producto')
bp(doc,'Técnica: ','Valoración de producto')
bp(doc,'Puntuación: ','20 puntos')
bp(doc,'Ubicación: ','Session 5 (Write It Right)')
doc.add_paragraph('')

h(doc,'Datos del Aprendiz',2)
tbl(doc,['Campo','Dato'],[['Nombre',''],['Ficha',''],['Fecha',''],['Estación de trabajo #','']],sc='4472C4')
doc.add_paragraph('')

h(doc,'Criterio 1: Formato del Email (4 puntos)',2)
tbl(doc,['Nivel','Pts','Descripción'],[
    ['Competente','4','Todos los elementos: From, To, Subject (#estación), Greeting, cuerpo, Closing, firma. Profesional.'],
    ['Adecuado','3','5-6 de 7 elementos. Falta 1-2 menores.'],
    ['En desarrollo','2','3-4 de 7 elementos. Formato parcial.'],
    ['Inicial','1','1-2 elementos. Formato irreconocible.'],
    ['No evidenciado','0','Sin formato o inentregable.'],
],sc='2E75B6')
bp(doc,'Elementos: ','☐ From  ☐ To  ☐ Subject (#estación)  ☐ Greeting  ☐ Cuerpo  ☐ Closing  ☐ Firma')
doc.add_paragraph('')

h(doc,'Criterio 2: Precisión Gramatical (6 puntos)',2)
tbl(doc,['Nivel','Pts','Descripción'],[
    ['Competente','5-6','IS/HAS/DEMONSTRATIVES correctos en ≥80%. Comparaciones correctas.'],
    ['Adecuado','3-4','Correctos en 60-79%. Errores menores.'],
    ['En desarrollo','1-2','Correctos en 40-59%. Confunde IS/HAS.'],
    ['No evidenciado','0','Correctos en <40% o ausentes.'],
],sc='2E75B6')
tbl(doc,['Estructura','Ejemplo','Eval'],[
    ['Verb To Be (is)','"The CPU is old."','☐ ✓ ☐ ✗'],
    ['Have/Has','"It has 4 GB of RAM."','☐ ✓ ☐ ✗'],
    ['Demonstratives','"This is an HDMI port."','☐ ✓ ☐ ✗'],
    ['Comparatives','"An SSD is faster than an HDD."','☐ ✓ ☐ ✗'],
    ['Request','"I need a new monitor."','☐ ✓ ☐ ✗'],
],sc='4472C4')
doc.add_page_break()

h(doc,'Criterio 3: Uso de Vocabulario Técnico (6 puntos)',2)
tbl(doc,['Nivel','Pts','Descripción'],[
    ['Competente','5-6','≥6 términos de los 20 key terms, en contexto correcto.'],
    ['Adecuado','3-4','4-5 términos correctamente.'],
    ['En desarrollo','1-2','2-3 términos o con errores.'],
    ['No evidenciado','0','0-1 términos o incorrectos.'],
],sc='2E75B6')
tbl(doc,['Término encontrado','¿Contexto correcto?'],[['1. _______________','☐ Sí ☐ No'],['2. _______________','☐ Sí ☐ No'],['3. _______________','☐ Sí ☐ No'],['4. _______________','☐ Sí ☐ No'],['5. _______________','☐ Sí ☐ No'],['6. _______________','☐ Sí ☐ No']],sc='4472C4')
doc.add_paragraph('')

h(doc,'Criterio 4: Claridad Comunicativa (4 puntos)',2)
tbl(doc,['Nivel','Pts','Descripción'],[
    ['Competente','4','IT Manager entendería problema + solicitud sin aclaraciones.'],
    ['Adecuado','3','Comprensible pero requiere 1 aclaración menor.'],
    ['En desarrollo','2','Parcialmente comprensible. Faltan detalles.'],
    ['Inicial','1','Difícil de entender.'],
    ['No evidenciado','0','Incomprensible o vacío.'],
],sc='2E75B6')
doc.add_paragraph('')

h(doc,'Resumen IE-02',2)
tbl(doc,['Criterio','Pts obtenidas','Pts máximas'],[
    ['1. Formato del email','','4'],['2. Precisión gramatical','','6'],
    ['3. Vocabulario técnico','','6'],['4. Claridad comunicativa','','4'],
    ['TOTAL','','20'],
],sc='1F3A5F')
st(doc,'Nivel: ☐ Competente (17-20)  ☐ Adecuado (12-16)  ☐ En desarrollo (7-11)  ☐ Inicial (0-6)')
bp(doc,'Comentario: ');line(doc,2)
bp(doc,'Instructor: ________________________  Firma: _____________')
doc.add_page_break()

# ═══ IE-03 ═══
h(doc,'IE-03: IT SUPPORT CALL — LISTA DE CHEQUEO DE SIMULACIÓN ORAL',1)
bp(doc,'Evidencia SENA: ','Desempeño')
bp(doc,'Técnica: ','Observación directa')
bp(doc,'Puntuación: ','20 puntos')
bp(doc,'Ubicación: ','Session 6 (The Help Desk)')
doc.add_paragraph('')

h(doc,'Datos del Aprendiz',2)
tbl(doc,['Campo','Dato'],[['Nombre',''],['Ficha',''],['Fecha',''],['Ronda observada','☐ 1  ☐ 2  ☐ 3'],['Rol','☐ Developer  ☐ IT Support'],['Compañero(a)','']],sc='4472C4')
doc.add_paragraph('')

h(doc,'Bloque A: Comunicación (8 puntos)',2)
tbl(doc,['#','Indicador','Logrado (2)','Parcial (1)','No logrado (0)'],[
    ['1','Inicia la conversación','Fórmula Skeleton: "IT Support, this is [name]"','Saluda sin fórmula','No inicia o solo español'],
    ['2','Describe el problema','≥2 problemas con IS/HAS correctos','1 problema o errores is/has','No describe o incomprensible'],
    ['3','Mantiene la conversación','Responde + hace preguntas ("What about...?")','Responde sin preguntas — flujo interrumpido','No responde o se detiene'],
    ['4','Cierra la conversación','Fórmula cierre + acuerdo claro','Cierra sin acuerdo o sin fórmula','Interrumpe abruptamente'],
],sc='2E75B6')
doc.add_paragraph('')

h(doc,'Bloque B: Competencia Lingüística (8 puntos)',2)
tbl(doc,['#','Indicador','Logrado (2)','Parcial (1)','No logrado (0)'],[
    ['5','Vocabulario técnico','≥3 términos en contexto correcto','1-2 términos o errores menores','Sin vocabulario técnico'],
    ['6','Estructuras gramaticales','IS/HAS/DEM correctos en ≥70%','Correctos en 30-69%','Correctos en <30%'],
    ['7','Pronunciación inteligible','Compañero entiende sin repetición','Entiende con 1-2 repeticiones','No entiende o traducción'],
    ['8','Frases del Skeleton','≥3 frases de forma natural','1-2 frases o lee mecánicamente','No usa frases'],
],sc='2E75B6')
doc.add_paragraph('')

h(doc,'Bloque C: Competencia Pragmática (4 puntos)',2)
tbl(doc,['#','Indicador','Logrado (2)','Parcial (1)','No logrado (0)'],[
    ['9','Negociación','Negocia NOW vs. ORDER. Acuerdo mutuo.','Menciona stock/order sin acuerdo','No negocia o acepta todo'],
    ['10','Transferencia','Aplica vocab + grammar + formato de la guía','Aplica parcialmente','Sin transferencia'],
],sc='2E75B6')
doc.add_page_break()

h(doc,'Resumen IE-03',2)
tbl(doc,['Bloque','Pts obtenidas','Pts máximas'],[
    ['A. Comunicación (1-4)','','8'],['B. Lingüística (5-8)','','8'],
    ['C. Pragmática (9-10)','','4'],['TOTAL','','20'],
],sc='1F3A5F')
st(doc,'Nivel: ☐ Competente (17-20)  ☐ Adecuado (12-16)  ☐ En desarrollo (7-11)  ☐ Inicial (0-6)')
doc.add_paragraph('')
bp(doc,'Frase destacada (literal):');st(doc,'"____________________________________________________________"')
bp(doc,'Fortaleza: ');line(doc)
bp(doc,'Área de mejora: ');line(doc)
bp(doc,'Recomendación: ');line(doc)
bp(doc,'Instructor: ________________________  Firma: _____________')
doc.add_page_break()

# ═══ TRIADA SUMMARY ═══
h(doc,'TRIADA DE EVIDENCIAS SENA — RESUMEN',1)
tbl(doc,['IE','Instrumento','Evidencia','Pts máx','Ubicación'],[
    ['IE-01','Cuestionario Técnico','Conocimiento','50','Session 7'],
    ['IE-02','Tech Request Email','Producto','20','Session 5'],
    ['IE-03','IT Support Call','Desempeño','20','Session 6'],
    ['','TOTAL','','90',''],
],sc='1F3A5F')

doc.add_paragraph('')
doc.add_paragraph('')
for txt in ['ADSO — GUÍA 1: The Hardware Specialist — Instrumentos de Evaluación SENA','Sistema de Prompts Maestros — LG Factory — FPI SENA — Bilingüismo','Instructor Sergio Cortés Perdomo · Marzo 2026']:
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(txt);r.font.size=Pt(9);r.font.name='Calibri';r.font.color.rgb=RGBColor(0x7F,0x8C,0x8D)

doc.save('/Users/Beppo/Projects/fpi-sena-factory/guides/ADSO-G1/ADSO — GUÍA 1 — Instrumentos de Evaluación SENA (IE-01 IE-02 IE-03).docx')
print('Done')
