#!/usr/bin/env python3
"""Generate Session 1 Build-Out as a Word document."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def set_first_row(table, headers, shade='1F3A5F'):
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, shade)

def make_table(doc, headers, rows, shade='1F3A5F'):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_first_row(table, headers, shade)
    for row_data in rows:
        row = table.add_row()
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(text))
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
    return table

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Calibri'
    return h

def add_bold_para(doc, label, value=''):
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    if value:
        run2 = p.add_run(value)
        run2.font.name = 'Calibri'
        run2.font.size = Pt(11)
    return p

def add_teacher_talk(doc, text):
    p = doc.add_paragraph()
    run = p.add_run('Teacher Talk: ')
    run.bold = True
    run.italic = True
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run2 = p.add_run(f'"{text}"')
    run2.italic = True
    run2.font.name = 'Calibri'
    run2.font.size = Pt(11)
    return p

def add_facilitation_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run('💡 ')
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.italic = True
    run2.font.name = 'Calibri'
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    return p

def add_checkpoint(doc, text):
    p = doc.add_paragraph()
    run = p.add_run('✓ Checkpoint: ')
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
    run2 = p.add_run(text)
    run2.font.name = 'Calibri'
    run2.font.size = Pt(11)
    return p

def add_icq(doc, question, expected):
    p = doc.add_paragraph()
    run = p.add_run('ICQ: ')
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run2 = p.add_run(f'"{question}"')
    run2.font.name = 'Calibri'
    run2.font.size = Pt(11)
    run3 = p.add_run(f' → Esperan: "{expected}"')
    run3.italic = True
    run3.font.name = 'Calibri'
    run3.font.size = Pt(11)
    run3.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    return p

# ════════════════════════════════════════
# TITLE
# ════════════════════════════════════════

for _ in range(4):
    doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('SESSION 1: THE WAKE-UP CALL')
run.bold = True
run.font.size = Pt(20)
run.font.name = 'Calibri'
run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('ADSO — GUÍA 1: The Hardware Specialist — Build-Out')
run.font.size = Pt(14)
run.font.name = 'Calibri'
run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

doc.add_paragraph('')

warn = doc.add_paragraph()
warn.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = warn.add_run('⚠️ Este documento es SOLO para el instructor. No distribuir a los aprendices.')
run.bold = True
run.font.size = Pt(12)
run.font.name = 'Calibri'
run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

doc.add_paragraph('')

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('Sistema de Prompts Maestros — LG Factory — FPI SENA — Bilingüismo')
run.font.size = Pt(10)
run.font.name = 'Calibri'
run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

meta2 = doc.add_paragraph()
meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta2.add_run('Instructor Sergio Cortés Perdomo · Marzo 2026')
run.font.size = Pt(10)
run.font.name = 'Calibri'
run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

doc.add_page_break()

# ════════════════════════════════════════
# SESSION HEADER
# ════════════════════════════════════════

add_heading_styled(doc, 'SESSION HEADER', level=1)

make_table(doc, ['Campo', 'Dato'], [
    ['Programa', 'Análisis y Desarrollo de Software (ADSO) — 228118'],
    ['Guía', 'Guía 1: The Hardware Specialist'],
    ['Session', '1: The Wake-Up Call'],
    ['Worksheets', 'PM-2.1 (The Spark) + PM-2.2 (The Gap Analysis)'],
    ['Duración', '180 minutos'],
    ['Habilidades foco', '— (sesión de diagnóstico y motivación)'],
    ['Habilidades soporte', '—'],
    ['Trabajo autónomo', 'Workbook Ch. 1 — My Hardware Profile (45 min)'],
    ['Siguiente sesión', 'Session 2: Read the Request (PM-2.3 — Reading)'],
])

doc.add_paragraph('')

# ════════════════════════════════════════
# MATERIALS CHECKLIST
# ════════════════════════════════════════

add_heading_styled(doc, 'MATERIALS CHECKLIST', level=1)

p = doc.add_paragraph('Marcar ANTES de entrar al aula:')
p.runs[0].italic = True

checklist = [
    'PM-2.1: The Spark — impresos (1 por estudiante + 2 extras)',
    'PM-2.2: The Gap Analysis — impresos (1 por estudiante + 2 extras)',
    'Canva slides 1-4 abiertas y listas en la computadora',
    'Proyector encendido y funcionando',
    'Tablero preparado según Board Plan',
    'Markers de colores (mínimo 2: negro + azul o rojo)',
    'Timer visible en pantalla o celular',
]
for item in checklist:
    p = doc.add_paragraph(f'☐  {item}')
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

doc.add_paragraph('')

# ════════════════════════════════════════
# BOARD PLAN
# ════════════════════════════════════════

add_heading_styled(doc, 'BOARD PLAN', level=1)

p = doc.add_paragraph('Preparar ANTES de que lleguen los aprendices:')
p.runs[0].italic = True

board_text = (
    'SESSION 1: THE WAKE-UP CALL\n\n'
    'Today\'s Question:\n'
    '"What is the most important part of a computer?"\n\n'
    '[LIVE ZONE 1 — vacío]              [LIVE ZONE 2 — vacío]\n'
    '(Se llena con respuestas del        (Se llena con los BLIND SPOTS\n'
    ' warm-up — Bloque A)                más comunes de la clase — Bloque C)\n\n'
    '                                         1. _________\n'
    '                                         2. _________\n'
    '                                         3. _________\n'
    '                                         4. _________\n'
    '                                         5. _________\n\n'
    'SURVIVAL WORDS (pre-escritas):\n'
    '• Crash = when a computer stops working suddenly\n'
    '• Deadline = the last day to finish something\n'
    '• Fix = to repair something'
)

p = doc.add_paragraph()
run = p.add_run(board_text)
run.font.name = 'Consolas'
run.font.size = Pt(9)

p2 = doc.add_paragraph()
run2 = p2.add_run('Durante la sesión: ')
run2.bold = True
run2.font.name = 'Calibri'
run2.font.size = Pt(10)
run3 = p2.add_run('El tablero izquierdo (LIVE ZONE 1) se llena con respuestas del warm-up. El tablero derecho (LIVE ZONE 2) se llena con los blind spots comunes en Bloque C. Las Survival Words permanecen visibles toda la sesión.')
run3.font.name = 'Calibri'
run3.font.size = Pt(10)

doc.add_page_break()

# ════════════════════════════════════════
# TIMELINE
# ════════════════════════════════════════

add_heading_styled(doc, 'MINUTE-BY-MINUTE TIMELINE', level=1)

make_table(doc,
    ['Tiempo', 'Dur.', 'Bloque', 'Actividad', 'Agrupación', 'Notas'],
    [
        ['0:00-0:03', '3 min', 'SET-UP', 'Warm-up: pregunta en tablero', 'Individual', 'Piensan en silencio'],
        ['0:03-0:05', '2 min', 'SET-UP', 'Warm-up: compartir con vecino', 'Pairs', 'Instructor circula'],
        ['0:05-0:10', '5 min', 'SET-UP', 'Warm-up: recoger respuestas', 'Plenary', 'Escribir en LIVE ZONE 1'],
        ['0:10-0:15', '5 min', 'SET-UP', 'Opening Script + Objective', 'Plenary', 'Presentar la guía'],
        ['0:15-0:20', '5 min', 'WHILE-A', 'Spark: leer escenario en voz alta', 'Plenary', 'Instructor lee, slide 2'],
        ['0:20-0:30', '10 min', 'WHILE-A', 'Spark: lectura individual', 'Individual', 'Silencio, worksheet PM-2.1'],
        ['0:30-0:35', '5 min', 'WHILE-A', 'Spark: presentar Pregunta Polémica', 'Plenary', 'Slide 3, tablero'],
        ['0:35-0:55', '20 min', 'WHILE-A', 'Spark: debate en grupos de 3', 'Groups of 3', 'Instructor circula'],
        ['0:55-1:05', '10 min', 'WHILE-A', 'Spark: conclusiones plenaria', 'Plenary', '3-4 grupos comparten'],
        ['1:05-1:15', '10 min', 'WHILE-A', 'Spark: Activity 1 (Personal Connection)', 'Individual', 'Rate 1-5'],
        ['1:15-1:25', '10 min', 'WHILE-A', 'Spark: Activity 3 (Need to Know)', 'Individual', 'Write 2 things'],
        ['1:25-1:40', '15 min', 'BREAK', 'BREAK — Stretch & Recharge', '—', ''],
        ['1:40-1:45', '5 min', 'WHILE-B', 'Transición: Gap Analysis intro', 'Plenary', ''],
        ['1:45-1:55', '10 min', 'WHILE-B', 'Gap: Activity 1 (Brainstorming)', 'Individual', 'Write English words'],
        ['1:55-2:10', '15 min', 'WHILE-B', 'Gap: Activity 2 (Blind Spots)', 'Individual', '5 situations'],
        ['2:10-2:25', '15 min', 'WHILE-B', 'Gap: comparar blind spots', 'Pairs', ''],
        ['2:25-2:40', '15 min', 'WHILE-B', 'Gap: Activity 3 (Learning Contract)', 'Individual', 'Escribir, firmar'],
        ['2:40-2:55', '15 min', 'WHILE-C', 'Síntesis: blind spots en tablero', 'Plenary', 'LIVE ZONE 2'],
        ['2:55-3:00', '5 min', 'WHILE-C', '✓ Checkpoint', 'Plenary', 'Thumbs up/down'],
        ['3:00-3:08', '8 min', 'WRAP-UP', 'Exit Ticket', 'Individual', '1 know + 1 don\'t know'],
        ['3:08-3:13', '5 min', 'WRAP-UP', 'Closing + Autónomo', 'Plenary', 'Workbook Ch. 1'],
        ['3:13-3:15', '2 min', 'WRAP-UP', 'Preview sesión siguiente', 'Plenary', ''],
    ],
)

p = doc.add_paragraph()
run = p.add_run('Total: 180 minutos ✓')
run.bold = True
run.font.name = 'Calibri'
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

doc.add_page_break()

# ════════════════════════════════════════
# SET-UP DETALLADO
# ════════════════════════════════════════

add_heading_styled(doc, 'SET-UP DETALLADO (15 min) — 0:00 a 0:15', level=2)

add_heading_styled(doc, 'Warm-up: "The Big Question" (10 min)', level=3)

add_bold_para(doc, 'Paso 1 — Pensar individualmente (3 min)')
p = doc.add_paragraph('Los aprendices entran al aula. El tablero ya muestra la pregunta: "What is the most important part of a computer?"')
p = doc.add_paragraph()
run = p.add_run('"Good morning everyone. Welcome to The Hardware Specialist. Look at the board. I want you to think — just think — about this question. No talking yet. You have 1 minute."')
run.italic = True
run.font.name = 'Calibri'
p2 = doc.add_paragraph('Timer: 1 minuto de silencio. Instructor espera.')

add_bold_para(doc, 'Paso 2 — Compartir con vecino (2 min)')
p = doc.add_paragraph()
run = p.add_run('"Now turn to the person next to you. Share your answer. You have 2 minutes. Go."')
run.italic = True
run.font.name = 'Calibri'
p2 = doc.add_paragraph('Timer: 2 minutos. Instructor circula, escucha las respuestas.')

add_bold_para(doc, 'Paso 3 — Recoger respuestas en plenaria (5 min)')
p = doc.add_paragraph()
run = p.add_run('"Ok, stop. Let\'s hear some answers. Who wants to share?"')
run.italic = True
run.font.name = 'Calibri'
doc.add_paragraph('Recoger 5-6 respuestas. Escribir cada una en LIVE ZONE 1 del tablero. No corregir ni evaluar — solo registrar.')
doc.add_paragraph('Ejemplos típicos: "CPU", "RAM", "the screen", "the brain", "the keyboard".')

add_facilitation_note(doc, 'Si la clase está muy callada y nadie quiere compartir, diga: "Ok, I\'ll start. I think the most important part is... the power supply. Because without power, NOTHING works. Agree or disagree?" Esto rompe el hielo.')

doc.add_paragraph('')
add_heading_styled(doc, 'Teacher Talk — Opening Script (3 min)', level=3)

p = doc.add_paragraph()
run = p.add_run('"Good morning everyone. My name is [nombre del instructor]. Welcome to Guía 1: The Hardware Specialist. For the next 8 sessions, you are going to learn the English of hardware — the parts inside your computer, how to describe them, how to request them, and how to talk about them with people from other countries. Today is Day 1. Before we start learning, I want to know — what do YOU already know? And what do you NOT know? Today is about discovery."')
run.italic = True
run.font.name = 'Calibri'

doc.add_paragraph('')
add_heading_styled(doc, 'Objective (escrito en tablero)', level=3)

p = doc.add_paragraph()
run = p.add_run('"Today you will: discover what you already know about hardware — and what you need to learn in this guide."')
run.italic = True
run.font.name = 'Calibri'

doc.add_paragraph('')
add_icq(doc, 'Are we starting with a test today?', 'No')
add_icq(doc, 'What are we doing today?', 'Discovering what we know / finding out what we don\'t know')

doc.add_page_break()

# ════════════════════════════════════════
# WHILE DETALLADO — BLOQUE A
# ════════════════════════════════════════

add_heading_styled(doc, 'WHILE DETALLADO (150 min) — 0:15 a 2:55', level=2)

add_heading_styled(doc, 'BLOQUE A — THE SPARK (70 min) — 0:15 a 1:25', level=3)

add_bold_para(doc, 'Worksheet: ', 'PM-2.1 (The Spark)')
add_bold_para(doc, 'Agrupación: ', 'Plenary → Individual → Groups of 3 → Individual')
add_bold_para(doc, 'Objetivo del bloque: ', 'El aprendiz identifica sus propios vacíos de conocimiento y se motiva para llenarlos durante la guía.')

doc.add_paragraph('')
add_heading_styled(doc, 'Spark: Lectura del escenario (15 min) — 0:15 a 0:30', level=3)

add_bold_para(doc, 'Instrucciones paso a paso:')
doc.add_paragraph('1. Proyectar slide 2 (imagen del escenario).')
doc.add_paragraph('2. Instructor lee el escenario en voz alta, con dramatismo moderado:')

p = doc.add_paragraph()
run = p.add_run('"It\'s Friday night. Your team has a big software project. The deadline is Monday. Suddenly, one computer in the office crashes. The screen goes black. There is smoke coming from the back of the machine. Nobody in the team knows what happened. The boss says: \'We need this computer working NOW. The project is inside.\' But nobody can open the computer. Nobody knows the name of the parts inside. Nobody can call technical support because... the support manual is in English."')
run.italic = True
run.font.name = 'Calibri'

doc.add_paragraph('3. Después de leer: "Now — read it again, silently, by yourself. Underline any words you don\'t understand. You have 5 minutes."')
doc.add_paragraph('4. Timer: 5 minutos de lectura silenciosa.')
doc.add_paragraph('5. "Do you have any words you don\'t understand? Ask me." — Responder 2-3 preguntas de vocabulario máximo.')

add_facilitation_note(doc, 'Los aprendices A1 pueden preguntar por "smoke", "deadline", "crashes". Defina con gestos o palabras simples: "Smoke = humo" (señalar), "Crash = when a computer dies suddenly" (gesto de apagón).')

doc.add_paragraph('')
add_heading_styled(doc, 'Spark: Pregunta Polémica + Debate (25 min) — 0:30 a 0:55', level=3)

add_bold_para(doc, 'Instrucciones paso a paso:')
doc.add_paragraph('1. Proyectar slide 3 con la Pregunta Polémica. Escribir también en el tablero:')
p = doc.add_paragraph()
run = p.add_run('"If you don\'t know the NAMES of the parts inside a computer in English... can you really call yourself a software developer?"')
run.bold = True
run.font.name = 'Calibri'

doc.add_paragraph('2. "This is the Controversial Question. There are NO wrong answers. I want your opinion. Form groups of 3 — now."')
doc.add_paragraph('3. Transición: 1 minuto para formar grupos.')
doc.add_paragraph('4. "Discuss this question in your group. Use these phrases" (señalar phrases del worksheet en slide 3):')
doc.add_paragraph('   - "I think... YES / NO because..."')
doc.add_paragraph('   - "In my opinion..."')
doc.add_paragraph('   - "I agree / I disagree because..."')
doc.add_paragraph('5. "You have 10 minutes. At the end, your group writes ONE conclusion on your worksheet. Go."')
doc.add_paragraph('6. Timer: 10 minutos. Instructor circula por los grupos.')

add_facilitation_note(doc, 'Mientras circula, NO entre en el debate como participante. Solo observe y tome nota de qué grupos tienen conclusiones claras. Si un grupo está en español, está bien — la actividad es de reflexión, no de producción lingüística.')

doc.add_paragraph('7. "Ok, stop. Let\'s hear from some groups. Group 3 — what did you decide?"')
doc.add_paragraph('8. 3-4 grupos comparten su conclusión (1 minuto por grupo). Instructor escribe palabras clave en el tablero.')
doc.add_paragraph('9. Después de escuchar: "Interesting. Some of you said YES, some said NO. The truth is — this guide exists because this is a real problem. Developers work with international teams. The manuals are in English. The support lines are in English. By the end of this guide, you will be able to make that call."')

doc.add_paragraph('')
add_heading_styled(doc, 'Spark: Activity 1 — Personal Connection (10 min) — 1:05 a 1:15', level=3)

add_bold_para(doc, 'Instrucciones paso a paso:')
doc.add_paragraph('1. "Now — open your PM-2.1 worksheet. Look at Activity 1. You see 3 questions with a scale of 1 to 5. Answer them HONESTLY. This is for YOU, not for me. You have 3 minutes."')
doc.add_paragraph('2. Timer: 3 minutos. Individual.')
doc.add_paragraph('3. "Done? Now share with the person next to you. Compare your numbers. Are they similar or different? You have 3 minutes."')
doc.add_paragraph('4. Timer: 3 minutos. Pairs.')
doc.add_paragraph('5. "Anyone have a 5 on the first question? And a 1 on the third?" — 2-3 voluntarios comparten.')

add_icq(doc, 'Are you writing sentences or just circling numbers?', 'Circling numbers')
add_icq(doc, 'Do you share with the whole class or just your partner?', 'Just my partner')

doc.add_paragraph('')
add_heading_styled(doc, 'Spark: Activity 3 — The "Need to Know" (10 min) — 1:15 a 1:25', level=3)

add_bold_para(doc, 'Instrucciones paso a paso:')
doc.add_paragraph('1. "Look at Activity 3 on your worksheet. Imagine you are the person in the scenario. The computer crashed. You need to call technical support in the United States. What do you need to learn TODAY to make that call? Write 2 things."')
doc.add_paragraph('2. "You have 5 minutes. In English or Spanish — whatever helps you think."')
doc.add_paragraph('3. Timer: 5 minutos. Individual.')
doc.add_paragraph('4. "Quick share — turn to your partner. Read your 2 things to each other. 2 minutes."')
doc.add_paragraph('5. Timer: 2 minutos. Pairs.')

add_facilitation_note(doc, 'Esta actividad es metacognitiva — los aprendices no necesitan escribir en inglés perfecto. Si escriben en español, está bien. El objetivo es que IDENTIFIQUEN sus vacíos, no que los formulen perfectamente.')

add_checkpoint(doc, '"Raise your hand if you wrote \'hardware vocabulary\' as one of your needs." — Si >70% lo escribió, confirma que la guía lo cubrirá.')

p = doc.add_paragraph()
run = p.add_run('Transition → BREAK: ')
run.bold = True
run.font.name = 'Calibri'
run2 = p.add_run('"Good work. You\'ve identified what you need. Now — take a break. 15 minutes. Come back ready to find out exactly what your blind spots are."')
run2.italic = True
run2.font.name = 'Calibri'

doc.add_page_break()

# ════════════════════════════════════════
# BREAK
# ════════════════════════════════════════

add_heading_styled(doc, 'BREAK — Stretch & Recharge (15 min) — 1:25 a 1:40', level=2)

p = doc.add_paragraph('Instructor toma nota mental de:')
items = [
    'Qué grupos tuvieron debates más interesantes (para referenciar en Bloque C)',
    'Qué aprendices parecen tener más o menos conocimiento (para agrupación posterior)',
    'Qué "Need to Know" aparecieron con más frecuencia',
]
for item in items:
    p = doc.add_paragraph(item, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

doc.add_page_break()

# ════════════════════════════════════════
# BLOQUE B — GAP ANALYSIS
# ════════════════════════════════════════

add_heading_styled(doc, 'BLOQUE B — THE GAP ANALYSIS (55 min) — 1:40 a 2:35', level=3)

add_bold_para(doc, 'Worksheet: ', 'PM-2.2 (The Gap Analysis)')
add_bold_para(doc, 'Agrupación: ', 'Plenary → Individual → Pairs → Individual')
add_bold_para(doc, 'Objetivo del bloque: ', 'El aprendiz diagnostica sus saberes previos, identifica sus puntos ciegos y crea un contrato de aprendizaje personalizado.')

doc.add_paragraph('')
add_heading_styled(doc, 'Transición + Gap Analysis intro (5 min) — 1:40 a 1:45', level=3)

p = doc.add_paragraph()
run = p.add_run('"Ok, welcome back. In the Spark, we talked about the big picture — why hardware English matters. Now we get SPECIFIC. I\'m going to show you the inside of a computer. You tell me: can you name the parts? Let\'s find out your blind spots."')
run.italic = True
run.font.name = 'Calibri'

doc.add_paragraph('1. Proyectar slide 4 (imagen ASCII del interior del computador de PM-2.2).')
doc.add_paragraph('2. "Look at this image. Components A, B, C, D, E. Can you name them in English? Don\'t answer yet — just look. This is what we\'re going to work with."')

doc.add_paragraph('')
add_heading_styled(doc, 'Gap Analysis: Activity 1 — Brainstorming (10 min) — 1:45 a 1:55', level=3)

add_bold_para(doc, 'Instrucciones paso a paso:')
doc.add_paragraph('1. "Open your PM-2.2 worksheet. Activity 1 — \'What I Know\'. Write ALL the computer words you already know in English. There is no minimum. Even 1 word is perfect. You have 5 minutes."')
doc.add_paragraph('2. Timer: 5 minutos. Individual. Silencio.')
doc.add_paragraph('3. "Done? Count your words. How many do you have? Tell your partner." — 30 segundos.')
doc.add_paragraph('4. "Who has more than 5 words? More than 10?" — Puños arriba. Celebrar los que tienen más.')

add_facilitation_note(doc, 'Algunos aprendices se bloquean porque piensan que tienen que saber "mucho". Refuerce: "Even 1 word is a starting point. We go from there." Si alguien tiene 0 palabras, dígale en voz baja: "Do you know \'computer\'? There — you have 1."')

doc.add_paragraph('')
add_heading_styled(doc, 'Gap Analysis: Activity 2 — Blind Spots (15 min) — 1:55 a 2:10', level=3)

add_bold_para(doc, 'Instrucciones paso a paso:')
doc.add_paragraph('1. "Now look at Activity 2 — \'What I Don\'t Know\'. There are 5 situations. Can you say them in English? If YES — write it. If NO — leave it blank or put a ❓. Be honest. No one is grading this."')
doc.add_paragraph('2. "You have 5 minutes. Go."')
doc.add_paragraph('3. Timer: 5 minutos. Individual. Silencio.')

add_bold_para(doc, 'Answer Key (para referencia del instructor — NO revelar a los aprendices todavía):')
doc.add_paragraph('Las respuestas correctas se revelan a lo largo de la guía. En esta sesión, el objetivo es que los aprendices IDENTIFIQUEN que no saben, no que aprendan las respuestas.')

make_table(doc,
    ['#', 'Situación', 'Respuesta correcta'],
    [
        ['1', 'Component that processes ALL information', 'CPU / Processor'],
        ['2', 'TEMPORARY memory of a computer', 'RAM'],
        ['3', '"This computer TIENE 16 gigas de RAM"', '"This computer has 16 gigabytes of RAM"'],
        ['4', '"The monitor IS a device for output"', '"The monitor is an output device"'],
        ['5', 'Difference between SSD and HDD', 'SSD is faster / solid state; HDD is mechanical / magnetic'],
    ],
    shade='2E75B6',
)

doc.add_paragraph('')
doc.add_paragraph('4. "Now — turn to your partner. Compare your Blind Spots. Which ones are blank for BOTH of you? You have 5 minutes."')
doc.add_paragraph('5. Timer: 5 minutos. Pairs.')
doc.add_paragraph('6. "Ok. How many pairs have at least 3 blank spots?" — Contar manos.')
doc.add_paragraph('7. "Good. These are your blind spots. And these are what we will fix in this guide."')

add_facilitation_note(doc, 'Si un par tiene TODAS las respuestas (lo cual es raro pero posible en aprendices con más nivel), dígales: "Great — you know a lot already. Your job in this guide is to help your classmates AND to go deeper. Can you explain the difference between SSD and HDD in 3 sentences? Write that."')

doc.add_paragraph('')
add_heading_styled(doc, 'Gap Analysis: Activity 3 — Learning Contract (15 min) — 2:10 a 2:25', level=3)

add_bold_para(doc, 'Instrucciones paso a paso:')
doc.add_paragraph('1. "Now — Activity 3. \'What I Need to Learn\'. Based on your Blind Spots, write YOUR Learning Contract. Complete the sentence: \'In this guide, I need to learn:\' — and write 3 or 4 things."')
doc.add_paragraph('2. "This is YOUR contract with yourself. You will sign it and date it."')
doc.add_paragraph('3. "You have 8 minutes. Go."')
doc.add_paragraph('4. Timer: 8 minutos. Individual.')
doc.add_paragraph('5. "Done? Read your contract to your partner. Do you have the same needs? You have 3 minutes."')
doc.add_paragraph('6. Timer: 3 minutos. Pairs.')
doc.add_paragraph('7. "Sign your contract. Date it. Keep it in your folder — you\'ll look at it again in Session 8."')

add_icq(doc, 'How many things do you write in the contract?', '3 or 4')
add_icq(doc, 'Do you sign it?', 'Yes')
add_icq(doc, 'When do you look at it again?', 'Session 8 / the last session')

add_facilitation_note(doc, 'El Learning Contract es un compromiso psicológico, no un documento legal. Refuerce la importancia: "This is a promise to yourself. When you finish this guide, you\'ll check — did I keep my promise?"')

doc.add_page_break()

# ════════════════════════════════════════
# BLOQUE C — SÍNTESIS
# ════════════════════════════════════════

add_heading_styled(doc, 'BLOQUE C — SÍNTESIS (15 min) — 2:40 a 2:55', level=3)

add_bold_para(doc, 'Agrupación: ', 'Pairs → Plenary')
add_bold_para(doc, 'Objetivo del bloque: ', 'Los aprendices ven sus vacíos colectivizados en el tablero, lo que normaliza el punto de partida y establece los objetivos de la guía.')

doc.add_paragraph('')
add_heading_styled(doc, 'Colectar Blind Spots en tablero (10 min) — 2:40 a 2:50', level=3)

add_bold_para(doc, 'Instrucciones paso a paso:')
doc.add_paragraph('1. "Ok, everyone — eyes on me. I want to know: which Blind Spots were the MOST common in this class. Raise your hand if you could NOT say \'CPU\' in English." — Contar y escribir el número en LIVE ZONE 2.')
doc.add_paragraph('2. Repetir para cada una de las 5 situaciones. Escribir los resultados en el tablero:')

result_text = (
    'BLIND SPOTS — CLASS RESULTS:\n'
    '1. CPU name:        [X] students couldn\'t say it\n'
    '2. RAM name:        [X] students couldn\'t say it\n'
    '3. "has" sentence:  [X] students couldn\'t say it\n'
    '4. "is" sentence:   [X] students couldn\'t say it\n'
    '5. SSD vs HDD:      [X] students couldn\'t say it'
)
p = doc.add_paragraph()
run = p.add_run(result_text)
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_paragraph('3. "Look at the board. These are YOUR blind spots — the things most of the class doesn\'t know YET. By Session 8, every single one of you will be able to do ALL of these things. That\'s the goal."')

add_facilitation_note(doc, 'Si la clase tiene niveles muy dispares, diga: "Some of you know more than others. That\'s normal. In this class, the people who know more will help the people who are learning. We\'re a team."')

doc.add_paragraph('')
add_heading_styled(doc, '✓ Checkpoint: Verificación (5 min) — 2:50 a 2:55', level=3)

p = doc.add_paragraph()
run = p.add_run('"Before we close — thumbs up if you understand what this guide is about. Thumbs sideways if you\'re not sure. Thumbs down if you\'re confused."')
run.italic = True
run.font.name = 'Calibri'

doc.add_paragraph('Si >20% thumbs sideways/down: "Ok, let me explain one more time. This guide teaches you the English of computer hardware. By the end, you can read emails about hardware, listen to phone calls, write requests, and have conversations. Today was about finding out where you start. Next session — we start learning."')

p = doc.add_paragraph()
run = p.add_run('Transition → WRAP-UP: ')
run.bold = True
run.font.name = 'Calibri'
run2 = p.add_run('"Good. Now let\'s close."')
run2.italic = True
run2.font.name = 'Calibri'

doc.add_page_break()

# ════════════════════════════════════════
# WRAP-UP
# ════════════════════════════════════════

add_heading_styled(doc, 'WRAP-UP DETALLADO (15 min) — 2:55 a 3:10', level=2)

add_heading_styled(doc, 'Exit Ticket (8 min) — 2:55 a 3:03', level=3)

add_bold_para(doc, 'Instrucciones:')
doc.add_paragraph('1. "On a clean piece of paper — or at the bottom of your worksheet — write TWO things:"')
doc.add_paragraph('2. "Thing 1: ONE thing you are SURE you know about hardware in English."')
doc.add_paragraph('3. "Thing 2: ONE thing you DEFINITELY DON\'T know."')
doc.add_paragraph('4. "You have 3 minutes. Go."')
doc.add_paragraph('5. Timer: 3 minutos. Individual.')
doc.add_paragraph('6. "Hand your paper to me on the way out. Or keep it — your choice."')

add_bold_para(doc, 'Criterio de éxito: ', 'El aprendiz escribe algo en ambos espacios (no importa si es en inglés o español, no importa si la respuesta es correcta — el objetivo es reflexión).')

add_facilitation_note(doc, 'NO revisar los exit tickets en clase. Lea después para entender el punto de partida de cada aprendiz. Use esta información para identificar quién necesita más apoyo en sesiones siguientes.')

doc.add_paragraph('')
add_heading_styled(doc, 'Teacher Talk — Closing Script (5 min) — 3:03 a 3:08', level=3)

p = doc.add_paragraph()
run = p.add_run('"Ok, let\'s close. Today you discovered what you know — and what you don\'t know. That\'s the first step. Nobody starts from zero — everyone has something. Now we build on it.\n\nYour homework: Workbook Chapter 1 — My Hardware Profile. Draw your REAL computer — the one at home, or the one here at SENA. Label 5 parts in English. Even if you\'re not sure of the name — try. Write one sentence about each part. You have 45 minutes. We review it next session.\n\nNext session — we READ. A real tech request email from a developer named Carlos. He has a problem with his workstation. Your job: understand the problem and decide what to do. See you next time."')
run.italic = True
run.font.name = 'Calibri'

add_icq(doc, 'What do you draw?', 'My real computer')
add_icq(doc, 'How many parts do you label?', '5')
add_icq(doc, 'Do the names have to be perfect?', 'No, try')

doc.add_paragraph('')
add_heading_styled(doc, 'Preview — 3:08 a 3:10', level=3)

p = doc.add_paragraph()
run = p.add_run('"Next session: Read the Request. Carlos\'s email. Hardware problems. Your first real English reading in this guide. Come ready."')
run.italic = True
run.font.name = 'Calibri'

doc.add_page_break()

# ════════════════════════════════════════
# ANSWER KEY
# ════════════════════════════════════════

add_heading_styled(doc, 'ANSWER KEY CONSOLIDADO', level=1)

add_heading_styled(doc, 'PM-2.1 — The Spark', level=2)

make_table(doc,
    ['Actividad', 'Ítem', 'Respuesta Correcta', 'Notas'],
    [
        ['Activity 1', 'Question 1', 'Escala personal (1-5)', 'No hay respuesta correcta — diagnóstico motivacional'],
        ['Activity 1', 'Question 2', 'Escala personal (1-5)', 'Idem'],
        ['Activity 1', 'Question 3', 'Escala personal (1-5)', 'Idem'],
        ['Activity 2', 'Conclusión del grupo', 'Abierta', 'No hay respuesta correcta — ejercicio de opinión'],
        ['Activity 3', 'Need to Know #1', 'Abierta', 'Ej: "hardware vocabulary", "how to describe parts"'],
        ['Activity 3', 'Need to Know #2', 'Abierta', 'Ej: "phone conversation", "email format"'],
    ],
    shade='2E75B6',
)

doc.add_paragraph('')
add_heading_styled(doc, 'PM-2.2 — The Gap Analysis', level=2)

make_table(doc,
    ['Actividad', 'Ítem', 'Respuesta Correcta', 'Alternativas aceptables'],
    [
        ['Diagrama', 'A', 'CPU / Processor', '—'],
        ['Diagrama', 'B', 'RAM / Memory', '—'],
        ['Diagrama', 'C', 'GPU / Graphics Card', 'Video Card'],
        ['Diagrama', 'D', 'Motherboard', 'Mainboard'],
        ['Diagrama', 'E', 'PSU / Power Supply', 'Power Supply Unit'],
        ['Activity 1', 'Brainstorming', 'Lista personal', 'No hay respuesta correcta — diagnóstico'],
        ['Activity 2', '#1 (processes ALL info)', 'CPU / Processor', '—'],
        ['Activity 2', '#2 (TEMPORARY memory)', 'RAM', 'Random Access Memory'],
        ['Activity 2', '#3 ("tiene 16 gigas")', '"This computer has 16 gigabytes of RAM"', '"It has 16 GB of RAM"'],
        ['Activity 2', '#4 ("monitor IS")', '"The monitor is an output device"', '"The monitor is a device for output"'],
        ['Activity 2', '#5 (SSD vs HDD)', 'SSD is faster; HDD is mechanical', 'Cualquier comparación válida'],
        ['Activity 3', 'Learning Contract', 'Lista personal (3-4 ítems)', 'No hay respuesta correcta — compromiso personal'],
    ],
    shade='2E75B6',
)

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Nota sobre el Exit Ticket: ')
run.bold = True
run.font.name = 'Calibri'
run2 = p.add_run('No tiene answer key. El objetivo es reflexión, no evaluación. El instructor lee después para conocer el punto de partida.')
run2.font.name = 'Calibri'

doc.add_page_break()

# ════════════════════════════════════════
# DIFFERENTIATION
# ════════════════════════════════════════

add_heading_styled(doc, 'DIFFERENTIATION NOTES', level=1)

add_heading_styled(doc, 'Fast Finishers', level=2)

ff = [
    ('Activity 1 (Brainstorming)', 'Si termina con >10 palabras en 3 minutos: "Can you put your words into categories? Input, Output, Storage, Internal. Write the category next to each word."'),
    ('Activity 2 (Blind Spots)', 'Si completa las 5 situaciones fácilmente: "For each one you KNOW — write the full sentence, not just the word."'),
    ('Activity 3 (Learning Contract)', 'Si termina antes: "Write 2 MORE things you want to learn that are NOT on the worksheet."'),
    ('Peer tutor role', '"Help your partner with Activity 2. Don\'t give the answer — ask questions to help them figure it out."'),
]
for title, desc in ff:
    add_bold_para(doc, f'{title}: ', desc)

doc.add_paragraph('')
add_heading_styled(doc, 'More Support Needed', level=2)

ms = [
    ('Activity 1 (Brainstorming)', 'Si tiene 0-1 palabras después de 2 minutos: "Open your phone. Search \'parts of a computer\' in images. Look at the labels. Write 3 words you see. It\'s ok."'),
    ('Activity 2 (Blind Spots)', 'Permitir uso de diccionario (celular o físico) para las 5 situaciones. No es trampa — el objetivo es que IDENTIFIQUEN los vacíos.'),
    ('Activity 3 (Learning Contract)', 'Proporcionar sentence starter: "In this guide, I need to learn: 1. The name of _________ in English. 2. How to say _________. 3. _________."'),
    ('Exit Ticket', 'Si no puede escribir en inglés, aceptar español. El objetivo es la reflexión, no la producción.'),
]
for title, desc in ms:
    add_bold_para(doc, f'{title}: ', desc)

doc.add_paragraph('')
add_heading_styled(doc, 'Para clases muy grandes (>25 aprendices)', level=2)

add_bold_para(doc, 'Debate (Spark): ', 'En lugar de 1 plenaria, usar 2 rondas: grupos de 3 comparten, luego cada grupo de 3 se junta con otro grupo de 3. Después 2-3 grupos grandes comparten con la plenaria.')
add_bold_para(doc, 'Blind Spots comparison: ', 'En lugar de parejas, usar grupos de 4. Es más rápido y genera más datos para la síntesis colectiva.')

doc.add_page_break()

# ════════════════════════════════════════
# SELF-CHECK
# ════════════════════════════════════════

add_heading_styled(doc, 'INSTRUCTOR SELF-CHECK', level=1)

p = doc.add_paragraph('Después de la sesión, responda estas 5 preguntas:')
p.runs[0].italic = True

doc.add_paragraph('')
questions = [
    '¿Todos los aprendices completaron el Exit Ticket? ¿Cuántos lo hicieron correctamente?',
    '¿Hubo algún momento donde la mayoría parecía perdida? ¿En qué actividad?',
    '¿Los tiempos se cumplieron o hubo desfases? ¿Dónde?',
    '¿Qué errores comunes o vacíos frecuentes observé que debo retomar en la siguiente sesión?',
    '¿El trabajo autónomo asignado (Workbook Ch. 1) es realista para mis aprendices? ¿Necesito ajustar?',
]
for i, q in enumerate(questions, 1):
    add_bold_para(doc, f'{i}. {q}')
    p = doc.add_paragraph('Respuesta: _______________________________________________')
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
    doc.add_paragraph('')

# ════════════════════════════════════════
# FOOTER
# ════════════════════════════════════════

doc.add_paragraph('')
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('SESSION 1: THE WAKE-UP CALL — BUILD-OUT')
run.font.size = Pt(9)
run.font.name = 'Calibri'
run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

footer2 = doc.add_paragraph()
footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer2.add_run('ADSO — GUÍA 1: The Hardware Specialist')
run.font.size = Pt(9)
run.font.name = 'Calibri'
run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

footer3 = doc.add_paragraph()
footer3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer3.add_run('Sistema de Prompts Maestros — LG Factory — FPI SENA — Bilingüismo')
run.font.size = Pt(9)
run.font.name = 'Calibri'
run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

footer4 = doc.add_paragraph()
footer4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer4.add_run('Instructor Sergio Cortés Perdomo · Marzo 2026')
run.font.size = Pt(9)
run.font.name = 'Calibri'
run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

# ════════════════════════════════════════
# SAVE
# ════════════════════════════════════════

output_path = '/Users/Beppo/Projects/fpi-sena-factory/guides/ADSO-G1/ADSO — GUÍA 1 — Session 1 — The Wake-Up Call — Build-Out.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
