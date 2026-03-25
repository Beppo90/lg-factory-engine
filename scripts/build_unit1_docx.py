import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE = "/Users/Beppo/Projects/fpi-sena-factory/guides/MARITIME-G1"

FILES = [
    ("PM-1.2 — Unidad 1 — Ship Overview — Scope & Sequence.md", "PM-1.2 — SCOPE & SEQUENCE + CURACIÓN"),
    ("U1 — PM-2.1 — The Spark — Worksheet.md", "PM-2.1 — THE SPARK"),
    ("U1 — PM-2.2 — The Gap Analysis — Worksheet.md", "PM-2.2 — THE GAP ANALYSIS"),
    ("U1 — PM-2.3 — The Master Anchor — Worksheet.md", "PM-2.3 — THE MASTER ANCHOR"),
    ("U1 — PM-2.4 — Writing — Worksheet.md", "PM-2.4 — WRITING"),
    ("U1 — PM-2.5 — Literacy and Vocabulary — Worksheet.md", "PM-2.5 — LITERACY & VOCABULARY"),
    ("U1 — PM-2.6 — The Auditory Anchor — Worksheet.md", "PM-2.6 — THE AUDITORY ANCHOR"),
    ("U1 — PM-2.7 — Pronunciation — Worksheet.md", "PM-2.7 — PRONUNCIATION"),
    ("U1 — PM-2.8 — Speaking — Worksheet.md", "PM-2.8 — SPEAKING"),
    ("U1 — PM-2.9 — Language Functions — Transversal.md", "PM-2.9 — LANGUAGE FUNCTIONS (TRANSVERSAL)"),
    ("U1 — PM-2.10 — Grammar — Worksheet.md", "PM-2.10 — GRAMMAR"),
]

def add_page_break(doc):
    doc.add_page_break()

def parse_markdown_line(doc, line, styles):
    line = line.rstrip('\n')

    # Skip horizontal rules
    if line.strip() == '---':
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run('─' * 60)
        run.font.size = Pt(6)
        run.font.color.rgb = RGBColor(180, 180, 180)
        return

    # Skip empty lines (add small space)
    if not line.strip():
        return

    # Headings
    if line.startswith('# '):
        h = line[2:].strip()
        p = doc.add_heading(h, level=1)
        for run in p.runs:
            run.font.color.rgb = RGBColor(31, 58, 95)
        return
    if line.startswith('## '):
        h = line[3:].strip()
        p = doc.add_heading(h, level=2)
        for run in p.runs:
            run.font.color.rgb = RGBColor(31, 58, 95)
        return
    if line.startswith('### '):
        h = line[4:].strip()
        p = doc.add_heading(h, level=3)
        for run in p.runs:
            run.font.color.rgb = RGBColor(68, 114, 196)
        return
    if line.startswith('#### '):
        h = line[5:].strip()
        p = doc.add_heading(h, level=4)
        return

    # Code blocks (lines starting with ```)
    if line.strip().startswith('```'):
        return

    # Table rows (lines starting with |)
    if line.strip().startswith('|') and '|' in line.strip()[1:]:
        cells = [c.strip() for c in line.strip().split('|')[1:-1]]
        # Skip separator rows
        if all(set(c) <= set('- :') for c in cells):
            return
        return ('table_row', cells)

    # Blockquotes
    if line.startswith('> '):
        text = line[2:].strip()
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.5)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.font.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(80, 80, 80)
        return

    # Bullet points
    if line.startswith('- ') or line.startswith('* '):
        text = line[2:].strip()
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        add_formatted_text(p, text)
        return

    # Numbered items
    import re
    num_match = re.match(r'^(\d+)\.\s', line)
    if num_match:
        text = line[num_match.end():].strip()
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        add_formatted_text(p, text)
        return

    # Regular paragraph
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    add_formatted_text(p, line)

def add_formatted_text(paragraph, text):
    """Handle basic markdown bold and inline code."""
    import re
    parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(100, 50, 50)
        else:
            paragraph.add_run(part)

def flush_table(doc, table_rows):
    """Create a table from collected rows."""
    if not table_rows:
        return
    max_cols = max(len(r) for r in table_rows)
    table = doc.add_table(rows=len(table_rows), cols=max_cols)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(table_rows):
        for j, cell_text in enumerate(row_data):
            if j < max_cols:
                cell = table.cell(i, j)
                cell.text = ''
                p = cell.paragraphs[0]
                add_formatted_text(p, cell_text)
                for run in p.runs:
                    run.font.size = Pt(9)
                    if i == 0:
                        run.bold = True

    doc.add_paragraph()  # spacing after table

def build_docx():
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ---- COVER PAGE ----
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('INGLÉS MARÍTIMO Y PORTUARIO')
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(31, 58, 95)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Unidad 1: Ship Overview')
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(68, 114, 196)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Guía de Aprendizaje Completa')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(100, 100, 100)

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('FPI SENA — Bilingüismo')
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Caribbean Maritime Lines (CML) — Cartagena, Colombia')
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Nivel CEFR: A1.1 — A1.2')
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Intensidad: 24h directa + 6h autónoma')
    run.font.size = Pt(12)

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Grammar Targets: Demonstratives (this/that/these/those) + Verb To Be')
    run.font.size = Pt(11)
    run.font.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('20 Key Terms: bridge, engine room, deck, hull, bow, stern, port side, starboard side, container ship, bulk carrier, tanker, lifeboat, life jacket, fire extinguisher, gangway, mooring lines, hatch, hold, superstructure, propeller')
    run.font.size = Pt(9)
    run.font.italic = True

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Instructor Sergio Cortés Perdomo · Marzo 2026')
    run.font.size = Pt(11)

    # ---- TABLE OF CONTENTS (manual) ----
    add_page_break(doc)
    doc.add_heading('CONTENIDO', level=1)

    toc_items = [
        "PM-1.2 — Scope & Sequence + Curación de Material Auténtico",
        "PM-2.1 — The Spark — Reflexión Inicial (Worksheet)",
        "PM-2.2 — The Gap Analysis — Contextualización (Worksheet)",
        "PM-2.3 — The Master Anchor — Reading Comprehension (Worksheet)",
        "PM-2.4 — Writing — Task-Based Production (Worksheet)",
        "PM-2.5 — Literacy & Vocabulary — Key Terms (Worksheet)",
        "PM-2.6 — The Auditory Anchor — Listening Comprehension (Worksheet)",
        "PM-2.7 — Pronunciation — Speaking Skills (Worksheet)",
        "PM-2.8 — Speaking — Production & Simulation (Worksheet)",
        "PM-2.9 — Language Functions — Communicative Competence (Transversal)",
        "PM-2.10 — Grammar — Structure Use (Worksheet)",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item, style='List Number')
        p.paragraph_format.space_after = Pt(4)

    # ---- CONTENT SECTIONS ----
    for filename, section_title in FILES:
        filepath = os.path.join(BASE, filename)
        if not os.path.exists(filepath):
            print(f"WARNING: {filepath} not found")
            continue

        add_page_break(doc)

        with open(filepath, 'r', encoding='utf-8') as f:
            lines = f.readlines()

        table_rows = []
        in_code_block = False

        for line in lines:
            stripped = line.strip()

            # Code block toggles
            if stripped.startswith('```'):
                if in_code_block:
                    in_code_block = False
                    # flush code block as formatted paragraph
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(1)
                else:
                    in_code_block = True
                continue

            if in_code_block:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                run = p.add_run(line.rstrip())
                run.font.name = 'Courier New'
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(60, 60, 60)
                continue

            # Collect table rows
            if stripped.startswith('|') and '|' in stripped[1:]:
                cells = [c.strip() for c in stripped.split('|')[1:-1]]
                if not all(set(c) <= set('- :') for c in cells):
                    table_rows.append(cells)
                    continue

            # Flush pending table before non-table line
            if table_rows:
                flush_table(doc, table_rows)
                table_rows = []

            parse_markdown_line(doc, line, {})

        # Flush any remaining table
        if table_rows:
            flush_table(doc, table_rows)

    # ---- FOOTER ----
    add_page_break(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('— FIN DE LA GUÍA —')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(31, 58, 95)

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Unidad 1: Ship Overview — Inglés Marítimo y Portuario')
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Caribbean Maritime Lines (CML) — FPI SENA — Bilingüismo')
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Instructor Sergio Cortés Perdomo · Marzo 2026')
    run.font.size = Pt(10)

    # Save
    output_path = os.path.join(BASE, "UNIDAD 1 — Ship Overview — Guía Completa.docx")
    doc.save(output_path)
    print(f"OK: {output_path}")
    print(f"Size: {os.path.getsize(output_path):,} bytes")

if __name__ == '__main__':
    build_docx()
