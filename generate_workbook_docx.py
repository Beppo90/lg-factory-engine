#!/usr/bin/env python3
"""Generate Workbook + Answer Key as Word documents."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def sh(cell,color):
    e=OxmlElement('w:shd');e.set(qn('w:fill'),color);e.set(qn('w:val'),'clear')
    cell._tc.get_or_add_tcPr().append(e)

def tbl(doc,headers,rows,sc='1F3A5F'):
    t=doc.add_table(rows=1,cols=len(headers));t.style='Table Grid';t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i];c.text='';r=c.paragraphs[0].add_run(str(h))
        r.font.size=Pt(10);r.font.name='Calibri';r.bold=True;r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF);sh(c,sc)
    for rd in rows:
        row=t.add_row()
        for i,v in enumerate(rd):
            c=row.cells[i];c.text='';r=c.paragraphs[0].add_run(str(v));r.font.size=Pt(9);r.font.name='Calibri'
    return t

def h(d,t,l=1):
    hd=d.add_heading(t,level=l)
    for r in hd.runs:r.font.name='Calibri'

def bp(d,l,v=''):
    p=d.add_paragraph();r=p.add_run(l);r.bold=True;r.font.name='Calibri';r.font.size=Pt(11)
    if v:r2=p.add_run(v);r2.font.name='Calibri';r2.font.size=Pt(11)

def ip(d,t):
    p=d.add_paragraph();r=p.add_run(t);r.italic=True;r.font.name='Calibri';r.font.size=Pt(11)

def st(d,t):
    p=d.add_paragraph(t)
    for r in p.runs:r.font.name='Calibri';r.font.size=Pt(11)

def line(d,n=1):
    for _ in range(n):
        p=doc.add_paragraph('____________________________________________________________')
        for r in p.runs:r.font.name='Calibri';r.font.size=Pt(11);r.font.color.rgb=RGBColor(0xBB,0xBB,0xBB)

def blank_lines(d,n=8):
    for _ in range(n):
        p=doc.add_paragraph('')
        p.paragraph_format.space_after=Pt(0)
        p.paragraph_format.space_before=Pt(0)
        r=p.add_run(' ')
        r.font.size=Pt(11)

# ════════════════════════════════════════
# WORKBOOK
# ════════════════════════════════════════

doc = Document()
for s in doc.sections:
    s.top_margin=Cm(2.5);s.bottom_margin=Cm(2.5);s.left_margin=Cm(2.5);s.right_margin=Cm(2.5)
style=doc.styles['Normal'];style.font.name='Calibri';style.font.size=Pt(11)

# TITLE
for _ in range(5):doc.add_paragraph('')
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('ADSO — GUÍA 1: THE HARDWARE SPECIALIST');r.bold=True;r.font.size=Pt(20);r.font.name='Calibri';r.font.color.rgb=RGBColor(0x1F,0x3A,0x5F)
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Workbook — Autonomous Work');r.font.size=Pt(16);r.font.name='Calibri';r.font.color.rgb=RGBColor(0x44,0x72,0xC4)
doc.add_paragraph('')
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Your practice outside the classroom');r.font.size=Pt(14);r.font.name='Calibri';r.italic=True
doc.add_paragraph('')
doc.add_paragraph('')
ip(doc,'Name: _________________________________  Ficha: _____________  Date: _____________')
doc.add_paragraph('')
for txt in ['Sistema de Prompts Maestros — LG Factory — FPI SENA — Bilingüismo','Instructor Sergio Cortés Perdomo · Marzo 2026']:
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(txt);r.font.size=Pt(10);r.font.name='Calibri';r.font.color.rgb=RGBColor(0x7F,0x8C,0x8D)
doc.add_page_break()

# CH 1
h(doc,'CHAPTER 1: MY HARDWARE PROFILE',1)
bp(doc,'Assigned after: ','Session 1 (The Wake-Up Call)')
bp(doc,'We review this in: ','Session 2 (Read the Request)')
bp(doc,'Estimated time: ','45 minutes')
doc.add_paragraph('')

h(doc,'Activity 1 — Draw Your Computer',2)
ip(doc,'Draw the computer you use at home or at SENA. Label at least 5 parts in English.')
ip(doc,'Use these words if you need help: CPU, RAM, monitor, keyboard, mouse, SSD, HDD, GPU, motherboard, PSU')
doc.add_paragraph('')
for _ in range(12):doc.add_paragraph('')
doc.add_paragraph('')

h(doc,'Activity 2 — Describe Each Part',2)
ip(doc,'Write ONE sentence about each part you labeled. Use "is" or "has."')
ip(doc,'Formula: "The [part] is [adjective]." OR "My [part] has [specification]."')
for i in range(1,6):
    bp(doc,f'{i}. ')
    line(doc)
doc.add_paragraph('')

h(doc,'Activity 3 — My English Level',2)
ip(doc,'Answer honestly. This is for YOU — to see your progress at the end of the guide.')
tbl(doc,['Question','Yes','A little','No'],[
    ['Can you name computer parts in English?','☐','☐','☐'],
    ['Can you describe a computer problem in English?','☐','☐','☐'],
    ['Can you write a short email in English about hardware?','☐','☐','☐'],
],sc='4472C4')
doc.add_page_break()

# CH 2
h(doc,'CHAPTER 2: READING EXTENSION',1)
bp(doc,'Assigned after: ','Session 2 (Read the Request)')
bp(doc,'We review this in: ','Session 3 (Tuning In)')
bp(doc,'Estimated time: ','45 minutes')
doc.add_paragraph('')

h(doc,'Activity 1 — Find "is" and "has"',2)
ip(doc,'Go back to Carlos\'s email (PM-2.3). Underline EVERY example of "is" and "has." Count them.')
tbl(doc,['How many "is"?','How many "has"?'],[['_______','_______']],sc='4472C4')
doc.add_paragraph('')

h(doc,'Activity 2 — Write About Your Computer',2)
ip(doc,'Write 3 NEW sentences about YOUR real computer — using the same structures Carlos used.')
ip(doc,'Formulas: "The [component] is [problem]." / "My [device] has [spec]." / "I need [a new component]."')
for label in ['Sentence 1 (describe a problem):','Sentence 2 (describe a feature):','Sentence 3 (make a request):']:
    bp(doc,label)
    line(doc,2)
doc.add_paragraph('')

h(doc,'Activity 3 — Vocabulary Recall',2)
ip(doc,'Without looking at your worksheet — write the 5 Toolbelt words from Session 2:')
for i,d in enumerate(['(document with technical info)','(to change a part for a better one)','(when two parts can work together)','(how fast something works)','(the money available to buy something)'],1):
    bp(doc,f'{i}. _________________________ ')
    ip(doc,d)
doc.add_page_break()

# CH 3
h(doc,'CHAPTER 3: VOCABULARY REINFORCEMENT',1)
bp(doc,'Assigned after: ','Session 3 (Tuning In)')
bp(doc,'We review this in: ','Session 4 (Say It Right, Build It Right)')
bp(doc,'Estimated time: ','45 minutes')
doc.add_paragraph('')

h(doc,'Activity 1 — Categorize the 20 Words',2)
ip(doc,'Put each word in the correct category.')
ip(doc,'Word Bank: CPU, RAM, GPU, Motherboard, PSU, SSD, HDD, Monitor, Keyboard, Mouse, USB port, HDMI port, Ethernet cable, Printer, Scanner, Gigabyte, Terabyte, Gigahertz, Compatible, Portable')
doc.add_paragraph('')
for cat in ['INPUT (send info TO computer)','OUTPUT (send info FROM computer)','STORAGE (save data)','INTERNAL (inside the computer)','CONNECTIVITY (ports and cables)','SPECS (units + adjectives)']:
    bp(doc,f'{cat}:')
    line(doc,3 if 'INTERNAL' in cat else 2 if 'SPECS' in cat else 2)
    doc.add_paragraph('')

doc.add_page_break()
h(doc,'Activity 2 — Crossword',2)
ip(doc,'Fill in the crossword using the 20 vocabulary words.')
doc.add_paragraph('')
ip(doc,'ACROSS →')
ip(doc,'1. The brain of the computer (3 letters) ___ ___ ___')
ip(doc,'3. Temporary memory, fast access (3 letters) ___ ___ ___')
ip(doc,'5. Shows images on screen (7 letters) ___ ___ ___ ___ ___ ___ ___')
ip(doc,'6. Stores data permanently, solid state (3 letters) ___ ___ ___')
ip(doc,'7. "When two parts can work together" (10 letters) ___ ___ ___ ___ ___ ___ ___ ___ ___ ___')
doc.add_paragraph('')
ip(doc,'DOWN ↓')
ip(doc,'2. Supplies energy to all parts (3 letters) ___ ___ ___')
ip(doc,'3. You type on this (8 letters) ___ ___ ___ ___ ___ ___ ___ ___')
ip(doc,'4. 1000 gigabytes = 1 ___ ___________')
ip(doc,'5. Pointing device (5 letters) ___ ___ ___ ___ ___')
ip(doc,'6. Used to print documents (7 letters) ___ ___ ___ ___ ___ ___ ___')
doc.add_paragraph('')

h(doc,'Activity 3 — My Ideal Workstation',2)
ip(doc,'Write a paragraph (5 sentences) describing your IDEAL workstation. Use at least 8 vocabulary words.')
ip(doc,'Formula to start: "My ideal workstation has _______. The _______ is _______. I need _______ because _______."')
doc.add_paragraph('')
for _ in range(7):line(doc)
doc.add_paragraph('')
ip(doc,'Circle the vocabulary words you used. Count them.')
bp(doc,'Words used: _______ / 8 minimum')
doc.add_page_break()

# CH 4
h(doc,'CHAPTER 4: GRAMMAR DRILL',1)
bp(doc,'Assigned after: ','Session 4 (Say It Right, Build It Right)')
bp(doc,'We review this in: ','Session 5 (Write It Right)')
bp(doc,'Estimated time: ','45 minutes')
doc.add_paragraph('')

h(doc,'Activity 1 — Fill in the Blank',2)
ip(doc,'Complete each sentence with IS, HAS, THIS, THAT, THESE, or THOSE.')
for q in [
    'The CPU _______ the brain of the computer.',
    'My laptop _______ 16 GB of RAM.',
    '_______ is an HDMI port. (pointing near you)',
    'The monitor _______ a VGA port. It is old.',
    '_______ are USB cables. (cables far from you)',
    'An SSD _______ faster than an HDD.',
    'The keyboard _______ compatible with the new system.',
    '_______ is a mechanical drive. (pointing to HDD on desk)',
    'This workstation _______ only 4 GB of RAM.',
    '_______ are old USB 2.0 cables. (pointing to cables far away)',
]:
    ip(doc,q)
    line(doc)
    doc.add_paragraph('')

doc.add_page_break()
h(doc,'Activity 2 — Syntax Bugs',2)
ip(doc,'Some are CORRECT ✓ and some have a BUG ✗. If ✗, write the CORRECT version.')
for q in [
    '"The RAM are 16 gigabytes."',
    '"The GPU is compatible with the motherboard."',
    '"This laptop have an SSD."',
    '"Those is USB cables."',
    '"An SSD is more faster than an HDD."',
]:
    ip(doc,q)
    bp(doc,'☐ Correct ✓ / ☐ Bug ✗ → Correct version: ')
    line(doc)
    doc.add_paragraph('')

h(doc,'Activity 3 — Write Your Own',2)
ip(doc,'Write 3 sentences about a computer in your classroom. Use all 3 structures.')
for label in ['Sentence 1 — Use IS:','Sentence 2 — Use HAS:','Sentence 3 — Use THIS/THAT/THESE/THOSE:']:
    bp(doc,label)
    line(doc)
    doc.add_paragraph('')
doc.add_page_break()

# CH 5
h(doc,'CHAPTER 5: WRITING DRAFT',1)
bp(doc,'Assigned after: ','Session 5 (Write It Right)')
bp(doc,'We review this in: ','Session 6 (The Help Desk)')
bp(doc,'Estimated time: ','60 minutes')
doc.add_paragraph('')

h(doc,'Step 1 — Revise Your Email',2)
ip(doc,'Look at the feedback your partner gave you in class.')
bp(doc,'What did your partner say you did WELL?')
line(doc)
bp(doc,'What SUGGESTION did they give?')
line(doc,2)
doc.add_paragraph('')

h(doc,'Step 2 — Write a Clean Draft',2)
ip(doc,'Rewrite your Tech Request email incorporating your partner\'s feedback. FINAL version.')
doc.add_paragraph('')
bp(doc,'From: _________________________________ — ADSO Apprentice')
bp(doc,'To: IT Department — SENA')
bp(doc,'Subject: Upgrade Request for Workstation #_______')
doc.add_paragraph('')
ip(doc,'Hello,')
doc.add_paragraph('')
for _ in range(5):line(doc)
doc.add_paragraph('')
ip(doc,'Thank you,')
doc.add_paragraph('')
line(doc)
doc.add_paragraph('')

h(doc,'Step 3 — Self-Check',2)
for item in ['I used "is" to describe a problem','I used "has" to describe a feature','I used "I need" to make my request','I included the Subject line with workstation number','My email has: Greeting → Problem → Details → Request → Closing']:
    p=doc.add_paragraph(f'☐  {item}')
    for r in p.runs:r.font.name='Calibri'
bp(doc,'How many checked? _______ / 5')
doc.add_paragraph('')

h(doc,'Step 4 — Read Aloud',2)
ip(doc,'Read your clean draft OUT LOUD — 3 times. Practice pronunciation.')
for i in range(1,4):
    bp(doc,f'Read {i}: _______')
bp(doc,'Did you understand yourself? ☐ Yes / ☐ Partially / ☐ Need more practice')
doc.add_page_break()

# CH 6
h(doc,'CHAPTER 6: SIMULATION PREP',1)
bp(doc,'Assigned after: ','Session 6 (The Help Desk)')
bp(doc,'We review this in: ','Session 7 (Prove What You Know)')
bp(doc,'Estimated time: ','45 minutes')
doc.add_paragraph('')

h(doc,'Activity 1 — Reflection',2)
ip(doc,'Write a paragraph (5 sentences) about your Help Desk simulation experience.')
ip(doc,'Answer: What was EASY? What was HARD? What would you do DIFFERENTLY?')
doc.add_paragraph('')
for _ in range(6):line(doc)
doc.add_paragraph('')

h(doc,'Activity 2 — Vocabulary Review',2)
ip(doc,'Write the 20 vocabulary words from memory. Then check with your Vocab Wall.')
tbl(doc,['#','My word','✓ Correct?'],[[str(i),'_________________________','☐'] for i in range(1,21)],sc='4472C4')
doc.add_paragraph('')
bp(doc,'Total correct: _______ / 20')
doc.add_page_break()

h(doc,'Activity 3 — Grammar Review',2)
ip(doc,'Write one sentence for EACH structure. About hardware.')
for label in ['To Be (IS):','Have/Has:','Demonstrative (THIS/THAT/THESE/THOSE):','Comparison:']:
    bp(doc,label)
    line(doc)
    doc.add_paragraph('')
doc.add_page_break()

# CH 7
h(doc,'CHAPTER 7: POST-QUIZ REVIEW',1)
bp(doc,'Assigned after: ','Session 7 (Prove What You Know)')
bp(doc,'We review this in: ','Session 8 (The Full Circle)')
bp(doc,'Estimated time: ','30 minutes')
doc.add_paragraph('')

h(doc,'Activity 1 — Correct Your Errors',2)
ip(doc,'Find 3 errors from your quiz. Write the wrong answer and the correct answer.')
for i in range(1,4):
    bp(doc,f'Error {i}:')
    bp(doc,'My wrong answer: ');line(doc)
    bp(doc,'The correct answer: ');line(doc)
    bp(doc,'Why: ');line(doc)
    doc.add_paragraph('')

h(doc,'Activity 2 — Improvement Plan',2)
ip(doc,'Write 3 sentences about what you will do DIFFERENTLY next time.')
for i in range(1,4):
    bp(doc,f'{i}.')
    line(doc,2)
    doc.add_paragraph('')

h(doc,'Activity 3 — Final Self-Assessment',2)
ip(doc,'Compare with your answers from Chapter 1, Activity 3.')
tbl(doc,['Question','Session 1','Now (Session 8)'],[
    ['Can you name computer parts in English?','☐Yes ☐Little ☐No','☐Yes ☐Little ☐No'],
    ['Can you describe a computer problem in English?','☐Yes ☐Little ☐No','☐Yes ☐Little ☐No'],
    ['Can you write a short email about hardware?','☐Yes ☐Little ☐No','☐Yes ☐Little ☐No'],
],sc='4472C4')
doc.add_paragraph('')
ip(doc,'Did you improve?')
p=doc.add_paragraph()
r=p.add_run('☐ Yes — I improved a lot!   ☐ A little — I still need practice   ☐ Not yet — but I will keep trying')
r.font.name='Calibri';r.font.size=Pt(11)
doc.add_page_break()

# FOOTER
doc.add_paragraph('')
for txt in ['ADSO — GUÍA 1: The Hardware Specialist — Workbook','Sistema de Prompts Maestros — LG Factory — FPI SENA — Bilingüismo','Instructor Sergio Cortés Perdomo · Marzo 2026']:
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(txt);r.font.size=Pt(9);r.font.name='Calibri';r.font.color.rgb=RGBColor(0x7F,0x8C,0x8D)

doc.save('/Users/Beppo/Projects/fpi-sena-factory/guides/ADSO-G1/ADSO — GUÍA 1 — The Hardware Specialist — Workbook.docx')
print('Workbook done')

# ════════════════════════════════════════
# ANSWER KEY
# ════════════════════════════════════════

ak = Document()
for s in ak.sections:
    s.top_margin=Cm(2.5);s.bottom_margin=Cm(2.5);s.left_margin=Cm(2.5);s.right_margin=Cm(2.5)
style=ak.styles['Normal'];style.font.name='Calibri';style.font.size=Pt(11)

for _ in range(5):ak.add_paragraph('')
p=ak.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('ADSO — GUÍA 1: THE HARDWARE SPECIALIST');r.bold=True;r.font.size=Pt(20);r.font.name='Calibri';r.font.color.rgb=RGBColor(0x1F,0x3A,0x5F)
p=ak.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Workbook — Answer Key');r.font.size=Pt(16);r.font.name='Calibri';r.font.color.rgb=RGBColor(0x44,0x72,0xC4)
doc.add_paragraph('')
p=ak.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('⚠️ SOLO PARA EL INSTRUCTOR. No distribuir a los aprendices.');r.bold=True;r.font.size=Pt(12);r.font.name='Calibri';r.font.color.rgb=RGBColor(0xC0,0x39,0x2B)
ak.add_paragraph('')
for txt in ['Sistema de Prompts Maestros — LG Factory — FPI SENA — Bilingüismo','Instructor Sergio Cortés Perdomo · Marzo 2026']:
    p=ak.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(txt);r.font.size=Pt(10);r.font.name='Calibri';r.font.color.rgb=RGBColor(0x7F,0x8C,0x8D)
ak.add_page_break()

# CH 1 AK
h(ak,'CHAPTER 1: MY HARDWARE PROFILE',1)
bp(ak,'Activity 1: ','No answer key — cada aprendiz dibuja su computador real. Evaluar ≥5 partes etiquetadas en inglés.')
bp(ak,'Activity 2: ','No answer key — evaluar uso correcto de "is" (describir) y "has" (specs/features). 5 oraciones coherentes.')
bp(ak,'Activity 3: ','No answer key — diagnóstico personal. NO revisar.')
ak.add_paragraph('')

h(ak,'CHAPTER 2: READING EXTENSION',1)
bp(ak,'Activity 1:')
tbl(ak,['Palabra','Apariciones'],[['is','12-15 veces'],['has','5-7 veces']],sc='2E75B6')
ip(ak,'Aceptable: cualquier recuento razonable (>8 "is" y >3 "has").')
bp(ak,'Activity 2: ','Ejemplos: "The CPU is old." / "My computer has 4 GB of RAM." / "I need a new monitor."')
bp(ak,'Activity 3: ','1. Spec Sheet  2. Upgrade  3. Compatible  4. Performance  5. Budget')
ak.add_paragraph('')

h(ak,'CHAPTER 3: VOCABULARY REINFORCEMENT',1)
bp(ak,'Activity 1:')
tbl(ak,['INPUT','OUTPUT','STORAGE','INTERNAL','CONNECTIVITY','SPECS'],[
    ['Keyboard','Monitor','SSD','CPU','USB port','Gigabyte'],
    ['Mouse','Printer','HDD','RAM','HDMI port','Terabyte'],
    ['Scanner','','','GPU','Ethernet cable','Gigahertz'],
    ['','','','Motherboard','','Compatible'],
    ['','','','PSU','','Portable'],
],sc='2E75B6')
bp(ak,'Activity 2 — Crossword:')
ip(ak,'ACROSS: 1=CPU, 3=RAM, 5=MONITOR, 6=SSD, 7=COMPATIBLE')
ip(ak,'DOWN: 2=PSU, 3=KEYBOARD, 4=TERABYTE, 5=MOUSE, 6=PRINTER')
bp(ak,'Activity 3: ','Evaluar ≥8 palabras de vocabulario, 5 oraciones coherentes.')
ak.add_page_break()

h(ak,'CHAPTER 4: GRAMMAR DRILL',1)
bp(ak,'Activity 1:')
tbl(ak,['#','Answer'],[['1','is'],['2','has'],['3','This'],['4','has'],['5','Those'],['6','is'],['7','is'],['8','That'],['9','has'],['10','Those']],sc='2E75B6')
bp(ak,'Activity 2:')
tbl(ak,['#','Verdict','Correct version'],[
    ['1','✗','"The RAM is 16 gigabytes."'],['2','✓','—'],['3','✗','"This laptop has an SSD."'],
    ['4','✗','"Those are USB cables."'],['5','✗','"An SSD is faster than an HDD."'],
],sc='2E75B6')
bp(ak,'Activity 3: ','Evaluar: ¿IS correcto? ¿HAS correcto? ¿Demonstrative correcto?')
ak.add_page_break()

h(ak,'CHAPTER 5: WRITING DRAFT',1)
bp(ak,'Step 1: ','No answer key — feedback del compañero.')
bp(ak,'Step 2 — Rúbrica:')
tbl(ak,['Criterio','0 pts','1 pt','2 pts'],[
    ['Format','Sin formato','Parcial','Completo'],['Grammar','0-1 correctas','2-3','4+'],
    ['Vocabulary','0-2 términos','3-5','6+'],['Clarity','Incomprensible','Parcial','Claro'],
],sc='2E75B6')
bp(ak,'Steps 3-4: ','Self-check y lectura — autoevaluación del aprendiz.')
ak.add_page_break()

h(ak,'CHAPTER 6: SIMULATION PREP',1)
bp(ak,'Activity 1: ','No answer key — reflexión. Evaluar 5 oraciones coherentes.')
bp(ak,'Activity 2: ','Las 20 palabras: CPU, RAM, GPU, Motherboard, PSU, SSD, HDD, Monitor, Keyboard, Mouse, USB port, HDMI port, Ethernet cable, Printer, Scanner, Gigabyte, Terabyte, Gigahertz, Compatible, Portable')
bp(ak,'Activity 3: ','Ejemplos: IS="The CPU is the processor." HAS="The laptop has 16 GB." DEM="This is an HDMI port." COMP="An SSD is faster than an HDD."')
ak.add_page_break()

h(ak,'CHAPTER 7: POST-QUIZ REVIEW',1)
bp(ak,'Activity 1: ','No answer key — evaluar: 3 errores reales identificados + respuesta correcta + razón.')
bp(ak,'Activity 2: ','No answer key — evaluar 3 oraciones coherentes sobre acciones de mejora.')
bp(ak,'Activity 3: ','No answer key — autoevaluación comparativa. NO revisar.')

ak.add_page_break()
ak.add_paragraph('')
for txt in ['ADSO — GUÍA 1: The Hardware Specialist — Workbook Answer Key','Sistema de Prompts Maestros — LG Factory — FPI SENA — Bilingüismo','Instructor Sergio Cortés Perdomo · Marzo 2026']:
    p=ak.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(txt);r.font.size=Pt(9);r.font.name='Calibri';r.font.color.rgb=RGBColor(0x7F,0x8C,0x8D)

ak.save('/Users/Beppo/Projects/fpi-sena-factory/guides/ADSO-G1/ADSO — GUÍA 1 — The Hardware Specialist — Workbook Answer Key.docx')
print('Answer Key done')
