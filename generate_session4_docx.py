#!/usr/bin/env python3
"""Generate Session 4 Build-Out as a Word document."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
for s in doc.sections:
    s.top_margin=Cm(2.5);s.bottom_margin=Cm(2.5);s.left_margin=Cm(2.5);s.right_margin=Cm(2.5)
style=doc.styles['Normal'];style.font.name='Calibri';style.font.size=Pt(11)

def sh(cell,color):
    e=OxmlElement('w:shd');e.set(qn('w:fill'),color);e.set(qn('w:val'),'clear')
    cell._tc.get_or_add_tcPr().append(e)

def tbl(doc,headers,rows,sc='1F3A5F'):
    t=doc.add_table(rows=1,cols=len(headers));t.style='Table Grid';t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i];c.text='';r=c.paragraphs[0].add_run(str(h))
        r.font.size=Pt(10);r.font.name='Calibri';r.bold=True;r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF);sh(c,sc)
    for rd in rows:
        row=t.add_row()
        for i,v in enumerate(rd):
            c=row.cells[i];c.text='';r=c.paragraphs[0].add_run(str(v));r.font.size=Pt(9);r.font.name='Calibri'
    return t

def h(d,t,l=1):
    hd=d.add_heading(t,level=l)
    for r in hd.runs:r.font.name='Calibri'

def bp(d,l,v=''):
    p=d.add_paragraph();r=p.add_run(l);r.bold=True;r.font.name='Calibri';r.font.size=Pt(11)
    if v:r2=p.add_run(v);r2.font.name='Calibri';r2.font.size=Pt(11)

def tt(d,t):
    p=d.add_paragraph();r=p.add_run('Teacher Talk: ');r.bold=True;r.italic=True;r.font.name='Calibri'
    r2=p.add_run(f'"{t}"');r2.italic=True;r2.font.name='Calibri'

def fn(d,t):
    p=d.add_paragraph();r=p.add_run('💡 ');r.font.name='Calibri'
    r2=p.add_run(t);r2.italic=True;r2.font.name='Calibri';r2.font.color.rgb=RGBColor(0x2E,0x75,0xB6)

def ck(d,t):
    p=d.add_paragraph();r=p.add_run('✓ Checkpoint: ');r.bold=True;r.font.name='Calibri';r.font.color.rgb=RGBColor(0x2E,0x7D,0x32)
    r2=p.add_run(t);r2.font.name='Calibri'

def icq(d,q,e):
    p=d.add_paragraph();r=p.add_run('ICQ: ');r.bold=True;r.font.name='Calibri'
    r2=p.add_run(f'"{q}"');r2.font.name='Calibri'
    r3=p.add_run(f' → Esperan: "{e}"');r3.italic=True;r3.font.name='Calibri';r3.font.color.rgb=RGBColor(0x7F,0x8C,0x8D)

def ip(d,t):
    p=d.add_paragraph();r=p.add_run(t);r.italic=True;r.font.name='Calibri';r.font.size=Pt(11)

def st(d,t):
    p=d.add_paragraph(t)
    for r in p.runs:r.font.name='Calibri';r.font.size=Pt(11)

def tr(d,l,t):
    p=d.add_paragraph();r=p.add_run(f'{l} ');r.bold=True;r.font.name='Calibri'
    r2=p.add_run(t);r2.italic=True;r2.font.name='Calibri'

# TITLE
for _ in range(4):doc.add_paragraph('')
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('SESSION 4: SAY IT RIGHT, BUILD IT RIGHT');r.bold=True;r.font.size=Pt(20);r.font.name='Calibri';r.font.color.rgb=RGBColor(0x1F,0x3A,0x5F)
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('ADSO — GUÍA 1: The Hardware Specialist — Build-Out');r.font.size=Pt(14);r.font.name='Calibri';r.font.color.rgb=RGBColor(0x44,0x72,0xC4)
doc.add_paragraph('')
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('⚠️ Este documento es SOLO para el instructor. No distribuir a los aprendices.');r.bold=True;r.font.size=Pt(12);r.font.name='Calibri';r.font.color.rgb=RGBColor(0xC0,0x39,0x2B)
doc.add_paragraph('')
for txt in ['Sistema de Prompts Maestros — LG Factory — FPI SENA — Bilingüismo','Instructor Sergio Cortés Perdomo · Marzo 2026']:
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(txt);r.font.size=Pt(10);r.font.name='Calibri';r.font.color.rgb=RGBColor(0x7F,0x8C,0x8D)
doc.add_page_break()

# HEADER
h(doc,'SESSION HEADER',1)
tbl(doc,['Campo','Dato'],[['Programa','ADSO — 228118'],['Guía','Guía 1: The Hardware Specialist'],['Session','4: Say It Right, Build It Right'],['Worksheets','PM-2.6 (Pronunciation) + PM-2.7 (Grammar)'],['Duración','180 minutos'],['Habilidades foco','P● G●'],['Habilidades soporte','V○'],['Trabajo autónomo','Workbook Ch. 4 — Grammar Drill (45 min)'],['Siguiente sesión','Session 5: Write It Right (PM-2.8)']])
doc.add_paragraph('')

# MATERIALS
h(doc,'MATERIALS CHECKLIST',1)
p=doc.add_paragraph('Marcar ANTES de entrar al aula:');p.runs[0].italic=True
for item in ['PM-2.6: Pronunciation worksheet (1 por estudiante + 2 extras)','PM-2.7: Grammar worksheet (1 por estudiante + 2 extras)','Canva slides 15-22 abiertas y listas','Proyector encendido y funcionando','Tablero preparado según Board Plan','Markers de colores (negro + rojo)','Timer visible']:
    p=doc.add_paragraph(f'☐  {item}')
    for r in p.runs:r.font.name='Calibri';r.font.size=Pt(11)
doc.add_paragraph('')

# BOARD PLAN
h(doc,'BOARD PLAN',1)
bt=('SESSION 4: SAY IT RIGHT, BUILD IT RIGHT\n'
    'Today: "Pronounce the words + build sentences with grammar"\n\n'
    '[PHONETIC ZONE]              [GRAMMAR FORMULAS]\n'
    '1. Processor /pro-CE-sor/    STRUCTURE 1: VERB TO BE\n'
    '2. Compatible /com-PA-ti-bol/ [Component] + IS + [definition]\n'
    '3. Monitor /MÓ-ni-tor/       "The CPU is the brain."\n'
    '4. Keyboard /QUÍI-bord/\n'
    '5. Hardware /JÁRD-uer/       STRUCTURE 2: HAVE / HAS\n'
    '6. Software /SÓFT-uer/       [Device] + HAS + [feature]\n'
    '7. Upgrade /ÓP-greid/        "The laptop has a keyboard."\n'
    '8. Gigabyte /GÍ-ga-bait/\n'
    '9. Ethernet /Í-zer-net/      STRUCTURE 3: DEMONSTRATIVES\n'
    '10. HDMI /eich-di-em-ÁI/     THIS/THAT + IS / THESE/THOSE + ARE\n\n'
    '[CHUNKS]                      [LIVE ZONE — vacío]\n'
    '"The CPU is"                  (respuestas y oraciones)\n'
    '"It has a"\n'
    '"This is an"\n'
    '"faster than"\n'
    '"compatible with"')
p=doc.add_paragraph();r=p.add_run(bt);r.font.name='Consolas';r.font.size=Pt(9)
doc.add_page_break()

# TIMELINE
h(doc,'MINUTE-BY-MINUTE TIMELINE',1)
tbl(doc,['Tiempo','Dur.','Bloque','Actividad','Agrupación','Notas'],[
    ['0:00-0:05','5','SET-UP','Lightning Round','Plenary','20 words fast'],
    ['0:05-0:10','5','SET-UP','Hard words identification','Plenary','4-5 difíciles'],
    ['0:10-0:15','5','SET-UP','Opening + Objective','Plenary',''],
    ['0:15-0:20','5','WHILE-A','Phonetic Toolbox intro','Plenary','10 words'],
    ['0:20-0:35','15','WHILE-A','Activity 1: Perception','Ind→Pairs','Syllables + stress'],
    ['0:35-0:50','15','WHILE-A','Activity 2: Chunking','Plenary→Pairs','5 chunks x 3'],
    ['0:50-1:05','15','WHILE-A','Activity 3: Voice Command','Pairs','Intelligibility'],
    ['1:05-1:20','15','BREAK','BREAK','—',''],
    ['1:20-1:30','10','WHILE-B','Syntax Blueprint','Plenary','3 structures'],
    ['1:30-1:50','20','WHILE-B','Activity 1: Syntax Bugs','Ind→Pairs→Plenary','6 sentences'],
    ['1:50-2:10','20','WHILE-B','Activity 2: Fill Ticket','Ind→Pairs→Plenary','12 blanks'],
    ['2:10-2:35','25','WHILE-B','Activity 3: Diagnostic','Ind→Pairs→Plenary','3 sentences'],
    ['2:35-2:50','15','WHILE-C','Integration drill','Pairs','Pron + Grammar'],
    ['2:50-2:55','5','WHILE-C','✓ Board Check','Plenary','1 correct, 1 bug'],
    ['2:55-3:03','8','WRAP-UP','Exit Ticket','Individual','3 sentences'],
    ['3:03-3:08','5','WRAP-UP','Closing + Autónomo','Plenary','Workbook Ch. 4'],
    ['3:08-3:10','2','WRAP-UP','Preview','Plenary',''],
])
p=doc.add_paragraph();r=p.add_run('Total: 180 minutos ✓');r.bold=True;r.font.name='Calibri';r.font.color.rgb=RGBColor(0x2E,0x7D,0x32)
doc.add_page_break()

# SET-UP
h(doc,'SET-UP DETALLADO (15 min) — 0:00 a 0:15',2)
h(doc,'Warm-up: "Vocabulary Lightning Round" (10 min)',3)
bp(doc,'Paso 1 — Rapid naming (5 min)')
ip(doc,'"Good morning everyone. Look at the screen. I point — you say the word. FAST. Ready?"')
st(doc,'1. Instructor señala 20 imágenes de componentes. Clase responde en coro. Ritmo: 1 cada 3-4 seg.')
st(doc,'2. Si no responden en 3 seg, instructor dice la palabra y clase repite.')
bp(doc,'Paso 2 — Identify hard words (5 min)')
st(doc,'3. "Which words were HARD to pronounce?" — Recoger 4-5 en tablero.')
st(doc,'4. Ejemplos típicos: Compatible, Processor, Ethernet, Gigabyte, Hardware.')
fn(doc,'Hispanohablantes: dificultad con /ð/ en "motherboard", schwa /ə/, estrés de "compatible."')
doc.add_paragraph('')
h(doc,'Teacher Talk — Opening Script (3 min)',3)
ip(doc,'"You know the words. You know the meaning. Today we make sure you can SAY them correctly and USE them in sentences. Two things: pronunciation — how to say the words. Grammar — how to build sentences with is, has, and this/that/these/those. By the end of today, you can describe any computer in English. Ready?"')
doc.add_paragraph('')
h(doc,'Objective (escrito en tablero)',3)
ip(doc,'"Today you will: (1) pronounce 10 hardware words correctly, and (2) build sentences using To Be, Have/Has, and Demonstratives."')
icq(doc,'Are we learning new words or practicing the ones we know?','Practicing')
icq(doc,'Two things — what are they?','Pronunciation and grammar')
doc.add_page_break()

# WHILE — BLOQUE A
h(doc,'WHILE DETALLADO (150 min) — 0:15 a 2:55',2)
h(doc,'BLOQUE A — PRONUNCIATION (50 min) — 0:15 a 1:05',3)
bp(doc,'Worksheet: ','PM-2.6 (Pronunciation & Speaking Skills)')
bp(doc,'Agrupación: ','Plenary → Individual → Pairs')
bp(doc,'Objetivo: ','Los aprendices pronuncian 10 palabras con sílabas y estrés correctos, producen chunks fluidos.')

doc.add_paragraph('')
h(doc,'Phonetic Toolbox: Introducción (5 min) — 0:15 a 0:20',3)
bp(doc,'Instrucciones paso a paso:')
st(doc,'1. Proyectar slide 16 (Phonetic Toolbox).')
st(doc,'2. Instructor presenta 10 palabras con guía fonética. Clase repite 2 veces cada una:')
st(doc,'   Processor /pro-CE-sor/, Compatible /com-PA-ti-bol/, Monitor /MÓ-ni-tor/')
st(doc,'   Keyboard /QUÍI-bord/, Hardware /JÁRD-uer/, Software /SÓFT-uer/')
st(doc,'   Upgrade /ÓP-greid/, Gigabyte /GÍ-ga-bait/, Ethernet /Í-zer-net/')
st(doc,'   HDMI /eich-di-em-ÁI/ (se deletrea, no lee como palabra)')
fn(doc,'"HDMI" es excepción — 4 letras, estrés en última. Los aprendices suelen decir "jádmi." Corrija: "It\'s H-D-M-I."')

doc.add_paragraph('')
h(doc,'Activity 1 — Perception: Syllables & Stress (15 min) — 0:20 a 0:35',3)
bp(doc,'Instrucciones paso a paso:')
st(doc,'1. "Count syllables. CIRCLE stressed syllable. 5 minutes." — Timer: 5 min. Individual.')
bp(doc,'Answer Key:')
tbl(doc,['Word','Syllables','Stressed'],[
    ['COM-PA-TI-BLE','4','PA (2nd)'],['PRO-CE-SSOR','3','CE (2nd)'],['MO-NI-TOR','3','MÓ (1st)'],
    ['KEY-BOARD','2','QUÍI (1st)'],['GI-GA-BYTE','3','GÍ (1st)'],['UP-GRADE','2','ÓP (1st)'],
],sc='2E75B6')
st(doc,'')
st(doc,'2. "Compare with partner. 2 minutes." — Timer: 2 min. Pairs.')
st(doc,'3. Revisar ítem por ítem. Marcar sílaba acentuada en Phonetic Zone con color rojo.')
st(doc,'4. "Practice. Say each word 2 times. Stress the right syllable. Partner. 3 minutes." — Timer: 3 min.')
ck(doc,'"Compatible — everyone together. Ready? Go." — Coro. Si incorrecto: "Again — com-PA-ti-bol."')

doc.add_paragraph('')
h(doc,'Activity 2 — Chunking Drill (15 min) — 0:35 a 0:50',3)
bp(doc,'Instrucciones paso a paso:')
st(doc,'1. "5 chunks. Say them FAST — no pauses."')
st(doc,'2. Para cada chunk: instructor dice lento → rápido → clase repite 3 veces.')
st(doc,'   "The CPU is" — todo junto | "It has a" — suena como una palabra')
st(doc,'   "This is an" — todo junto | "faster than" — sin pausa')
st(doc,'   "compatible with" — todo junto')
st(doc,'3. "Pair up. Take turns. 3 times each chunk. 5 minutes." — Timer: 5 min. Pairs.')
st(doc,'4. "Who can say all 5 fast?" — 1-2 voluntarios. Feedback.')
fn(doc,'"It has a" es el más difícil — "has" casi desaparece. "It sounds like ONE word: it-JÁ-sa."')

doc.add_paragraph('')
h(doc,'Activity 3 — Voice Command: Radio Check (15 min) — 0:50 a 1:05',3)
bp(doc,'Instrucciones paso a paso:')
st(doc,'1. "You are an IT Specialist. Read this report OUT LOUD to your partner. Partner marks: YES / PARTIALLY / NO."')
st(doc,'2. Proyectar slide 17 (reporte del IT Specialist con componentes).')
st(doc,'3. "Student A reads, Student B checks. Swap. 6 minutes." — Timer: 6 min. Pairs.')
st(doc,'4. "How many had YES for both?" — Manos arriba.')
st(doc,'5. "PARTIALLY — what was the hard word?" — 2-3 respuestas.')
st(doc,'6. Instructor corrige 2-3 errores comunes en plenaria.')
fn(doc,'Errores típicos: "gigabytes" estrés incorrecto, "compatible" demasiadas sílabas, "HDMI" como palabra.')
tr(doc,'Transition → BREAK:','"Great pronunciation. Take a break — 15 minutes. Come back ready for grammar."')
doc.add_page_break()

# BREAK
h(doc,'BREAK — Stretch & Recharge (15 min) — 1:05 a 1:20',2)
doc.add_paragraph('Instructor toma nota de:')
for item in ['Qué palabras causaron más dificultad','Qué chunks necesitan más práctica','Preparar Grammar Formulas si no están']:
    p=doc.add_paragraph(item,style='List Bullet')
    for r in p.runs:r.font.name='Calibri';r.font.size=Pt(11)
doc.add_page_break()

# BLOQUE B — GRAMMAR
h(doc,'BLOQUE B — GRAMMAR (60 min) — 1:20 a 2:35',3)
bp(doc,'Worksheet: ','PM-2.7 (Grammar & Structure Use)')
bp(doc,'Agrupación: ','Plenary → Individual → Pairs → Plenary')
bp(doc,'Objetivo: ','Los aprendices dominan 3 estructuras (To Be, Have/Has, Demonstratives) mediante identificación de errores, práctica controlada y producción aplicada.')

doc.add_paragraph('')
h(doc,'Transición + Syntax Blueprint (10 min) — 1:20 a 1:30',3)
tt(doc,'Welcome back. You can say the words. Now — build SENTENCES. Three structures. Everything you need to describe hardware. Look at the board.')
bp(doc,'Instrucciones paso a paso:')
st(doc,'1. Structure 1 — VERB TO BE: "[Component] + IS + [definition/adjective]" → "The CPU is the brain." — Copiar 1 min.')
st(doc,'2. Structure 2 — HAVE/HAS: "[Device] + HAS + [feature]" → "The laptop has a keyboard." — Copiar 1 min.')
st(doc,'3. Structure 3 — DEMONSTRATIVES: "THIS/THAT + IS / THESE/THOSE + ARE" → "This is an HDMI port." — Copiar 1 min.')
st(doc,'4. "Read all three formulas to your partner. 1 minute." — Timer: 1 min. Pairs.')
fn(doc,'A1 suelen confundir "is" y "has": "IS = what it IS. HAS = what it CONTAINS."')

doc.add_paragraph('')
h(doc,'Activity 1 — Syntax Bugs (20 min) — 1:30 a 1:50',3)
bp(doc,'Instrucciones paso a paso:')
st(doc,'1. "6 sentences. Correct ✓ or Bug ✗. Find the bug, select correct version. 5 minutes." — Timer: 5 min. Individual.')
bp(doc,'Answer Key:')
tbl(doc,['#','Sentence','Verdict','Corrección / Explicación'],[
    ['1','"The RAM are 16 GB."','✗','"The RAM **is**" — RAM is singular'],
    ['2','"The CPU is the processor."','✓','—'],
    ['3','"This laptop have an SSD."','✗','"This laptop **has**" — singular → has'],
    ['4','"Those is USB cables."','✗','"Those **are**" — plural → are'],
    ['5','"The monitor has an HDMI port."','✓','—'],
    ['6','"An SSD is more faster."','✗','"An SSD is **faster**" — no "more" with -er'],
],sc='2E75B6')
st(doc,'')
st(doc,'2. "Compare with partner. 3 minutes." — Timer: 3 min. Pairs.')
st(doc,'3. Revisar 6 ítems. Para cada bug: escribir incorrecto en rojo ✗, correcto en verde ✓. Preguntar "Why?"')
fn(doc,'Ítem 6: "más rápido" en español = "more fast" en inglés. Pero "faster" ya incluye el "more."')
doc.add_page_break()

h(doc,'Activity 2 — Fill the Ticket (20 min) — 1:50 a 2:10',3)
bp(doc,'Instrucciones paso a paso:')
st(doc,'1. "Support Ticket #0042. Fill IS, HAS, THIS, THAT, THESE, THOSE. 7 minutes." — Timer: 7 min. Individual.')
bp(doc,'Answer Key (12 blanks):')
tbl(doc,['Blank','Answer','Rule'],[
    ['1. "workstation _______ slow"','is','To Be — describing'],
    ['2. "CPU _______ an old i3"','is','To Be — defining'],
    ['3. "_______ processor"','This','Demonstrative — singular, near'],
    ['4. "_______ not compatible"','is','To Be — adjective'],
    ['5. "computer _______ 4 GB"','has','Have/Has — specs'],
    ['6. "_______ _______ not enough"','This is','Demonstrative + To Be'],
    ['7. "storage _______ an HDD"','is','To Be — defining'],
    ['8. "_______ _______ slow drives"','Those are','Demonstrative — plural, far'],
    ['9. "_______ _______ the keyboard"','This is','Demonstrative — pointing'],
    ['10. "It _______ OK"','is','To Be — adjective'],
    ['11. "_______ cables over there"','Those','Demonstrative — plural, far'],
    ['12. "_______ _______ old USB 2.0"','Those are','Demonstrative + To Be'],
],sc='2E75B6')
st(doc,'')
st(doc,'2. "Compare with partner. 3 minutes." — Timer: 3 min. Pairs.')
st(doc,'3. Revisar blank por blank en plenaria.')
fn(doc,'Demonstratives THESE/THOSE (plural) son más difíciles que THIS/THAT. "THIS/THAT = one. THESE/THOSE = more than one. THIS/THESE = near. THAT/THOSE = far."')

doc.add_paragraph('')
h(doc,'Activity 3 — Diagnostic Report (25 min) — 2:10 a 2:35',3)
bp(doc,'Instrucciones paso a paso:')
st(doc,'1. "You are a QA Tester. Write 3 sentences about a computer in this classroom. Use IS, HAS, and one DEMONSTRATIVE. 7 minutes." — Timer: 7 min. Individual.')
st(doc,'2. "Share with partner. Check: is/has/demonstrative correct? 2 minutes." — Timer: 2 min. Pairs.')
st(doc,'3. "Who wants to share one sentence?" — 4-5 voluntarios. Instructor escribe en LIVE ZONE.')
st(doc,'4. Para cada oración: correcta → confirmar. Con error → corregir colectivamente.')
fn(doc,'Puente entre gramática controlada y producción libre. Los errores son NORMALES. Corrija como práctica, no examen.')
tr(doc,'Transition → Bloque C:','"Good sentences. Now let\'s put pronunciation AND grammar together."')
doc.add_page_break()

# BLOQUE C — INTEGRATION
h(doc,'BLOQUE C — INTEGRATION (20 min) — 2:35 a 2:55',3)
bp(doc,'Objetivo: ','Los aprendices integran pronunciación y gramática en producción oral fluida.')

doc.add_paragraph('')
h(doc,'Pronunciation + Grammar Combined Drill (15 min) — 2:35 a 2:50',3)
bp(doc,'Instrucciones paso a paso:')
st(doc,'1. "Your 3 sentences from Activity 3. Read ALOUD to partner. Partner checks TWO things:"')
st(doc,'   Check 1: Grammar correct? (is/has/demonstrative)')
st(doc,'   Check 2: Pronunciation intelligible?')
st(doc,'2. "Student A reads first. Student B checks. Swap. 6 minutes." — Timer: 6 min. Pairs.')
st(doc,'3. "Who had PERFECT grammar AND pronunciation?" — 2-3 parejas.')
st(doc,'4. "Let\'s hear one." — 2-3 voluntarios leen. Instructor da feedback pronunciación + gramática.')
st(doc,'5. Corrige 2-3 errores comunes en plenaria. Escribe incorrecto en LIVE ZONE, clase corrige.')

doc.add_paragraph('')
h(doc,'✓ Board Check (5 min) — 2:50 a 2:55',3)
st(doc,'1. Instructor escribe 2 oraciones:')
ip(doc,'✓ "This monitor has an HDMI port and it is compatible with the GPU."\n✗ "Those cable is old and the keyboard have a broken key."')
st(doc,'2. "Which is correct? What\'s wrong with the second?" → "Those cables ARE" / "keyboard HAS"')
tr(doc,'Transition → WRAP-UP:','"Excellent work. Let\'s close."')
doc.add_page_break()

# WRAP-UP
h(doc,'WRAP-UP DETALLADO (15 min) — 2:55 a 3:10',2)
h(doc,'Exit Ticket (8 min) — 2:55 a 3:03',3)
bp(doc,'Instrucciones:')
st(doc,'1. "Write 3 sentences. Each uses ONE structure:"')
st(doc,'   Sentence 1: Use IS | Sentence 2: Use HAS | Sentence 3: Use THIS/THAT/THESE/THOSE')
st(doc,'2. "All about hardware. 3 minutes." — Timer: 3 min. Individual.')
st(doc,'3. "Show partner. Check grammar. 30 seconds." — Timer: 30 seg. Pairs.')
bp(doc,'Ejemplos válidos:')
st(doc,'- IS: "The CPU is old." / "The monitor is small."')
st(doc,'- HAS: "The laptop has 16 GB of RAM." / "The keyboard has a broken key."')
st(doc,'- DEM: "This is an HDMI port." / "Those are USB cables."')
doc.add_paragraph('')
h(doc,'Teacher Talk — Closing Script (5 min) — 3:03 a 3:08',3)
ip(doc,'"Good work today. Pronunciation — how to say the words. Grammar — how to build sentences. These are YOUR tools.\n\nHomework: Workbook Chapter 4. 10 fill-in-blank sentences (is/has/demonstratives) + 5 Syntax Bug corrections. 45 minutes.\n\nNext session — you WRITE. Your own Tech Request email. Everything you\'ve learned: vocabulary, pronunciation, grammar. See you next time."')
icq(doc,'How many fill-in-blank sentences?','10')
icq(doc,'How many Syntax Bug corrections?','5')
doc.add_paragraph('')
h(doc,'Preview — 3:08 a 3:10',3)
ip(doc,'"Next session: Write It Right. Your own Tech Request email. Using Carlos\'s formulas and today\'s grammar. Come ready to write."')
doc.add_page_break()

# ANSWER KEY
h(doc,'ANSWER KEY CONSOLIDADO',1)
h(doc,'PM-2.6 — Pronunciation',2)
tbl(doc,['Actividad','Ítem','Respuesta'],[
    ['Activity 1','COM-PA-TI-BLE','4 syll, stress PA (2nd)'],['Activity 1','PRO-CE-SSOR','3 syll, stress CE (2nd)'],
    ['Activity 1','MO-NI-TOR','3 syll, stress MÓ (1st)'],['Activity 1','KEY-BOARD','2 syll, stress QUÍI (1st)'],
    ['Activity 1','GI-GA-BYTE','3 syll, stress GÍ (1st)'],['Activity 1','UP-GRADE','2 syll, stress ÓP (1st)'],
],sc='2E75B6')
doc.add_paragraph('')
h(doc,'PM-2.7 — Grammar',2)
tbl(doc,['Actividad','Ítem','Respuesta'],[
    ['Act 1','1. RAM are','✗ → "RAM is"'],['Act 1','2. CPU is the processor','✓'],
    ['Act 1','3. laptop have','✗ → "laptop has"'],['Act 1','4. Those is','✗ → "Those are"'],
    ['Act 1','5. monitor has','✓'],['Act 1','6. more faster','✗ → "faster"'],
    ['Act 2','Blank 1','is'],['Act 2','Blank 2','is'],['Act 2','Blank 3','This'],
    ['Act 2','Blank 4','is'],['Act 2','Blank 5','has'],['Act 2','Blank 6','This is'],
    ['Act 2','Blank 7','is'],['Act 2','Blank 8','Those are'],['Act 2','Blank 9','This is'],
    ['Act 2','Blank 10','is'],['Act 2','Blank 11','Those'],['Act 2','Blank 12','Those are'],
],sc='2E75B6')
doc.add_page_break()

# DIFFERENTIATION
h(doc,'DIFFERENTIATION NOTES',1)
h(doc,'Fast Finishers',2)
for t,d in [
    ('Perception','"Write phonetic guide for 4 MORE words: Motherboard, Scanner, Printer, Gigahertz."'),
    ('Chunking','"Create 3 NEW chunks: \'The SSD is fast,\' \'compatible with the,\' \'I need a.\'"'),
    ('Voice Command','"Write a SECOND report about a DIFFERENT workstation. 5 sentences."'),
    ('Syntax Bugs','"Write 4 NEW sentences — 2 correct ✓, 2 bugs ✗. Give to partner to fix."'),
    ('Diagnostic','"Write 3 MORE sentences about your phone or laptop at home."'),
]:bp(doc,f'{t}: ',d)
doc.add_paragraph('')
h(doc,'More Support Needed',2)
for t,d in [
    ('Pronunciation','"Cheat sheet" con fonética simplificada (aproximación en español). Visible toda la sesión.'),
    ('Perception','Proporcionar sílabas ya separadas. Aprendiz solo marca sílaba acentuada.'),
    ('Chunking','1 min de práctica individual con instructor antes de pair work.'),
    ('Syntax Bugs','Reducir a 4 ítems (2 correctos + 2 bugs más obvios: #1 y #4).'),
    ('Fill the Ticket','Word Bank visible: IS, HAS, THIS, THAT, THESE, THOSE. Elegir de 6 opciones.'),
    ('Diagnostic','Sentence starters: "The _______ is _______." / "This/That is a _______."'),
]:bp(doc,f'{t}: ',d)
doc.add_page_break()

# SELF-CHECK
h(doc,'INSTRUCTOR SELF-CHECK',1)
doc.add_paragraph('')
for i,q in enumerate([
    '¿Todos completaron Activity 1 (Perception)? ¿Cuántos tuvieron 6 ítems correctos?',
    '¿Hubo algún momento donde la mayoría parecía perdida? ¿En pronunciación o gramática?',
    '¿Los tiempos se cumplieron o hubo desfases? ¿Dónde?',
    '¿Qué estructuras causaron más errores? ¿"is" vs "has" o demonstratives?',
    '¿Los aprendices integraron pronunciación Y gramática en Bloque C? ¿Listos para Session 5?',
],1):
    bp(doc,f'{i}. {q}')
    p=doc.add_paragraph('Respuesta: _______________________________________________')
    for r in p.runs:r.font.name='Calibri';r.font.size=Pt(11)
    doc.add_paragraph('')

# FOOTER
doc.add_paragraph('')
for txt in ['SESSION 4: SAY IT RIGHT, BUILD IT RIGHT — BUILD-OUT','ADSO — GUÍA 1: The Hardware Specialist','Sistema de Prompts Maestros — LG Factory — FPI SENA — Bilingüismo','Instructor Sergio Cortés Perdomo · Marzo 2026']:
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(txt);r.font.size=Pt(9);r.font.name='Calibri';r.font.color.rgb=RGBColor(0x7F,0x8C,0x8D)

doc.save('/Users/Beppo/Projects/fpi-sena-factory/guides/ADSO-G1/ADSO — GUÍA 1 — Session 4 — Say It Right, Build It Right — Build-Out.docx')
print('Done')
