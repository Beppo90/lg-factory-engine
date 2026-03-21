#!/usr/bin/env python3
"""Generate Session 7 & 8 Build-Outs as Word documents."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def sh(cell,color):
    e=OxmlElement('w:shd');e.set(qn('w:fill'),color);e.set(qn('w:val'),'clear')
    cell._tc.get_or_add_tcPr().append(e)

def tbl(doc,headers,rows,sc='1F3A5F'):
    t=doc.add_table(rows=1,cols=len(headers));t.style='Table Grid';t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i];c.text='';r=c.paragraphs[0].add_run(str(h))
        r.font.size=Pt(10);r.font.name='Calibri';r.bold=True;r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF);sh(c,sc)
    for rd in rows:
        row=t.add_row()
        for i,v in enumerate(rd):
            c=row.cells[i];c.text='';r=c.paragraphs[0].add_run(str(v));r.font.size=Pt(9);r.font.name='Calibri'
    return t

def h(d,t,l=1):
    hd=d.add_heading(t,level=l)
    for r in hd.runs:r.font.name='Calibri'

def bp(d,l,v=''):
    p=d.add_paragraph();r=p.add_run(l);r.bold=True;r.font.name='Calibri';r.font.size=Pt(11)
    if v:r2=p.add_run(v);r2.font.name='Calibri';r2.font.size=Pt(11)

def tt(d,t):
    p=d.add_paragraph();r=p.add_run('Teacher Talk: ');r.bold=True;r.italic=True;r.font.name='Calibri'
    r2=p.add_run(f'"{t}"');r2.italic=True;r2.font.name='Calibri'

def fn(d,t):
    p=d.add_paragraph();r=p.add_run('💡 ');r.font.name='Calibri'
    r2=p.add_run(t);r2.italic=True;r2.font.name='Calibri';r2.font.color.rgb=RGBColor(0x2E,0x75,0xB6)

def ck(d,t):
    p=d.add_paragraph();r=p.add_run('✓ Checkpoint: ');r.bold=True;r.font.name='Calibri';r.font.color.rgb=RGBColor(0x2E,0x7D,0x32)
    r2=p.add_run(t);r2.font.name='Calibri'

def icq(d,q,e):
    p=d.add_paragraph();r=p.add_run('ICQ: ');r.bold=True;r.font.name='Calibri'
    r2=p.add_run(f'"{q}"');r2.font.name='Calibri'
    r3=p.add_run(f' → Esperan: "{e}"');r3.italic=True;r3.font.name='Calibri';r3.font.color.rgb=RGBColor(0x7F,0x8C,0x8D)

def ip(d,t):
    p=d.add_paragraph();r=p.add_run(t);r.italic=True;r.font.name='Calibri';r.font.size=Pt(11)

def st(d,t):
    p=d.add_paragraph(t)
    for r in p.runs:r.font.name='Calibri';r.font.size=Pt(11)

def tr(d,l,t):
    p=d.add_paragraph();r=p.add_run(f'{l} ');r.bold=True;r.font.name='Calibri'
    r2=p.add_run(t);r2.italic=True;r2.font.name='Calibri'

def make_title(doc, session_name, subtitle):
    for _ in range(4):doc.add_paragraph('')
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(session_name);r.bold=True;r.font.size=Pt(20);r.font.name='Calibri';r.font.color.rgb=RGBColor(0x1F,0x3A,0x5F)
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(subtitle);r.font.size=Pt(14);r.font.name='Calibri';r.font.color.rgb=RGBColor(0x44,0x72,0xC4)
    doc.add_paragraph('')
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('⚠️ Este documento es SOLO para el instructor. No distribuir a los aprendices.');r.bold=True;r.font.size=Pt(12);r.font.name='Calibri';r.font.color.rgb=RGBColor(0xC0,0x39,0x2B)
    doc.add_paragraph('')
    for txt in ['Sistema de Prompts Maestros — LG Factory — FPI SENA — Bilingüismo','Instructor Sergio Cortés Perdomo · Marzo 2026']:
        p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(txt);r.font.size=Pt(10);r.font.name='Calibri';r.font.color.rgb=RGBColor(0x7F,0x8C,0x8D)
    doc.add_page_break()

def make_materials(doc, items):
    h(doc,'MATERIALS CHECKLIST',1)
    p=doc.add_paragraph('Marcar ANTES de entrar al aula:');p.runs[0].italic=True
    for item in items:
        p=doc.add_paragraph(f'☐  {item}')
        for r in p.runs:r.font.name='Calibri';r.font.size=Pt(11)
    doc.add_paragraph('')

def make_timeline(doc, rows):
    h(doc,'MINUTE-BY-MINUTE TIMELINE',1)
    tbl(doc,['Tiempo','Dur.','Bloque','Actividad','Agrupación','Notas'],rows)
    p=doc.add_paragraph();r=p.add_run('Total: 180 minutos ✓');r.bold=True;r.font.name='Calibri';r.font.color.rgb=RGBColor(0x2E,0x7D,0x32)
    doc.add_page_break()

def make_selfcheck(doc, questions):
    h(doc,'INSTRUCTOR SELF-CHECK',1)
    doc.add_paragraph('')
    for i,q in enumerate(questions,1):
        bp(doc,f'{i}. {q}')
        p=doc.add_paragraph('Respuesta: _______________________________________________')
        for r in p.runs:r.font.name='Calibri';r.font.size=Pt(11)
        doc.add_paragraph('')

def make_footer(doc, title):
    doc.add_paragraph('')
    for txt in [title,'ADSO — GUÍA 1: The Hardware Specialist','Sistema de Prompts Maestros — LG Factory — FPI SENA — Bilingüismo','Instructor Sergio Cortés Perdomo · Marzo 2026']:
        p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(txt);r.font.size=Pt(9);r.font.name='Calibri';r.font.color.rgb=RGBColor(0x7F,0x8C,0x8D)

def setup_section(doc, title):
    h(doc,title,1)
    doc.add_paragraph('')

# ════════════════════════════════════════
# SESSION 7
# ════════════════════════════════════════

doc7 = Document()
for s in doc7.sections:
    s.top_margin=Cm(2.5);s.bottom_margin=Cm(2.5);s.left_margin=Cm(2.5);s.right_margin=Cm(2.5)
style=doc7.styles['Normal'];style.font.name='Calibri';style.font.size=Pt(11)

make_title(doc7,'SESSION 7: PROVE WHAT YOU KNOW','ADSO — GUÍA 1: The Hardware Specialist — Build-Out')

h(doc7,'SESSION HEADER',1)
tbl(doc7,['Campo','Dato'],[['Programa','ADSO — 228118'],['Guía','Guía 1: The Hardware Specialist'],['Session','7: Prove What You Know'],['Worksheets','PM-4.2 (Cuestionario Técnico — IE-01)'],['Duración','180 minutos'],['Habilidades foco','R● L● V● G● W●'],['Habilidades soporte','—'],['Trabajo autónomo','Workbook Ch. 7 — Post-Quiz Review (30 min)'],['Siguiente sesión','Session 8: The Full Circle'],['Plan B','Si audio no disponible: instructor lee script']])
doc7.add_paragraph('')

make_materials(doc7, ['PM-4.2: Cuestionario Técnico — impresos (1 por estudiante + 3 extras)','Audio file Section 3 — REPRODUCIR antes de clase','Speakers — VERIFICAR que funcionan','Canva slides 32-33 abiertas','Proyector encendido','Tablero preparado según Board Plan','Markers (negro + rojo)','Timer visible','Papel de respuestas extra','Transcripción impresa (Plan B — 2-3 copias)'])
bp(doc7,'NOTA — Seguridad: ','NO dejar cuestionarios sin supervisión. Recoger TODOS al final. Answer Key OCULTO durante aplicación.')
doc7.add_paragraph('')

h(doc7,'BOARD PLAN',1)
bt=('SESSION 7: PROVE WHAT YOU KNOW\n\n'
    '⏱ QUIZ TODAY — 50 points — 5 sections\n\n'
    'Section 1: Reading Comprehension (10 pts)\n'
    'Section 2: Writing Task (10 pts)\n'
    'Section 3: Listening Comprehension (10 pts)\n'
    'Section 4: Key Vocabulary HOTS (10 pts)\n'
    'Section 5: Grammar & Structure HOTS (10 pts)\n\n'
    '⏱ SCHEDULE:\n'
    'Quiz: 0:15 - 2:15 (120 min)\n'
    'Break: 2:15 - 2:25 (10 min)\n'
    'Review: 2:25 - 2:55 (30 min)\n'
    'Close: 2:55 - 3:10 (15 min)')
p=doc7.add_paragraph();r=p.add_run(bt);r.font.name='Consolas';r.font.size=Pt(9)
doc7.add_page_break()

make_timeline(doc7, [
    ['0:00-0:05','5','SET-UP','5 rapid-fire Qs','Ind→Plenary','Activation'],
    ['0:05-0:10','5','SET-UP','Opening + sections','Plenary',''],
    ['0:10-0:15','5','SET-UP','Distribute + header','Individual',''],
    ['0:15-0:35','20','QUIZ','Section 1: Reading','Individual','Silent'],
    ['0:35-1:05','30','QUIZ','Section 2: Writing','Individual','Silent'],
    ['1:05-1:25','20','QUIZ','Section 3: Listening','Individual','Audio 3x'],
    ['1:25-1:45','20','QUIZ','Section 4: Vocab HOTS','Individual','Silent'],
    ['1:45-2:15','30','QUIZ','Section 5: Grammar HOTS','Individual','Silent'],
    ['2:15-2:20','5','QUIZ','Collect quizzes','Individual','All collected'],
    ['2:20-2:30','10','BREAK','BREAK','—',''],
    ['2:30-2:55','25','REVIEW','Feedback: common errors','Plenary','Key corrections'],
    ['2:55-3:03','8','WRAP-UP','Exit Ticket','Individual','1 thing learned'],
    ['3:03-3:08','5','WRAP-UP','Closing + Autónomo','Plenary','Workbook Ch. 7'],
    ['3:08-3:10','2','WRAP-UP','Preview','Plenary','Last session'],
])

h(doc7,'SET-UP DETALLADO (15 min)',2)
h(doc7,'Warm-up: "Quick Review" (5 min)',3)
st(doc7,'1. 5 preguntas en tablero/proyectadas (slide 32):')
st(doc7,'   1. What is the brain of the computer? 2. Write a sentence with "has"')
st(doc7,'   3. What does "compatible" mean? 4. "An SSD is _____ than an HDD."')
st(doc7,'   5. What are the 5 parts of a Tech Request email?')
st(doc7,'2. "Write answers. 2 minutes." — Timer: 2 min. Revisar rápidamente.')
st(doc7,'3. "NO grading. Brain warm-up. Quiz starts NOW."')
doc7.add_paragraph('')
h(doc7,'Opening + Sections (5 min)',3)
ip(doc7,'"Today is quiz day. 50 points. 5 sections. 120 minutes. Read carefully. Use your time well. If you finish early — check your answers."')
st(doc7,'5 Sections: (1) Reading 10pts (2) Writing 10pts (3) Listening 10pts (4) Vocab HOTS 10pts (5) Grammar HOTS 10pts')
icq(doc7,'How many sections?','5')
icq(doc7,'How many points total?','50')
icq(doc7,'Can you pause the audio?','No')
doc7.add_paragraph('')
h(doc7,'Distribute + Header (5 min)',3)
st(doc7,'1. Distribuir boca abajo. "Flip when everyone has one. Fill header: Name, Ficha, Date. DON\'T start yet."')
st(doc7,'2. "Ready? Begin. Section 1: Reading. 20 minutes. Go."')
doc7.add_page_break()

h(doc7,'QUIZ DETALLADO (120 min)',2)
h(doc7,'Section 1: Reading (20 min)',3)
bp(doc7,'Mientras circula: ')
st(doc7,'- NO responde sobre contenido — solo aclara instrucciones')
st(doc7,'- Verifica: sin celular, sin notas')
fn(doc7,'Si preguntan "What does [word] mean?": "I can\'t help with content. Read carefully — the answer is in the text."')
doc7.add_paragraph('')
h(doc7,'Section 2: Writing (30 min)',3)
st(doc7,'- Verifica que usen Language Bank y Skeleton Structure del cuestionario (scaffolding integrado)')
st(doc7,'- Toma nota de quién parece bloqueado')
doc7.add_paragraph('')
h(doc7,'Section 3: Listening (20 min)',3)
st(doc7,'1. "Put down pencils. First listen — no writing." — Reproducir audio 1x.')
st(doc7,'2. "Second listen — now you can write." — Reproducir audio 2x.')
st(doc7,'3. "Third listen — last time. Finish answers." — Reproducir audio 3x.')
st(doc7,'4. "10 more minutes to finish." — Timer: 10 min.')
bp(doc7,'Plan B: ','Instructor lee script con dos voces. Primera: solo escuchar. Segunda: con transcripción (recoger después). Tercera: sin transcripción.')
doc7.add_paragraph('')
h(doc7,'Section 4: Vocabulary HOTS (20 min)',3)
st(doc7,'- Tareas HOTS requieren justificación escrita')
h(doc7,'Section 5: Grammar HOTS (30 min)',3)
st(doc7,'- Dar tiempo completo. Anunciar: "15 min left" / "5 min left"')
doc7.add_paragraph('')
h(doc7,'Collect (5 min)',3)
st(doc7,'"Time. Close. Pass to front." — Recoger TODOS.')
doc7.add_page_break()

h(doc7,'BREAK (10 min)',2)
st(doc7,'Instructor: revisa rápidamente, identifica 3-4 errores comunes por sección, prepara feedback.')
doc7.add_page_break()

h(doc7,'REVIEW (25 min)',2)
h(doc7,'Feedback: Common Errors',3)
st(doc7,'1. "I\'m NOT giving grades today. I\'ll show you common mistakes so you learn from them."')
st(doc7,'2. Revisar patrones por sección (NO todas las respuestas):')
bp(doc7,'Section 1: ','Errores de scanning — "The answer was IN THE TEXT."')
bp(doc7,'Section 3: ','Gap fill — "compatible" no "compitable."')
bp(doc7,'Section 4: ','Odd One Out — "You must JUSTIFY with a sentence."')
bp(doc7,'Section 5: ','Error Log — "Those cable IS → Those cables ARE. Plural."')
st(doc7,'3. "Any questions about specific questions?" — 2-3 preguntas. Guiar sin revelar respuestas.')
fn(doc7,'Feedback es FORMATIVO: "Common mistakes — let\'s learn" NO "You got this wrong."')
doc7.add_page_break()

h(doc7,'WRAP-UP (15 min)',2)
h(doc7,'Exit Ticket (8 min)',3)
st(doc7,'1. "Write ONE thing you learned in this guide that you will use in your career. English or Spanish. 2 min." — Timer: 2 min.')
st(doc7,'2. "Share with partner. 30 sec." — Timer: 30 seg.')
doc7.add_paragraph('')
h(doc7,'Closing (5 min)',3)
ip(doc7,'"Today you proved what you know. I\'ll grade and give results next session.\n\nHomework: Workbook Ch. 7. Correct errors from quiz. Write 3 sentences about what to do differently. 30 minutes.\n\nNext session — LAST session. The Full Circle. Review, reflection, feedback."')
doc7.add_page_break()

h(doc7,'DISTRIBUTION & RÚBRICA',1)
tbl(doc7,['Sección','Habilidad','Pts','Tipo'],[
    ['1. Reading','Receptiva','10','Remember/Understand/Analyze'],
    ['2. Writing','Productiva','10','Apply/Create'],
    ['3. Listening','Receptiva','10','Remember/Understand/Analyze'],
    ['4. Vocabulary HOTS','Productiva','10','Apply/Analyze/Evaluate'],
    ['5. Grammar HOTS','Productiva','10','Apply/Analyze/Evaluate'],
    ['Total','','50',''],
],sc='2E75B6')
doc7.add_paragraph('')
tbl(doc7,['Criterio','0 pts','1 pt','2 pts','3 pts'],[
    ['Format (2)','Sin formato','Parcial','Completo','—'],
    ['Grammar (3)','0-1 correctas','2-3 correctas','4-5 correctas','6+ correctas'],
    ['Vocabulary (3)','0-2 términos','3-4 términos','5-6 términos','7+ términos'],
    ['Clarity (2)','Incomprensible','Parcial','Claro','—'],
],sc='2E75B6')
doc7.add_page_break()

h(doc7,'DIFFERENTIATION NOTES',1)
for t,d in [('Timed accommodations','Dar 30 min extra en espacio separado si hay necesidades especiales documentadas.'),
    ('Reading support','NO diccionario. Pero parafrasear INSTRUCCIONES si no se entiende QUÉ hacer.'),
    ('Writing blocked','Si nada después de 5 min: señalar Language Bank + Skeleton. Guiar, no dar respuesta.'),
    ('Students who finish early','No salir. Revisar respuestas o leer.'),
]:bp(doc7,f'{t}: ',d)
doc7.add_page_break()

make_selfcheck(doc7, [
    '¿Todos completaron las 5 secciones? ¿Quiénes no terminaron?',
    '¿Hubo problemas con audio Section 3? ¿Necesité Plan B?',
    '¿Qué secciones tuvieron errores más comunes?',
    '¿El feedback fue útil? ¿Aprendices preguntaron sobre errores?',
    '¿Qué reforzar en Session 8 basado en patrones de error?',
])
make_footer(doc7,'SESSION 7: PROVE WHAT YOU KNOW — BUILD-OUT')

doc7.save('/Users/Beppo/Projects/fpi-sena-factory/guides/ADSO-G1/ADSO — GUÍA 1 — Session 7 — Prove What You Know — Build-Out.docx')
print('Session 7 done')

# ════════════════════════════════════════
# SESSION 8
# ════════════════════════════════════════

doc8 = Document()
for s in doc8.sections:
    s.top_margin=Cm(2.5);s.bottom_margin=Cm(2.5);s.left_margin=Cm(2.5);s.right_margin=Cm(2.5)
style=doc8.styles['Normal'];style.font.name='Calibri';style.font.size=Pt(11)

make_title(doc8,'SESSION 8: THE FULL CIRCLE','ADSO — GUÍA 1: The Hardware Specialist — Build-Out')

h(doc8,'SESSION HEADER',1)
tbl(doc8,['Campo','Dato'],[['Programa','ADSO — 228118'],['Guía','Guía 1: The Hardware Specialist'],['Session','8: The Full Circle'],['Worksheets','Ninguno (revisión, reflexión, cierre)'],['Duración','180 minutos'],['Habilidades foco','— (consolidación)'],['Habilidades soporte','—'],['Trabajo autónomo','Ninguno'],['Siguiente sesión','Ninguno (fin de guía)']])
doc8.add_paragraph('')

make_materials(doc8, ['PM-2.2: Gap Analysis original (aprendices traen el suyo de Session 1)','Canva slides 34-38 abiertas','Proyector encendido','Tablero preparado según Board Plan','Markers (negro + verde)','Timer visible','Guide Feedback Forms — impresos (1 por estudiante)','Imágenes componentes hardware (slides)','Resultados quiz (opcional — para entregar)'])
doc8.add_paragraph('')

h(doc8,'BOARD PLAN',1)
bt=('SESSION 8: THE FULL CIRCLE\n'
    'Today: "Review, reflect, and close the guide"\n\n'
    '[ONE WORD ZONE — vacío]      [THE JOURNEY]\n'
    '(cada aprendiz dice 1 word)   S1: What is the most\n'
    '                               important part?\n'
    '                              S2: We READ\n'
    '                              S3: We LISTENED\n'
    '                              S4: We PRONOUNCED\n'
    '                              S5: We WROTE\n'
    '                              S6: We SPOKE\n'
    '                              S7: We PROVED\n'
    '                              S8: WE MADE IT ⭐')
p=doc8.add_paragraph();r=p.add_run(bt);r.font.name='Consolas';r.font.size=Pt(9)
doc8.add_page_break()

make_timeline(doc8, [
    ['0:00-0:10','10','SET-UP','One Word warm-up','Plenary','Every student'],
    ['0:10-0:15','5','SET-UP','Opening + Objective','Plenary',''],
    ['0:15-0:25','10','WHILE-A','Component Identification','Plenary','Point & say'],
    ['0:25-0:40','15','WHILE-A','Grammar Relay','Teams of 4','Sentence race'],
    ['0:40-0:55','15','WHILE-A','Pronunciation Check','Plenary','5 hard words'],
    ['0:55-1:10','15','BREAK','BREAK','—',''],
    ['1:10-1:15','5','WHILE-B','Transición: Gap revisit','Plenary',''],
    ['1:15-1:50','35','WHILE-B','Gap Analysis Revisited','Ind→Pairs→Plenary','Session 1 vs Now'],
    ['1:50-2:05','15','WHILE-B','Celebration of growth','Plenary','Individual victories'],
    ['2:05-2:10','5','WHILE-C','Challenge scenario intro','Plenary',''],
    ['2:10-2:30','20','WHILE-C','Group planning','Groups of 4','Spec sheet + call'],
    ['2:30-2:45','15','WHILE-C','Presentations','Plenary','2-3 groups'],
    ['2:45-3:00','15','WHILE-D','Feedback Loop','Individual','Anonymous form'],
    ['3:00-3:08','8','WRAP-UP','Closing ceremony','Plenary','Final words'],
    ['3:08-3:10','2','WRAP-UP','Class dismissed','Plenary',''],
])

h(doc8,'SET-UP (15 min)',2)
h(doc8,'Warm-up: "One Word" (10 min)',3)
st(doc8,'1. "Every person says ONE word learned in this guide. No repeats. I\'ll start: Compatible."')
st(doc8,'2. Ir persona por persona. Escribir en ONE WORD ZONE del tablero.')
st(doc8,'3. Si no puede pensar: "What\'s NOT on the board?" / "What did you use in your Tech Request?"')
fn(doc8,'Clase grande (>20): hacer en grupos de 4 primero (4 palabras únicas por grupo), luego recoger las mejores en plenaria.')
st(doc8,'4. "ALL those words. YOU learned them. In 8 sessions. That\'s real."')
doc8.add_paragraph('')
h(doc8,'Opening (5 min)',3)
ip(doc8,'"Last session. Session 8. We started with: \'What is the most important part of a computer?\' Today we close the circle. Review. Reflect. You tell me — what did this guide give you? What would you change?"')
ip(doc8,'"Today you will: review everything, reflect on growth, and give feedback."')
doc8.add_page_break()

h(doc8,'WHILE — BLOQUE A (40 min)',2)
h(doc8,'Component Identification (10 min)',3)
st(doc8,'1. Proyectar slide 35 (imágenes sin etiquetas).')
st(doc8,'2. "I point — you say the word AND use it in a sentence. Fast." — 15-20 imágenes, 30 seg c/u.')
fn(doc8,'Si nadie sabe: dar primera letra. "Starts with M... shows you the image" → Monitor.')
doc8.add_paragraph('')
h(doc8,'Grammar Relay (15 min)',3)
st(doc8,'1. "Teams of 4. Write as many CORRECT sentences using To Be, Have/Has, Demonstratives. About hardware. 5 min." — Timer: 5 min.')
st(doc8,'2. "Count your sentences. Team with MOST correct wins."')
st(doc8,'3. Teams leen mejores 3. Instructor verifica gramática públicamente. Celebrar ganadores.')
doc8.add_paragraph('')
h(doc8,'Pronunciation Check (15 min)',3)
st(doc8,'1. "5 hard words: Compatible, Processor, Ethernet, Gigabyte, Hardware." — Coro cada una.')
st(doc8,'2. "3 volunteers recite the IT report from Session 4 from memory." — 3 voluntarios.')
st(doc8,'3. Feedback: "Great pronunciation. Much better than Session 4."')
tr(doc8,'Transition → BREAK:','"Break 15 min. Come back to see how far you\'ve come."')
doc8.add_page_break()

h(doc8,'BREAK (15 min)',2)
st(doc8,'Instructor prepara: resultados quiz (opcional), Blind Spots originales.')
doc8.add_page_break()

h(doc8,'WHILE — BLOQUE B (50 min)',2)
h(doc8,'Gap Analysis Revisited (35 min)',3)
tt(doc8,'"Take out your PM-2.2 from Session 1. Your Blind Spots. Your Learning Contract. Did you keep your promise?"')
st(doc8,'1. "Look at Activity 2. Can you do them NOW? Mark ✓ or ❓. 5 min." — Timer: 5 min.')
st(doc8,'2. "Compare with partner. How many ✓ each? 3 min." — Timer: 3 min.')
st(doc8,'3. "ALL 5 ✓?" — Manos. "4 ✓?" — Manos. Celebrar.')
st(doc8,'4. Instructor escribe en tablero: Session 1 vs Session 8 comparison.')
st(doc8,'5. "Look at the board. THAT is progress."')
doc8.add_paragraph('')
h(doc8,'Learning Contract Completion (10 min)',3)
st(doc8,'1. "Activity 3. Check off your contract items. Write completion status. 3 min." — Timer: 3 min.')
st(doc8,'2. "Share with partner. 2 min." — Timer: 2 min.')
st(doc8,'3. "Who kept their promise?" — 3-4 voluntarios.')
doc8.add_paragraph('')
h(doc8,'Celebration of Growth (5 min)',3)
st(doc8,'"One volunteer — ONE thing you can do NOW that you COULDN\'T do 8 sessions ago." — 3-4 voluntarios.')
st(doc8,'Instructor celebra cada uno.')
doc8.add_page_break()

h(doc8,'WHILE — BLOQUE C (40 min)',2)
h(doc8,'The Challenge Scenario (5 min)',3)
ip(doc8,'"You are at your first job. Supervisor: \'New developer starts Monday. Set up their workstation. Write the spec sheet. Call IT if you need parts.\'"')
st(doc8,'1. "Uses EVERYTHING: vocab, grammar, pronunciation, reading, writing, speaking."')
st(doc8,'2. "Groups of 4. Prepare 3 things: (1) Spec Sheet (2) Phone Call practice (3) 2-min presentation."')
doc8.add_paragraph('')
h(doc8,'Group Planning (20 min)',3)
st(doc8,'1. "Groups of 4. NOW." — 1 min.')
st(doc8,'2. "15 minutes. Divide work." — Timer: 15 min. Instructor circula.')
st(doc8,'3. "5 minutes left." — Timer: 5 min.')
doc8.add_paragraph('')
h(doc8,'Presentations (15 min)',3)
st(doc8,'1. 2-3 grupos presentan (2 min cada uno). Clase aplaude.')
st(doc8,'2. Instructor da 1 comentario positivo por presentación.')
fn(doc8,'Celebratorio, no evaluativo. Sin nota. Integración de todo lo aprendido.')
doc8.add_page_break()

h(doc8,'WHILE — BLOQUE D: FEEDBACK LOOP (15 min)',2)
st(doc8,'1. Distribuir Guide Feedback Forms.')
st(doc8,'2. "ANONYMOUS. Don\'t write name. Honest opinion."')
bp(doc8,'Formulario:','3 preguntas: (1) What WORKED? (2) What DIDN\'T work? (3) What would you CHANGE?')
st(doc8,'3. "5 minutes." — Timer: 5 min.')
st(doc8,'4. "Fold and put in box as you leave."')
fn(doc8,'Feedback anónimo es más honesto. NO leer en clase.')
doc8.add_page_break()

h(doc8,'WRAP-UP: CLOSING CEREMONY (10 min)',2)
ip(doc8,'"8 sessions ago, you walked in. Some didn\'t know \'CPU.\' Some couldn\'t write a sentence. Some were afraid to speak.\n\nNow: you read emails, listened to calls, learned 20 words, pronounced them, built sentences, wrote emails, had phone conversations, took a quiz, presented to the class.\n\nThat\'s real. That\'s yours.\n\nThe hardware English you learned is the same English used in companies worldwide. You\'re ready.\n\nThank you for your work. Thank you for your effort. This is just Guide 1 — 5 more to go. Each builds on what you learned here.\n\nWell done."')
st(doc8,'"Class dismissed. Take your worksheets — they\'re yours."')
doc8.add_page_break()

h(doc8,'DIFFERENTIATION NOTES',1)
for t,d in [
    ('One Word','Permitir palabra repetida si no puede pensar una única — objetivo es participación.'),
    ('Grammar Relay','Permitir Grammar Formulas del tablero como referencia.'),
    ('Gap Analysis','Si muchos ❓: enfocar en los ✓. "Look how many you CAN do."'),
    ('Challenge','Emparejar fuertes con apoyo dentro del grupo.'),
]:bp(doc8,f'{t}: ',d)
doc8.add_page_break()

make_selfcheck(doc8, [
    '¿Todos participaron en One Word?',
    '¿Gap Analysis mostró mejora significativa? ¿Qué % tenía los 5 ✓?',
    '¿Grupos completaron el Challenge? ¿Integraron vocab + grammar + formato?',
    '¿Qué feedback recibí en los Forms? ¿Patrones?',
    '¿Qué ajustaría para la próxima ejecución de esta guía?',
])
make_footer(doc8,'SESSION 8: THE FULL CIRCLE — BUILD-OUT')

doc8.save('/Users/Beppo/Projects/fpi-sena-factory/guides/ADSO-G1/ADSO — GUÍA 1 — Session 8 — The Full Circle — Build-Out.docx')
print('Session 8 done')
