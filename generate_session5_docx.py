#!/usr/bin/env python3
"""Generate Session 5 Build-Out as a Word document."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
for s in doc.sections:
    s.top_margin=Cm(2.5);s.bottom_margin=Cm(2.5);s.left_margin=Cm(2.5);s.right_margin=Cm(2.5)
style=doc.styles['Normal'];style.font.name='Calibri';style.font.size=Pt(11)

def sh(cell,color):
    e=OxmlElement('w:shd');e.set(qn('w:fill'),color);e.set(qn('w:val'),'clear')
    cell._tc.get_or_add_tcPr().append(e)

def tbl(doc,headers,rows,sc='1F3A5F'):
    t=doc.add_table(rows=1,cols=len(headers));t.style='Table Grid';t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i];c.text='';r=c.paragraphs[0].add_run(str(h))
        r.font.size=Pt(10);r.font.name='Calibri';r.bold=True;r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF);sh(c,sc)
    for rd in rows:
        row=t.add_row()
        for i,v in enumerate(rd):
            c=row.cells[i];c.text='';r=c.paragraphs[0].add_run(str(v));r.font.size=Pt(9);r.font.name='Calibri'
    return t

def h(d,t,l=1):
    hd=d.add_heading(t,level=l)
    for r in hd.runs:r.font.name='Calibri'

def bp(d,l,v=''):
    p=d.add_paragraph();r=p.add_run(l);r.bold=True;r.font.name='Calibri';r.font.size=Pt(11)
    if v:r2=p.add_run(v);r2.font.name='Calibri';r2.font.size=Pt(11)

def tt(d,t):
    p=d.add_paragraph();r=p.add_run('Teacher Talk: ');r.bold=True;r.italic=True;r.font.name='Calibri'
    r2=p.add_run(f'"{t}"');r2.italic=True;r2.font.name='Calibri'

def fn(d,t):
    p=d.add_paragraph();r=p.add_run('💡 ');r.font.name='Calibri'
    r2=p.add_run(t);r2.italic=True;r2.font.name='Calibri';r2.font.color.rgb=RGBColor(0x2E,0x75,0xB6)

def ck(d,t):
    p=d.add_paragraph();r=p.add_run('✓ Checkpoint: ');r.bold=True;r.font.name='Calibri';r.font.color.rgb=RGBColor(0x2E,0x7D,0x32)
    r2=p.add_run(t);r2.font.name='Calibri'

def icq(d,q,e):
    p=d.add_paragraph();r=p.add_run('ICQ: ');r.bold=True;r.font.name='Calibri'
    r2=p.add_run(f'"{q}"');r2.font.name='Calibri'
    r3=p.add_run(f' → Esperan: "{e}"');r3.italic=True;r3.font.name='Calibri';r3.font.color.rgb=RGBColor(0x7F,0x8C,0x8D)

def ip(d,t):
    p=d.add_paragraph();r=p.add_run(t);r.italic=True;r.font.name='Calibri';r.font.size=Pt(11)

def st(d,t):
    p=d.add_paragraph(t)
    for r in p.runs:r.font.name='Calibri';r.font.size=Pt(11)

def tr(d,l,t):
    p=d.add_paragraph();r=p.add_run(f'{l} ');r.bold=True;r.font.name='Calibri'
    r2=p.add_run(t);r2.italic=True;r2.font.name='Calibri'

# TITLE
for _ in range(4):doc.add_paragraph('')
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('SESSION 5: WRITE IT RIGHT');r.bold=True;r.font.size=Pt(20);r.font.name='Calibri';r.font.color.rgb=RGBColor(0x1F,0x3A,0x5F)
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('ADSO — GUÍA 1: The Hardware Specialist — Build-Out');r.font.size=Pt(14);r.font.name='Calibri';r.font.color.rgb=RGBColor(0x44,0x72,0xC4)
doc.add_paragraph('')
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('⚠️ Este documento es SOLO para el instructor. No distribuir a los aprendices.');r.bold=True;r.font.size=Pt(12);r.font.name='Calibri';r.font.color.rgb=RGBColor(0xC0,0x39,0x2B)
doc.add_paragraph('')
for txt in ['Sistema de Prompts Maestros — LG Factory — FPI SENA — Bilingüismo','Instructor Sergio Cortés Perdomo · Marzo 2026']:
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(txt);r.font.size=Pt(10);r.font.name='Calibri';r.font.color.rgb=RGBColor(0x7F,0x8C,0x8D)
doc.add_page_break()

h(doc,'SESSION HEADER',1)
tbl(doc,['Campo','Dato'],[['Programa','ADSO — 228118'],['Guía','Guía 1: The Hardware Specialist'],['Session','5: Write It Right'],['Worksheets','PM-2.8 (Writing Skills & Pragmatics)'],['Duración','180 minutos'],['Habilidades foco','W●'],['Habilidades soporte','G○ V○ R○'],['Trabajo autónomo','Workbook Ch. 5 — Writing Draft (60 min)'],['Siguiente sesión','Session 6: The Help Desk (PM-2.9)']])
doc.add_paragraph('')

h(doc,'MATERIALS CHECKLIST',1)
p=doc.add_paragraph('Marcar ANTES de entrar al aula:');p.runs[0].italic=True
for item in ['PM-2.8: Writing worksheet (1 por estudiante + 2 extras)','Canva slides 23-27 abiertas y listas','Proyector encendido y funcionando','Tablero preparado según Board Plan','Markers (negro + verde)','Timer visible','Hojas blancas extra']:
    p=doc.add_paragraph(f'☐  {item}')
    for r in p.runs:r.font.name='Calibri';r.font.size=Pt(11)
doc.add_paragraph('')

h(doc,'BOARD PLAN',1)
bt=('SESSION 5: WRITE IT RIGHT\n'
    'Today: "Write your OWN Tech Request email"\n\n'
    '[WRITING FORMULAS]             [EMAIL STRUCTURE]\n'
    'DESCRIBE: "The [X] is [Y]."   Greeting → Problem\n'
    'REQUEST: "I need [X]."         → Details → Request\n'
    'COMPARE: "An X is faster      → Closing\n'
    '          than a Y."\n\n'
    '[LANGUAGE BANK]                [LIVE ZONE — vacío]\n'
    'has — is — need               (ejemplos de aprendices\n'
    'not compatible — is old        y correcciones)\n'
    'This is')
p=doc.add_paragraph();r=p.add_run(bt);r.font.name='Consolas';r.font.size=Pt(9)
doc.add_page_break()

h(doc,'MINUTE-BY-MINUTE TIMELINE',1)
tbl(doc,['Tiempo','Dur.','Bloque','Actividad','Agrupación','Notas'],[
    ['0:00-0:05','5','SET-UP','Formula Recall','Individual','3 structures memory'],
    ['0:05-0:10','5','SET-UP','Volunteers write','Plenary','Board formulas'],
    ['0:10-0:15','5','SET-UP','Opening + Objective','Plenary',''],
    ['0:15-0:20','5','WHILE-A','Blueprint Model','Plenary','Display email'],
    ['0:20-0:40','20','WHILE-A','Activity 1: Pragmatic Analysis','Ind→Pairs→Plenary','Match purpose'],
    ['0:40-0:45','5','WHILE-A','Discussion: genre conventions','Pairs→Plenary','Why problem first?'],
    ['0:45-1:00','15','BREAK','BREAK','—',''],
    ['1:00-1:05','5','WHILE-B','Transición: Guided Drafting','Plenary',''],
    ['1:05-1:35','30','WHILE-B','Fill the Skeleton','Individual→Plenary','8 blanks'],
    ['1:35-1:40','5','WHILE-B','Transición: Final Task','Plenary',''],
    ['1:40-2:10','30','WHILE-C','Final Task: Write email','Individual','Instructor circulates'],
    ['2:10-2:12','2','WHILE-C','Auditor\'s Checklist','Individual','Self-check'],
    ['2:12-2:17','5','WHILE-C','✓ Checkpoint','Individual','Hand count'],
    ['2:17-2:32','15','WHILE-D','Peer Review','Pairs','Exchange + feedback'],
    ['2:32-2:37','5','WHILE-D','Share feedback','Pairs→Plenary','2-3 pairs'],
    ['2:37-2:42','5','WHILE-D','Revise + Submit','Individual','Final changes'],
    ['2:42-2:50','8','WRAP-UP','Exit Ticket','Individual','2 sentences'],
    ['2:50-2:55','5','WRAP-UP','Closing + Autónomo','Plenary','Workbook Ch. 5'],
    ['2:55-2:57','2','WRAP-UP','Preview','Plenary',''],
])
p=doc.add_paragraph();r=p.add_run('Total: 180 minutos ✓');r.bold=True;r.font.name='Calibri';r.font.color.rgb=RGBColor(0x2E,0x7D,0x32)
doc.add_page_break()

h(doc,'SET-UP DETALLADO (15 min) — 0:00 a 0:15',2)
h(doc,'Warm-up: "Formula Recall" (10 min)',3)
bp(doc,'Paso 1 — Recreate from memory (5 min)')
ip(doc,'"Good morning everyone. Take out a clean piece of paper. No worksheets — just paper."')
st(doc,'1. "Last session we learned 3 grammar structures. From MEMORY — write the 3 formulas. No looking at PM-2.7. 2 minutes." — Timer: 2 min. Individual.')
bp(doc,'Paso 2 — Volunteers on board (5 min)')
st(doc,'2. "Formula 1?" — Voluntario escribe en tablero.')
st(doc,'3. "Formula 2?" — Voluntario escribe.')
st(doc,'4. "Formula 3?" — Voluntario escribe.')
st(doc,'5. Instructor confirma o corrige.')
st(doc,'6. "Which sentences from Carlos\'s email used these? Tell partner. 1 min." — Timer: 1 min.')
fn(doc,'Si no recuerdan: el Bloque A retoma con Blueprint Model. El warm-up es activación, no evaluación.')
doc.add_paragraph('')
h(doc,'Teacher Talk — Opening Script (3 min)',3)
ip(doc,'"You\'ve read Carlos\'s email. Listened to his call. Learned vocabulary, pronunciation, grammar. TODAY — you write your own Tech Request. You are a developer at SENA. Your workstation has a problem. You write to IT. In English. Let\'s go."')
doc.add_paragraph('')
h(doc,'Objective',3)
ip(doc,'"Today you will: write your own Tech Request email in English — using the structures, vocabulary, and format you\'ve learned."')
icq(doc,'Reading or writing today?','Writing')
icq(doc,'In what language?','English')
doc.add_page_break()

h(doc,'WHILE DETALLADO (150 min) — 0:15 a 2:55',2)
h(doc,'BLOQUE A — DECONSTRUCT (30 min) — 0:15 a 0:45',3)
bp(doc,'Worksheet: ','PM-2.8, Activity 1')
bp(doc,'Agrupación: ','Plenary → Individual → Pairs → Plenary')
bp(doc,'Objetivo: ','Los aprendices analizan la estructura pragmática del email de Carlos.')

doc.add_paragraph('')
h(doc,'Blueprint Model: Display (5 min) — 0:15 a 0:20',3)
st(doc,'1. Proyectar slide 23 (email simplificado de Carlos).')
st(doc,'2. Instructor lee en voz alta señalando cada parte.')
st(doc,'3. "SHORT. CLEAR. PROFESSIONAL. Today you write one just like this — but about YOUR computer."')

doc.add_paragraph('')
h(doc,'Activity 1 — Pragmatic Analysis (20 min) — 0:20 a 0:40',3)
bp(doc,'Instrucciones paso a paso:')
st(doc,'1. "Match each part with its PURPOSE. 3 minutes." — Timer: 3 min. Individual.')
bp(doc,'Answer Key:')
tbl(doc,['Part','Purpose'],[
    ['1. "Hello,"','C. Greeting'],['2. "My workstation has a problem."','D. Introduce the problem'],
    ['3. "The CPU is old."','A. Describe the problem'],['4. "I need a new CPU..."','B. Request what you need'],
    ['5. "Thank you,"','E. Close politely'],
],sc='2E75B6')
st(doc,'')
st(doc,'2. "Compare with partner. 2 minutes." — Timer: 2 min. Pairs.')
st(doc,'3. Revisar ítem por ítem en plenaria.')
st(doc,'4. "IMPORTANT: Why does the email start with the problem, NOT with \'How are you?\' Partner. 1 min."')
st(doc,'5. "Why?" — 2-3 voluntarios.')
st(doc,'6. Instructor: "It\'s a TECH REQUEST. IT person is busy. Problem first. Greeting → Problem → Request → Closing."')
st(doc,'7. Escribir Email Structure en tablero.')
fn(doc,'Hispanohablantes escriben emails largos con "Espero que estés bien." En Tech Request en inglés: directo al problema. "SHORT. Problem first. Always."')
doc.add_page_break()

h(doc,'BREAK — Stretch & Recharge (15 min) — 0:45 a 1:00',2)
doc.add_paragraph('Instructor toma nota de:')
for item in ['Qué aprendices necesitan más apoyo con formato','Verificar Email Structure + Writing Formulas en tablero']:
    p=doc.add_paragraph(item,style='List Bullet')
    for r in p.runs:r.font.name='Calibri';r.font.size=Pt(11)
doc.add_page_break()

h(doc,'BLOQUE B — GUIDED DRAFTING (35 min) — 1:00 a 1:35',3)
bp(doc,'Worksheet: ','PM-2.8, Activity 2')
bp(doc,'Agrupación: ','Plenary → Individual → Plenary')
bp(doc,'Objetivo: ','Los aprendices practican el formato con soporte (Language Bank) antes de escribir libremente.')

doc.add_paragraph('')
h(doc,'Transición (5 min) — 1:00 a 1:05',3)
tt(doc,'Welcome back. You know the structure. Now let\'s PRACTICE filling it in — before you write from scratch.')

doc.add_paragraph('')
h(doc,'Fill the Skeleton (30 min) — 1:05 a 1:35',3)
bp(doc,'Instrucciones paso a paso:')
st(doc,'1. "Different developer. Monitor is old, only VGA. Mouse is broken. Fill in the blanks using Language Bank. 8 minutes." — Timer: 8 min. Individual.')
bp(doc,'Answer Key (8 blanks):')
tbl(doc,['Blank','Answer'],[
    ['1. "workstation _______ a problem"','has'],['2. "monitor _______"','is'],
    ['3. "_______" (second blank)','old'],['4. "It only _______ VGA port"','has'],
    ['5. "_______ _______ _______ with GPU"','This is not compatible'],
    ['6. "keyboard _______ OK"','is'],['7. "mouse _______ broken"','is'],
    ['8. "I _______ a new monitor"','need'],
],sc='2E75B6')
st(doc,'')
st(doc,'2. "Let\'s check." — Blank por blank. Escribir en LIVE ZONE.')
st(doc,'3. "Read the COMPLETE email aloud to partner. 1 minute." — Timer: 1 min. Pairs.')
st(doc,'4. Revisar 2-3 ejemplos con clase. ¿Gramática correcta? ¿Estructura correcta?')
fn(doc,'PUENTE: practican con soporte antes de escribir libremente. Si no puede completar el skeleton, necesita más apoyo en Bloque C.')
tr(doc,'Transition → Bloque C:', '"You filled in the blanks. Now — the real thing. YOUR email. YOUR computer. YOUR problem."')
doc.add_page_break()

h(doc,'BLOQUE C — THE FINAL TASK (50 min) — 1:35 a 2:25',3)
bp(doc,'Worksheet: ','PM-2.8, Activity 3')
bp(doc,'Agrupación: ','Individual (instructor circulates)')
bp(doc,'Objetivo: ','Los aprendices escriben un email completo sobre un computador real.')

doc.add_paragraph('')
h(doc,'Writing Time (30 min) — 1:35 a 2:05',3)
bp(doc,'Instrucciones paso a paso:')
st(doc,'1. "YOUR turn. Write YOUR OWN Tech Request. Real computer. Real problem. Real request."')
st(doc,'2. "Must have: From, To, Subject, Greeting, Problem, Request, Closing."')
st(doc,'3. "Use the Writing Formulas on the board. Use the vocabulary. Use the grammar."')
st(doc,'4. "30 minutes. ALONE. No talking. Go." — Timer: 30 min. Individual. Silencio.')
bp(doc,'Mientras circula, el instructor:')
st(doc,'- Lee por encima del hombro (sin interrumpir)')
st(doc,'- Ayuda bloqueados: "What computer? What\'s wrong? Start with: \'My workstation has a problem.\'"')
st(doc,'- Toma nota mental de errores comunes')
st(doc,'- NO corrige durante escritura — deja que fluya')
fn(doc,'A1 pueden bloquearse ante "página en blanco." Si nada después de 5 min: "Start with the Subject line. What workstation number?" Una vez empieza formato, contenido fluye.')
doc.add_paragraph('')
h(doc,'Auditor\'s Checklist + ✓ Checkpoint (5 min) — 2:05 a 2:10',3)
st(doc,'1. "Self-check. Auditor\'s Checklist at bottom of worksheet. Mark each item. 2 minutes." — Timer: 2 min.')
st(doc,'2. ✓ "Raise hand if ALL 5 items checked." — Contar manos.')
bp(doc,'El Checklist:')
st(doc,'- [ ] Used "is" to describe a problem')
st(doc,'- [ ] Used "has" to describe a feature')
st(doc,'- [ ] Used "I need" for the request')
st(doc,'- [ ] Included Subject line with workstation number')
st(doc,'- [ ] Email has: Greeting → Problem → Details → Request → Closing')

doc.add_paragraph('')
h(doc,'Peer Review (15 min) — 2:10 a 2:25',3)
bp(doc,'Instrucciones paso a paso:')
st(doc,'1. "Exchange email with partner. Read it. Check the Auditor\'s Checklist for THEM. 3 minutes." — Timer: 3 min.')
st(doc,'2. "Feedback: ONE positive + ONE suggestion. Write at bottom of their paper. 2 minutes." — Timer: 2 min.')
st(doc,'3. "Give back. Read feedback. 1 minute." — Timer: 1 min.')
st(doc,'4. "What did your partner do well?" — 2-3 voluntarios.')
st(doc,'5. "What suggestion did you get?" — 2-3 voluntarios.')
fn(doc,'Peer Review es formativo, no sumativo. "This is not a grade. This is help."')
doc.add_paragraph('')
h(doc,'✓ Checkpoint: Submit Draft (5 min) — 2:25 a 2:30',3)
st(doc,'1. "One quick change based on feedback? 2 minutes." — Timer: 2 min.')
st(doc,'2. "Submit final draft to me." — Instructor recoge.')
tr(doc,'Transition → WRAP-UP:', '"Good work. Let\'s close."')
doc.add_page_break()

h(doc,'WRAP-UP DETALLADO (15 min) — 2:30 a 2:45',2)
h(doc,'Exit Ticket (8 min) — 2:30 a 2:38',3)
bp(doc,'Instrucciones:')
st(doc,'1. "On clean paper — ONE sentence from your email with \'is\' and ONE with \'has.\' The exact sentences you wrote. 2 minutes." — Timer: 2 min.')
st(doc,'2. "Show partner. Grammatically correct? 30 sec." — Timer: 30 seg.')
bp(doc,'Ejemplos válidos:')
st(doc,'- IS: "The CPU is old." / "The monitor is small."')
st(doc,'- HAS: "It has a VGA port." / "The workstation has 4 GB of RAM."')
doc.add_paragraph('')
h(doc,'Teacher Talk — Closing Script (5 min) — 2:38 a 2:43',3)
ip(doc,'"You wrote a real email. In English. About a real computer with a real problem. That\'s a real skill.\n\nHomework: Workbook Ch. 5. (1) Revise your email — clean second draft. (2) Read it aloud 3 times for fluency. 60 minutes.\n\nNext session — you SPEAK. Help Desk simulation. Phone conversation. Your email is your preparation. Come ready to talk."')
icq(doc,'What do you do first in homework?','Revise email — clean draft')
icq(doc,'What do you do second?','Read it aloud 3 times')
icq(doc,'Why read aloud?','For the speaking simulation next session')
doc.add_paragraph('')
h(doc,'Preview — 2:43 a 2:45',3)
ip(doc,'"Next session: The Help Desk. Simulation day. Phone conversation in English. Come ready to talk."')
doc.add_page_break()

h(doc,'ANSWER KEY CONSOLIDADO',1)
h(doc,'PM-2.8 — Writing Skills & Pragmatics',2)
tbl(doc,['Actividad','Ítem','Respuesta'],[
    ['Activity 1','1. "Hello,"','C. Greeting'],
    ['Activity 1','2. "workstation has a problem"','D. Introduce the problem'],
    ['Activity 1','3. "CPU is old"','A. Describe the problem'],
    ['Activity 1','4. "I need a new CPU"','B. Request what you need'],
    ['Activity 1','5. "Thank you,"','E. Close politely'],
    ['Activity 2','Blank 1','has'],['Activity 2','Blank 2','is'],
    ['Activity 2','Blank 3','old'],['Activity 2','Blank 4','has'],
    ['Activity 2','Blank 5','This is not compatible'],
    ['Activity 2','Blank 6','is'],['Activity 2','Blank 7','is'],
    ['Activity 2','Blank 8','need'],
],sc='2E75B6')
doc.add_paragraph('')
h(doc,'Activity 3 — Rúbrica para evaluación del instructor',2)
tbl(doc,['Criterio','0 pts','1 pt','2 pts'],[
    ['Formato','Sin formato','Formato parcial','Formato completo'],
    ['Grammar Accuracy','0-1 estructuras','2-3 correctas','4+ correctas'],
    ['Vocabulary Use','0-2 términos','3-5 términos','6+ términos'],
    ['Clarity','Incomprensible','Parcial','Claramente comprensible'],
],sc='2E75B6')
doc.add_page_break()

h(doc,'DIFFERENTIATION NOTES',1)
h(doc,'Fast Finishers',2)
for t,d in [
    ('Fill the Skeleton','"Write SECOND skeleton for GPU old + Ethernet broken."'),
    ('Final Task','"Write SECOND email — you are IT MANAGER responding to Carlos. Approve or reject."'),
    ('Peer tutor','"Guide your partner — ask questions, don\'t write for them."'),
    ('Extra challenge','"Add a budget line. Research real prices online."'),
]:bp(doc,f'{t}: ',d)
doc.add_paragraph('')
h(doc,'More Support Needed',2)
for t,d in [
    ('Skeleton','Language Bank como tarjeta física sobre escritorio.'),
    ('Final Task','Email Template pre-escrito — aprendiz solo llena espacios.'),
    ('Blocked after 10 min','Permitir usar Activity 2 como modelo. Cambiar detalles, mantener estructura.'),
    ('Peer Review','Emparejar niveles similares.'),
]:bp(doc,f'{t}: ',d)
doc.add_page_break()

h(doc,'INSTRUCTOR SELF-CHECK',1)
doc.add_paragraph('')
for i,q in enumerate([
    '¿Todos entregaron un draft? ¿Cuántos cumplen las 5 partes del formato?',
    '¿Hubo momento donde la mayoría parecía perdida? ¿En qué actividad?',
    '¿Los tiempos se cumplieron? ¿Dónde hubo desfases?',
    '¿Qué errores gramaticales fueron más comunes? ¿Necesito reforzar para Session 7?',
    '¿El Peer Review produjo feedback útil? ¿Los aprendices hicieron cambios?',
],1):
    bp(doc,f'{i}. {q}')
    p=doc.add_paragraph('Respuesta: _______________________________________________')
    for r in p.runs:r.font.name='Calibri';r.font.size=Pt(11)
    doc.add_paragraph('')

doc.add_paragraph('')
for txt in ['SESSION 5: WRITE IT RIGHT — BUILD-OUT','ADSO — GUÍA 1: The Hardware Specialist','Sistema de Prompts Maestros — LG Factory — FPI SENA — Bilingüismo','Instructor Sergio Cortés Perdomo · Marzo 2026']:
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(txt);r.font.size=Pt(9);r.font.name='Calibri';r.font.color.rgb=RGBColor(0x7F,0x8C,0x8D)

doc.save('/Users/Beppo/Projects/fpi-sena-factory/guides/ADSO-G1/ADSO — GUÍA 1 — Session 5 — Write It Right — Build-Out.docx')
print('Done')
