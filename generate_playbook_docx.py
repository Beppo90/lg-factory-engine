#!/usr/bin/env python3
"""Generate ADSO Guide 1 Playbook Outline as a Word document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Style defaults ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ── Helper functions ──
def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table_row(table, cells, bold=False, shade=None):
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        if bold:
            run.bold = True
        if shade:
            set_cell_shading(cell, shade)
    return row

def set_first_row(table, cells, shade='4472C4'):
    for i, text in enumerate(cells):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, shade)

def make_table(doc, headers, rows, shade='4472C4'):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_first_row(table, headers, shade)
    for row_data in rows:
        add_table_row(table, row_data)
    return table

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Calibri'
    return h

def add_bold_para(doc, label, value=''):
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    if value:
        run2 = p.add_run(value)
        run2.font.name = 'Calibri'
        run2.font.size = Pt(11)
    return p

def add_teacher_talk(doc, text):
    p = doc.add_paragraph()
    run = p.add_run('Teacher Talk: ')
    run.bold = True
    run.italic = True
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run2 = p.add_run(f'"{text}"')
    run2.italic = True
    run2.font.name = 'Calibri'
    run2.font.size = Pt(11)
    return p

# ════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('ADSO — GUÍA 1: THE HARDWARE SPECIALIST')
run.bold = True
run.font.size = Pt(20)
run.font.name = 'Calibri'
run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Playbook Outline (Session Map)')
run.font.size = Pt(16)
run.font.name = 'Calibri'
run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

doc.add_paragraph('')

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub2.add_run("Instructor's Session Map")
run.font.size = Pt(14)
run.font.name = 'Calibri'

doc.add_paragraph('')
doc.add_paragraph('')

# Warning box
warn = doc.add_paragraph()
warn.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = warn.add_run('⚠️ Este documento es SOLO para el instructor. No distribuir a los aprendices.')
run.bold = True
run.font.size = Pt(12)
run.font.name = 'Calibri'
run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

doc.add_paragraph('')

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('Sistema de Prompts Maestros — LG Factory — FPI SENA — Bilingüismo')
run.font.size = Pt(10)
run.font.name = 'Calibri'
run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

meta2 = doc.add_paragraph()
meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta2.add_run('Instructor Sergio Cortés Perdomo · Marzo 2026')
run.font.size = Pt(10)
run.font.name = 'Calibri'
run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

doc.add_page_break()

# ════════════════════════════════════════
# ENCABEZADO INSTITUCIONAL
# ════════════════════════════════════════

add_heading_styled(doc, 'ENCABEZADO INSTITUCIONAL', level=1)

make_table(doc, ['Campo', 'Dato'], [
    ['Programa', 'Análisis y Desarrollo de Software (ADSO) — 228118'],
    ['Guía', 'Guía 1: The Hardware Specialist'],
    ['Macro-Temática', "The Developer's Ecosystem (Hardware, OS & Environment)"],
    ['Nivel CEFR', 'A1.1 — A1.2'],
    ['Intensidad directa', '24 horas'],
    ['Intensidad autónoma', '6 horas'],
    ['Duración por sesión', '3 horas (180 minutos)'],
    ['Número de sesiones', '8'],
    ['Instructor', 'Sergio Cortés Perdomo'],
])

doc.add_paragraph('')

# ════════════════════════════════════════
# PANORAMA GENERAL
# ════════════════════════════════════════

add_heading_styled(doc, 'PANORAMA GENERAL', level=1)

make_table(doc, 
    ['Session', 'Nombre', 'Worksheets', 'Foco', 'Habilidades', 'Autónomo'],
    [
        ['1', 'The Wake-Up Call', 'PM-2.1 + PM-2.2', 'Motivación: el developer sin PC + diagnóstico de saberes previos', '—', 'Workbook Ch. 1 (45 min)'],
        ['2', 'Read the Request', 'PM-2.3', 'Input pesado: Tech Request Email de Carlos Ramírez', 'R● V○', 'Workbook Ch. 2 (45 min)'],
        ['3', 'Tuning In', 'PM-2.4 + PM-2.5', 'Listening de la llamada + trabajo de vocabulario hardware', 'L● V● R○', 'Workbook Ch. 3 (45 min)'],
        ['4', 'Say It Right, Build It Right', 'PM-2.6 + PM-2.7', 'Pronunciación + gramática: To Be, Have/Has, Demonstratives', 'P● G● V○', 'Workbook Ch. 4 (45 min)'],
        ['5', 'Write It Right', 'PM-2.8', 'Producción escrita: redactar Tech Request propio', 'W● G○ V○ R○', 'Workbook Ch. 5 (60 min)'],
        ['6', 'The Help Desk', 'PM-2.9', 'Simulación oral: Help Desk con problem + stock cards', 'S● V○ G○', 'Workbook Ch. 6 (45 min)'],
        ['7', 'Prove What You Know', 'PM-4.2', 'Cuestionario Técnico (50 pts) + retroalimentación', 'R● L● V● G● W●', 'Workbook Ch. 7 (30 min)'],
        ['8', 'The Full Circle', 'Review + Closure', 'Recuperación, extensión, cierre circular, feedback loop', '—', '—'],
    ],
    shade='1F3A5F'
)

doc.add_page_break()

# ════════════════════════════════════════
# SESSION DETAIL TEMPLATE
# ════════════════════════════════════════

def add_session(doc, num, name, worksheets, duration, focus_skills, support_skills,
                setup_time, warmup, objective, teacher_open,
                while_blocks, wrap_time, exit_ticket, teacher_close, autonomous,
                log_worksheets, log_materials, log_grouping, log_canva, log_plan_b=''):
    
    add_heading_styled(doc, f'SESSION {num}: {name}', level=2)
    
    add_bold_para(doc, 'Worksheets: ', worksheets)
    add_bold_para(doc, 'Duración: ', f'{duration} minutos')
    p = doc.add_paragraph()
    run = p.add_run('Habilidades foco: ')
    run.bold = True
    run.font.name = 'Calibri'
    p.add_run(f'{focus_skills} | ').font.name = 'Calibri'
    run2 = p.add_run('Soporte: ')
    run2.bold = True
    run2.font.name = 'Calibri'
    p.add_run(support_skills).font.name = 'Calibri'
    
    doc.add_paragraph('')
    
    # SET-UP
    add_heading_styled(doc, f'SET-UP ({setup_time} min)', level=3)
    add_bold_para(doc, 'Warm-up: ', warmup)
    add_bold_para(doc, 'Objective: ', objective)
    add_teacher_talk(doc, teacher_open)
    
    doc.add_paragraph('')
    
    # WHILE
    while_total = duration - setup_time - wrap_time
    add_heading_styled(doc, f'WHILE ({while_total} min)', level=3)
    for block_name, block_time, block_desc in while_blocks:
        add_bold_para(doc, f'{block_name} ({block_time} min): ', block_desc)
    
    doc.add_paragraph('')
    
    # WRAP-UP
    add_heading_styled(doc, f'WRAP-UP ({wrap_time} min)', level=3)
    add_bold_para(doc, 'Exit Ticket: ', exit_ticket)
    add_teacher_talk(doc, teacher_close)
    add_bold_para(doc, 'Trabajo Autónomo: ', autonomous)
    
    doc.add_paragraph('')
    
    # LOGISTICS BOX
    add_heading_styled(doc, 'LOGISTICS BOX', level=3)
    log_rows = [
        ['Worksheets', log_worksheets],
        ['Materiales', log_materials],
        ['Agrupación', log_grouping],
        ['Recursos Canva', log_canva],
    ]
    if log_plan_b:
        log_rows.append(['Plan B', log_plan_b])
    make_table(doc, ['Campo', 'Detalle'], log_rows, shade='2E75B6')
    
    doc.add_page_break()


# SESSION 1
add_session(doc, 1, 'THE WAKE-UP CALL',
    'PM-2.1 (The Spark) + PM-2.2 (The Gap Analysis)', 180, '—', '—',
    15,
    'No previous session — first session. Instructor writes on board: "What is the most important part of a computer?" Students think individually (1 min), share with neighbor (2 min). Collect 5-6 answers on board.',
    'Today you will discover what you already know about hardware — and what you need to learn in this guide.',
    'Good morning everyone. Welcome to The Hardware Specialist. Before we start learning, I want to know — what do YOU already know? Today is about discovery.',
    [
        ('Bloque A', 60, 'The Spark (PM-2.1): Read scenario aloud (developer, Friday night, crashed PC, deadline Monday). Students read individually. Controversial Question debate in groups of 3: "If you don\'t know the names of the parts inside a computer in English... can you really call yourself a software developer?" Each group shares conclusion. Activity 1: Personal Connection (rate 1-5). Activity 3: "Need to Know" — write 2 things to learn. Agrupación: Individual → Groups of 3 → Plenary.'),
        ('Bloque B', 15, 'Break / Transition. Teacher Talk: "Ok, great debate. Now let\'s find out what you ACTUALLY know — and what your blind spots are."'),
        ('Bloque C', 60, 'The Gap Analysis (PM-2.2): Activity 1: Brainstorming (write all English computer words). Activity 2: Blind Spots (5 situations — can you say them in English?). Activity 3: Learning Contract (based on blind spots, write what they need to learn). Students sign the contract. Agrupación: Individual → Pairs → Plenary.'),
        ('Bloque D', 15, 'Synthesis: Instructor collects most common blind spots on the board. These become the "targets" for the guide. Teacher Talk: "Look at the board. These are YOUR blind spots. By the end of this guide, you will be able to do ALL of these things."'),
    ],
    15,
    'Write ONE thing you are sure you know about hardware in English, and ONE thing you definitely don\'t know.',
    'Before we go — one thing you know, one thing you don\'t. Next session we read a real tech request from a developer.',
    'Workbook Ch. 1 — My Hardware Profile: draw your real computer at home or SENA, label 5 parts in English (even if wrong), write 1 sentence about each. (45 min estimado)',
    'PM-2.1 (printed) + PM-2.2 (printed)',
    'Whiteboard/markers, printed worksheets',
    'Individual → Groups of 3 → Pairs → Plenary',
    'Slides 1-4 (Scenario image, Controversial Question, Gap Analysis diagram)',
)

# SESSION 2
add_session(doc, 2, 'READ THE REQUEST',
    'PM-2.3 (The Master Anchor — Reading)', 180, 'R●', 'V○',
    15,
    '"Quick Recall" — Students share with partner: "What did you write in your Learning Contract last session? What do you need to learn?" (3 min). Then 3 volunteers share with plenary.',
    'Today you will read a REAL tech request email from a developer. You will extract information, make decisions, and learn the structures you need to write your OWN request.',
    'Good morning everyone. Last session we discovered your blind spots. Today we read a real email from Carlos — a developer who has a problem with his workstation. Your job: understand the problem and decide what to do.',
    [
        ('Bloque A', 20, 'Pre-Reading: Toolbelt (5 key words: Spec Sheet, Upgrade, Compatible, Performance, Budget). Instructor introduces each word with a simple example. Students write in worksheet. Quick check: instructor says word, students say meaning. Agrupación: Plenary.'),
        ('Bloque B', 60, 'While-Reading: First read silently (5 min). Gist question: "What does Carlos want?" (1 min pair discussion). Second read: Activity 1 — Information Transfer table (Current Spec vs. Requested Spec). Individual → pairs → plenary review.'),
        ('Bloque C', 15, 'Break / Transition. Teacher Talk: "You understood what Carlos needs. Now YOU are the IT Manager — and you don\'t have enough money."'),
        ('Bloque D', 40, 'Post-Reading: Activity 2 — Technical Decision. IT Managers with only $800 (not $1,200). Which TWO upgrades are most important? Circle, justify with sentence from text. Groups of 4 compare and debate. Each group presents.'),
        ('Bloque E', 15, 'Activity 3: Deconstruct for Future Output. Find exact sentences Carlos used to: describe a problem, request something, compare two things. These become "writing formulas" for Session 5. Teacher Talk: "Save these sentences. You will use this SAME structure to write your own Tech Request later."'),
    ],
    15,
    "Write Carlos's problem in YOUR words — 2 sentences using 'is' and 'has'.",
    'Before we go — summarize Carlos\'s problem in your own words. Next session we LISTEN to Carlos on the phone.',
    'Workbook Ch. 2 — Reading Extension: re-read email and underline ALL examples of "is" and "has" (count them). Write 3 new sentences about own computer using same structure. (45 min estimado)',
    'PM-2.3 (printed)',
    'Whiteboard/markers, printed worksheet, projector (optional for displaying email)',
    'Individual → Pairs → Groups of 4 → Plenary',
    'Slides 5-7 (Toolbelt words, email on screen, budget decision scenario)',
)

# SESSION 3
add_session(doc, 3, 'TUNING IN',
    'PM-2.4 (The Auditory Anchor — Listening) + PM-2.5 (Vocabulary & Language Function)', 180, 'L● V●', 'R○',
    20,
    '"Quick Recall" — In pairs, each student names 3 hardware components from Carlos\'s email (reciclaje PM-2.3). Then: "Can you say them in a sentence with \'is\' or \'has\'?" Quick round — 6 volunteers share 1 sentence each.',
    'Today you will LISTEN to Carlos on the phone calling IT Support — and learn the 20 key vocabulary words you need for hardware.',
    'Good morning everyone. Last session we READ Carlos\'s email. Today we\'re going to LISTEN to Carlos on the phone — and learn the exact words you need for hardware.',
    [
        ('Bloque A', 50, 'Listening (PM-2.4): Pre-listening: Audio Cues (explain "gig" = Gigabyte, "lemme check" = let me check, is/has sound similar at speed). First listen (gist): "Who is calling? What is the problem?" — individual, pair check. Second listen (detail): Activity 1 — The Triage (match problem with component). Third listen: Activity 2 — Immediate Action (which action first? justify). Activity 3: Noticing & Chunk Extraction — find exact phrases for: greet, describe problem, confirm solution. These become "speaking blueprint" for Session 6. Agrupación: Individual → Pairs → Plenary.'),
        ('Bloque B', 15, 'Break / Transition. Teacher Talk: "Great job with the listening. Now we\'re going to focus on the vocabulary — the words Carlos used on the phone."'),
        ('Bloque C', 60, 'Vocabulary (PM-2.5): Toolbelt visual categorization: 8 key chunks on board. Activity 1: Logical Association (match component with function — 5 items). Activity 2: Gap Fill — The System Log (7 blanks, word bank). Activity 3: Micro-Production — describe SENA workstation using 2 sentences with Toolbelt words. Agrupación: Individual → Small groups → Plenary.'),
        ('Bloque D', 20, 'Reciclaje: Return to Reading text (PM-2.3). Students underline the 20 key terms. Compare: which words appeared in email AND listening? Mini-concordance on board.'),
    ],
    15,
    'Write 3 hardware components and 1 sentence about each using "has" or "is".',
    'Before we go — write down three components and one sentence for each. Next session we work on pronunciation and grammar.',
    'Workbook Ch. 3 — Vocabulary Reinforcement: categorization (Input/Output/Storage/Internal) + crossword + "My Ideal Workstation" paragraph (5 sentences). (45 min estimado)',
    'PM-2.4 (printed) + PM-2.5 (printed)',
    'Audio file (TTS at 0.85x), speakers or individual headphones',
    'Individual → Pairs → Small groups → Plenary',
    'Slides 8-14 (Listening visuals + vocabulary images + categorization chart)',
    'If audio unavailable: Instructor reads script at natural pace, students follow printed transcript',
)

# SESSION 4
add_session(doc, 4, 'SAY IT RIGHT, BUILD IT RIGHT',
    'PM-2.6 (Pronunciation & Speaking Skills) + PM-2.7 (Grammar & Structure Use)', 180, 'P● G●', 'V○',
    15,
    '"Vocabulary Lightning Round" — Instructor points at images on Canva slides, students say the word aloud. 20 words, fast pace. Then: "Which words are HARD to pronounce?" Collect 4-5 difficult words on board.',
    'Today you will learn HOW to pronounce the hardware words correctly AND how to build sentences with the grammar you need.',
    'Good morning everyone. You know the words. Now let\'s make sure you can SAY them and USE them in real sentences. Today is about accuracy.',
    [
        ('Bloque A', 50, 'Pronunciation (PM-2.6): Phonetic Toolbox: 10 words with phonetic guide + stress marks. Activity 1: Perception — count syllables, circle stressed syllable (6 words). Activity 2: Chunking Drill — 5 chunks ("The CPU is", "It has a", "This is an", "faster than", "compatible with"). Say each 3 times fast. Instructor models, students repeat, pairs practice. Activity 3: Voice Command — read IT report aloud to partner. Partner checks intelligibility (YES/PARTIALLY/NO). Swap roles. Agrupación: Plenary → Individual → Pairs.'),
        ('Bloque B', 15, 'Break / Transition. Teacher Talk: "Now you can say the words. Let\'s build the sentences. This is the grammar you need."'),
        ('Bloque C', 60, 'Grammar (PM-2.7): Syntax Blueprint: 3 structures on board (To Be, Have/Has, Demonstratives). Activity 1: Syntax Bugs — 6 sentences, find bugs, select correct version. Individual → pair → plenary. Activity 2: Fill the Ticket — Support Ticket #0042 (IS, HAS, THIS, THAT, THESE, THOSE). Activity 3: Applied Production — 3 sentences about real SENA workstation using IS, HAS, one demonstrative. Agrupación: Individual → Pairs → Plenary.'),
        ('Bloque D', 20, 'Integration: Pronunciation + Grammar combined. Students read their 3 sentences ALOUD to partner. Partner checks: (1) Grammar correct? (2) Pronunciation intelligible? Quick feedback round.'),
    ],
    15,
    'Write one sentence using "is", one using "has", and one using "this/that/these/those" — all about hardware.',
    'Before we go — three sentences, three structures. Next session you WRITE your own Tech Request.',
    'Workbook Ch. 4 — Grammar Drill: 10 fill-in-blank sentences (is/has/demonstratives) + 5 "Syntax Bug" corrections from scratch. (45 min estimado)',
    'PM-2.6 (printed) + PM-2.7 (printed)',
    'Whiteboard/markers, printed worksheets, Canva slides with phonetic guides',
    'Plenary → Individual → Pairs → Plenary',
    'Slides 15-22 (Phonetic toolbox, chunking drill, grammar formulas, Syntax Bugs)',
)

# SESSION 5
add_session(doc, 5, 'WRITE IT RIGHT',
    'PM-2.8 (Writing Skills & Pragmatics)', 180, 'W●', 'G○ V○ R○',
    15,
    '"Formula Recall" — On board, students reconstruct from memory the 3 grammar structures (To Be, Have/Has, Demonstratives). 3 volunteers write them. Then: "Which sentences from Carlos\'s email used these structures?" Quick recall of "writing formulas" from Session 2.',
    'Today you will WRITE your own Tech Request email. By the end of this session, you will have a complete, professional email in English.',
    'Good morning everyone. You\'ve read Carlos\'s email. You\'ve listened to his phone call. You\'ve learned the vocabulary, pronunciation, and grammar. TODAY — you write your own. This is where everything comes together.',
    [
        ('Bloque A', 30, 'Deconstruct (PM-2.8, Activity 1): Blueprint Model: display Carlos\'s simplified email on screen. Activity 1: Pragmatic Analysis — match each part with its purpose (5 items: greeting, introduce problem, describe problem, request, close). Discuss: "Why does the email start with the problem, not with \'How are you\'?" Genre conventions. Agrupación: Individual → Pairs → Plenary.'),
        ('Bloque B', 15, 'Break / Transition. Teacher Talk: "Now you know the structure. Let\'s practice filling it in before you write from scratch."'),
        ('Bloque C', 45, 'Guided Drafting (PM-2.8, Activity 2): Fill the Skeleton — different developer, monitor old, only VGA, mouse broken. Students complete email using Language Bank (has, is, need, not compatible, is old, This is). Instructor reviews 2-3 examples — focus on grammar accuracy and format. Agrupación: Individual → Plenary.'),
        ('Bloque D', 45, 'The Final Task (PM-2.8, Activity 3): Write YOUR OWN Tech Request for a SENA workstation. Complete email: From, To, Subject, Greeting, Problem, Request, Closing. Auditor\'s Checklist for self-check. Instructor circulates. Agrupación: Individual.'),
        ('Bloque E', 15, 'Peer Review: Exchange emails with partner. Partner checks Auditor\'s Checklist: ✓ "is" for problem, ✓ "has" for feature, ✓ "I need" for request, ✓ Subject line, ✓ Greeting → Problem → Details → Request → Closing. Return with 1 positive + 1 suggestion.'),
    ],
    15,
    'Submit the final draft of YOUR Tech Request email to the instructor.',
    'Hand in your email. I\'ll review them and give feedback next session. Next session — you SPEAK. The Help Desk simulation.',
    'Workbook Ch. 5 — Writing Draft: revise Tech Request based on peer feedback. Write SECOND version (clean draft). Practice reading aloud 3 times for fluency (Session 6 prep). (60 min estimado)',
    'PM-2.8 (printed)',
    'Whiteboard/markers, printed worksheet, projector (Blueprint Model)',
    'Individual → Pairs → Plenary → Individual → Pairs',
    'Slides 23-27 (Blueprint Model, Language Bank, Skeleton Structure, Auditor\'s Checklist)',
)

# SESSION 6
add_session(doc, 6, 'THE HELP DESK',
    'PM-2.9 (Speaking Production & Simulation — The Mission)', 180, 'S●', 'V○ G○',
    20,
    '"Chunk Recall" — Students recall speaking chunks from Session 3: Greet, Describe Problem, Confirm Solution. Write on board from memory. Review Skeleton Script (Opening, Describing, Asking, Confirming, Closing). Practice each chunk as class.',
    'Today you will have a REAL phone conversation in English. You are either a Developer with a problem or IT Support with the solution. The problem only gets solved when BOTH people agree.',
    'Good morning everyone. Today is simulation day. You\'ve read, listened, learned vocabulary, practiced pronunciation, written an email — now you TALK. The goal is NOT perfect English. The goal is: your partner UNDERSTANDS and the problem GETS SOLVED.',
    [
        ('Bloque A', 20, 'Briefing: Distribute Problem Cards and Stock Cards. Explain roles (Role A: Developer, Role B: IT Support). Read Skeleton Script together. Rules: (1) Problem solved only when both agree, (2) Try English first, (3) Can use gestures/images/Spanish if stuck. Demo round with volunteer. Agrupación: Plenary.'),
        ('Bloque B', 50, 'Round 1 — Cue Cards FACE-UP. Pairs perform simulation. Instructor circulates, takes notes. After 10 min, early finishers help other pairs. Minimum 2 rounds per pair (swap cards). Agrupación: Pairs.'),
        ('Bloque C', 15, 'Break / Debrief Round 1. Teacher Talk: "Good work. What was HARD? What was EASY?" Collect 3-4 responses. Quick correction of 2-3 common errors.'),
        ('Bloque D', 40, 'Round 2 — Cue Cards FACE-DOWN (from memory). New pairs. Same or swapped roles. No card support — rely on memory and Skeleton Script. Instructor observes improvement. Agrupación: Pairs.'),
        ('Bloque E', 20, 'Round 3 — Role Swap. Developers become IT Support and vice versa. New cards. Final round face-down for challenge. Agrupación: Pairs.'),
    ],
    15,
    'Communication Deliverable checklist: ✓ Developer described ALL problems, ✓ IT Support asked at least ONE clarifying question, ✓ Both agreed on NOW vs. ORDER, ✓ Conversation had Greeting → Problem → Details → Solution → Closing.',
    'You did it. You had a real phone conversation in English about hardware. Next session — the quiz. Prove what you know.',
    'Workbook Ch. 6 — Simulation Prep: reflection paragraph (5 sentences) about what was easy/hard. Review ALL vocab and grammar for quiz. (45 min estimado)',
    'PM-2.9 (printed: Problem Cards + Stock Cards + Skeleton Script)',
    'Printed cue cards (cut before class), timer visible on screen',
    'Plenary → Pairs → Pairs → Pairs',
    'Slides 28-31 (Mission brief, Skeleton Script, observation criteria)',
    'If class is too large for pair simulation: Use fishbowl format — 2 volunteers perform while rest observe with checklist. Rotate every 5 min.',
)

# SESSION 7
add_session(doc, 7, 'PROVE WHAT YOU KNOW',
    'PM-4.2 (Cuestionario Técnico — IE-01)', 180, 'R● L● V● G● W●', '—',
    15,
    '"Quick Review" — 5 rapid-fire questions on board (1 per PM covered). Students write answers, then reveal. No grading — just activation. Example: "What is the brain of the computer?" / "Write a sentence with \'has\' and a hardware component."',
    'Today is the Cuestionario Técnico. 50 points. 5 sections. Show me what you\'ve learned.',
    'Ok everyone. Today is quiz day. You\'ve worked hard for 6 sessions. Now prove what you know. Read carefully. Use your time well. If you finish early, check your answers.',
    [
        ('Bloque A', 120, 'Cuestionario Técnico (PM-4.2): 5 sections, 50 points. Section 1: Reading Comprehension (10 pts). Section 2: Writing Task (10 pts). Section 3: Listening Comprehension (10 pts — audio played 3 times). Section 4: Key Vocabulary HOTS (10 pts). Section 5: Grammar & Structure HOTS (10 pts). Instructor plays audio for Section 3, circulates. Strict silence. Agrupación: Individual.'),
        ('Bloque B', 30, 'Immediate Feedback: Instructor reviews key answers (focus on most common errors from circulation). Students self-assess Sections 4 and 5 where possible. Collect quizzes for formal grading. Teacher Talk: "Let\'s look at the most common mistakes. This is how you learn — not from the grade, but from understanding what went wrong."'),
    ],
    15,
    'Write ONE thing you learned in this guide that you will use in your career as a developer.',
    'Next session is the last one. We close the circle — review everything, and you tell me what worked and what didn\'t.',
    'Workbook Ch. 7 — Post-Quiz Review: correct errors from feedback. Write 3 sentences about what to do differently next time. (30 min estimado)',
    'PM-4.2 (printed: Cuestionario Técnico — one per student)',
    'Audio file for Section 3 (TTS), speakers, printed quizzes, answer sheets',
    'Individual (strict exam conditions) → Plenary (feedback)',
    'Slides 32-33 (Quiz instructions, timer)',
    'If audio unavailable for Section 3: Instructor reads listening script at natural pace. Students follow printed transcript (remove after second reading).',
)

# SESSION 8
add_session(doc, 8, 'THE FULL CIRCLE',
    'None (Review + Closure session)', 180, '—', '—',
    15,
    '"One Word" — Each student says ONE word in English learned in this guide. No repeats. Go around the room. Write all words on board.',
    'Today we close the circle. You review, you reflect, and you tell me what this guide gave you — and what it didn\'t.',
    'Good morning everyone. Last session. Look at all those words on the board — YOU learned them. Today we make sure they stick.',
    [
        ('Bloque A', 40, 'Vocabulary & Grammar Recovery: (1) Component identification — instructor points at images, students say word + sentence. (2) Grammar relay — teams of 4 compete to write most correct sentences using To Be, Have/Has, Demonstratives in 5 min. (3) Pronunciation check — 5 "hard" words from Session 4. Agrupación: Plenary → Teams → Individual.'),
        ('Bloque B', 15, 'Break / Transition. Teacher Talk: "Now let\'s see how far you\'ve come since Session 1."'),
        ('Bloque C', 40, 'Gap Analysis Revisited (reciclaje PM-2.2): Students look at original Blind Spots from Session 1. For each: can they do it NOW? Check ✓ or still ❓. Fill Learning Contract "completion" column. Compare with partner. Celebrate growth. Agrupación: Individual → Pairs → Plenary.'),
        ('Bloque D', 40, 'Extension & Application: Scenario — "You are at your first job. Your supervisor says: \'The new developer starts Monday. Set up their workstation. Write me the spec sheet and call IT if you need parts.\' What do you do?" Groups of 4 plan steps, present in English (2 min per group). Integrates ALL skills. Agrupación: Groups of 4 → Plenary.'),
        ('Bloque E', 15, 'Feedback Loop: Students complete Guide Feedback Form (anonymous). What worked? What didn\'t? What would you change?'),
    ],
    15,
    'None (final session).',
    'You came in not knowing how to name the parts of a computer in English. Now you can read a tech request, listen to a phone call, write an email, and have a conversation about hardware. That\'s real. That\'s yours. Well done.',
    'None.',
    'PM-2.2 (revisit original printed worksheet from Session 1)',
    'Whiteboard/markers, images of hardware components, Guide Feedback Forms (printed)',
    'Plenary → Teams → Individual → Pairs → Groups of 4 → Plenary',
    'Slides 34-38 (Review activities, Gap Analysis revisit, extension scenario)',
)

# ════════════════════════════════════════
# SKILLS PROGRESSION MAP
# ════════════════════════════════════════

add_heading_styled(doc, 'SKILLS PROGRESSION MAP', level=1)

p = doc.add_paragraph()
run = p.add_run('● = Habilidad foco de la sesión | ○ = Habilidad de soporte/reciclaje | — = No activa')
run.font.size = Pt(9)
run.font.name = 'Calibri'
run.italic = True

p2 = doc.add_paragraph()
run2 = p2.add_run('R = Reading | L = Listening | V = Vocabulary | P = Pronunciation | G = Grammar | W = Writing | S = Speaking')
run2.font.size = Pt(9)
run2.font.name = 'Calibri'
run2.italic = True

make_table(doc,
    ['Session', 'R', 'L', 'V', 'P', 'G', 'W', 'S'],
    [
        ['1', '—', '—', '—', '—', '—', '—', '—'],
        ['2', '●', '—', '○', '—', '—', '—', '—'],
        ['3', '○', '●', '●', '—', '—', '—', '—'],
        ['4', '—', '—', '○', '●', '●', '—', '—'],
        ['5', '○', '—', '○', '—', '○', '●', '—'],
        ['6', '—', '—', '—', '—', '—', '○', '●'],
        ['7', '●', '●', '●', '—', '●', '●', '—'],
        ['8', '—', '—', '○', '○', '○', '—', '—'],
    ],
    shade='1F3A5F'
)

doc.add_page_break()

# ════════════════════════════════════════
# MAPA DE TRABAJO AUTÓNOMO
# ════════════════════════════════════════

add_heading_styled(doc, 'MAPA DE TRABAJO AUTÓNOMO', level=1)

make_table(doc,
    ['Asignado en', 'Revisado en', 'Actividad', 'Ref. Workbook', 'Tiempo'],
    [
        ['Session 1', 'Session 2', 'My Hardware Profile: draw real computer, label 5 parts in English, write 1 sentence each', 'Ch. 1', '45 min'],
        ['Session 2', 'Session 3', 'Reading Extension: underline all "is"/"has" in email, write 3 new sentences about own computer', 'Ch. 2', '45 min'],
        ['Session 3', 'Session 4', 'Vocabulary Reinforcement: categorization + crossword + "My Ideal Workstation" paragraph (5 sentences)', 'Ch. 3', '45 min'],
        ['Session 4', 'Session 5', 'Grammar Drill: 10 fill-in-blank + 5 Syntax Bug corrections', 'Ch. 4', '45 min'],
        ['Session 5', 'Session 6', 'Writing Draft: revise Tech Request (clean draft) + read aloud 3 times for fluency', 'Ch. 5', '60 min'],
        ['Session 6', 'Session 7', 'Simulation Prep: reflection paragraph (5 sentences) + review all vocab/grammar for quiz', 'Ch. 6', '45 min'],
        ['Session 7', 'Session 8', 'Post-Quiz Review: correct errors + write 3 sentences about improvement', 'Ch. 7', '30 min'],
    ],
    shade='1F3A5F'
)

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Total trabajo autónomo: 315 min ≈ 5.25 horas (cerca de las 6 horas asignadas; el resto se absorbe en repaso informal)')
run.italic = True
run.font.size = Pt(10)
run.font.name = 'Calibri'

doc.add_page_break()

# ════════════════════════════════════════
# LISTA MAESTRA DE MATERIALES
# ════════════════════════════════════════

add_heading_styled(doc, 'LISTA MAESTRA DE MATERIALES', level=1)

add_heading_styled(doc, 'Impresos', level=2)
impresos = [
    'PM-2.1: The Spark (1 por estudiante)',
    'PM-2.2: The Gap Analysis (1 por estudiante)',
    'PM-2.3: The Master Anchor — Reading (1 por estudiante)',
    'PM-2.4: The Auditory Anchor — Listening (1 por estudiante)',
    'PM-2.5: Vocabulary & Language Function (1 por estudiante)',
    'PM-2.6: Pronunciation & Speaking Skills (1 por estudiante)',
    'PM-2.7: Grammar & Structure Use (1 por estudiante)',
    'PM-2.8: Writing Skills & Pragmatics (1 por estudiante)',
    'PM-2.9: Problem Cards (3 diseños, cortar antes de clase) — 1 set por pareja',
    'PM-2.9: Stock Cards (3 diseños, cortar antes de clase) — 1 set por pareja',
    'PM-2.9: Skeleton Script (1 por estudiante)',
    'PM-4.2: Cuestionario Técnico (1 por estudiante)',
    'Guide Feedback Forms — Session 8 (1 por estudiante)',
]
for item in impresos:
    p = doc.add_paragraph(item, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(10)

add_heading_styled(doc, 'Digitales', level=2)
digitales = [
    'Audio file PM-2.4: Phone call TTS (0.85x speed, 2 voices, pauses)',
    'Audio file PM-4.2 Section 3: Listening TTS (0.85x speed)',
    'Canva presentation (38 slides total, por sesión)',
]
for item in digitales:
    p = doc.add_paragraph(item, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(10)

add_heading_styled(doc, 'Equipamiento', level=2)
equipamiento = [
    'Proyector o pantalla para Canva slides',
    'Speakers or individual headphones (Session 3 y 7)',
    'Whiteboard + markers',
]
for item in equipamiento:
    p = doc.add_paragraph(item, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(10)

add_heading_styled(doc, 'Aula', level=2)
aula = [
    'Disposición para trabajo en parejas (mesas facing o frente a frente)',
    'Espacio para grupos de 3-4 (Session 1, 4, 8)',
    'Área de simulación clara para Session 6 (pares simultáneos sin interferencia)',
]
for item in aula:
    p = doc.add_paragraph(item, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(10)

doc.add_paragraph('')

# ════════════════════════════════════════
# FOOTER
# ════════════════════════════════════════

footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('ADSO — GUÍA 1: THE HARDWARE SPECIALIST — PLAYBOOK OUTLINE (SESSION MAP)')
run.font.size = Pt(9)
run.font.name = 'Calibri'
run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

footer2 = doc.add_paragraph()
footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = footer2.add_run('Sistema de Prompts Maestros — LG Factory — FPI SENA — Bilingüismo')
run2.font.size = Pt(9)
run2.font.name = 'Calibri'
run2.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

footer3 = doc.add_paragraph()
footer3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = footer3.add_run('Instructor Sergio Cortés Perdomo · Marzo 2026')
run3.font.size = Pt(9)
run3.font.name = 'Calibri'
run3.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

# ════════════════════════════════════════
# SAVE
# ════════════════════════════════════════

output_path = '/Users/Beppo/Projects/fpi-sena-factory/guides/ADSO-G1/ADSO — GUÍA 1 — The Hardware Specialist — Playbook Outline (Session Map).docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')