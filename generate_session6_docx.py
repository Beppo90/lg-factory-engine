#!/usr/bin/env python3
"""Generate Session 6 Build-Out as a Word document."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
for s in doc.sections:
    s.top_margin=Cm(2.5);s.bottom_margin=Cm(2.5);s.left_margin=Cm(2.5);s.right_margin=Cm(2.5)
style=doc.styles['Normal'];style.font.name='Calibri';style.font.size=Pt(11)

def sh(cell,color):
    e=OxmlElement('w:shd');e.set(qn('w:fill'),color);e.set(qn('w:val'),'clear')
    cell._tc.get_or_add_tcPr().append(e)

def tbl(doc,headers,rows,sc='1F3A5F'):
    t=doc.add_table(rows=1,cols=len(headers));t.style='Table Grid';t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i];c.text='';r=c.paragraphs[0].add_run(str(h))
        r.font.size=Pt(10);r.font.name='Calibri';r.bold=True;r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF);sh(c,sc)
    for rd in rows:
        row=t.add_row()
        for i,v in enumerate(rd):
            c=row.cells[i];c.text='';r=c.paragraphs[0].add_run(str(v));r.font.size=Pt(9);r.font.name='Calibri'
    return t

def h(d,t,l=1):
    hd=d.add_heading(t,level=l)
    for r in hd.runs:r.font.name='Calibri'

def bp(d,l,v=''):
    p=d.add_paragraph();r=p.add_run(l);r.bold=True;r.font.name='Calibri';r.font.size=Pt(11)
    if v:r2=p.add_run(v);r2.font.name='Calibri';r2.font.size=Pt(11)

def tt(d,t):
    p=d.add_paragraph();r=p.add_run('Teacher Talk: ');r.bold=True;r.italic=True;r.font.name='Calibri'
    r2=p.add_run(f'"{t}"');r2.italic=True;r2.font.name='Calibri'

def fn(d,t):
    p=d.add_paragraph();r=p.add_run('💡 ');r.font.name='Calibri'
    r2=p.add_run(t);r2.italic=True;r2.font.name='Calibri';r2.font.color.rgb=RGBColor(0x2E,0x75,0xB6)

def ck(d,t):
    p=d.add_paragraph();r=p.add_run('✓ Checkpoint: ');r.bold=True;r.font.name='Calibri';r.font.color.rgb=RGBColor(0x2E,0x7D,0x32)
    r2=p.add_run(t);r2.font.name='Calibri'

def icq(d,q,e):
    p=d.add_paragraph();r=p.add_run('ICQ: ');r.bold=True;r.font.name='Calibri'
    r2=p.add_run(f'"{q}"');r2.font.name='Calibri'
    r3=p.add_run(f' → Esperan: "{e}"');r3.italic=True;r3.font.name='Calibri';r3.font.color.rgb=RGBColor(0x7F,0x8C,0x8D)

def ip(d,t):
    p=d.add_paragraph();r=p.add_run(t);r.italic=True;r.font.name='Calibri';r.font.size=Pt(11)

def st(d,t):
    p=d.add_paragraph(t)
    for r in p.runs:r.font.name='Calibri';r.font.size=Pt(11)

def tr(d,l,t):
    p=d.add_paragraph();r=p.add_run(f'{l} ');r.bold=True;r.font.name='Calibri'
    r2=p.add_run(t);r2.italic=True;r2.font.name='Calibri'

# TITLE
for _ in range(4):doc.add_paragraph('')
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('SESSION 6: THE HELP DESK');r.bold=True;r.font.size=Pt(20);r.font.name='Calibri';r.font.color.rgb=RGBColor(0x1F,0x3A,0x5F)
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('ADSO — GUÍA 1: The Hardware Specialist — Build-Out');r.font.size=Pt(14);r.font.name='Calibri';r.font.color.rgb=RGBColor(0x44,0x72,0xC4)
doc.add_paragraph('')
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('⚠️ Este documento es SOLO para el instructor. No distribuir a los aprendices.');r.bold=True;r.font.size=Pt(12);r.font.name='Calibri';r.font.color.rgb=RGBColor(0xC0,0x39,0x2B)
doc.add_paragraph('')
for txt in ['Sistema de Prompts Maestros — LG Factory — FPI SENA — Bilingüismo','Instructor Sergio Cortés Perdomo · Marzo 2026']:
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(txt);r.font.size=Pt(10);r.font.name='Calibri';r.font.color.rgb=RGBColor(0x7F,0x8C,0x8D)
doc.add_page_break()

h(doc,'SESSION HEADER',1)
tbl(doc,['Campo','Dato'],[['Programa','ADSO — 228118'],['Guía','Guía 1: The Hardware Specialist'],['Session','6: The Help Desk'],['Worksheets','PM-2.9 (Speaking Simulation)'],['Duración','180 minutos'],['Habilidades foco','S●'],['Habilidades soporte','V○ G○'],['Trabajo autónomo','Workbook Ch. 6 — Simulation Prep (45 min)'],['Siguiente sesión','Session 7: Prove What You Know (PM-4.2)'],['Plan B','Fishbowl si clase >25: 2 actúan, resto observa con checklist']])
doc.add_paragraph('')

h(doc,'MATERIALS CHECKLIST',1)
p=doc.add_paragraph('Marcar ANTES de entrar al aula:');p.runs[0].italic=True
for item in ['PM-2.9: Problem Cards — IMPRESOS y CORTADOS (3 diseños × copias suficientes)','PM-2.9: Stock Cards — IMPRESOS y CORTADOS (3 diseños × copias suficientes)','PM-2.9: Skeleton Script — impresos (1 por estudiante)','PM-2.9: Communication Deliverable — impresos (1 por estudiante)','Canva slides 28-31 abiertas y listas','Proyector encendido','Tablero preparado según Board Plan','Markers (negro + verde)','Timer visible','Reloj de pared visible']:
    p=doc.add_paragraph(f'☐  {item}')
    for r in p.runs:r.font.name='Calibri';r.font.size=Pt(11)
doc.add_paragraph('')
bp(doc,'NOTA — Preparación de tarjetas: ','Cortar ANTES de clase. Mezclar combinaciones (no siempre Stock A + Problem A). Tener sets para MITAD Developer + MITAD IT Support.')
doc.add_page_break()

h(doc,'BOARD PLAN',1)
bt=('SESSION 6: THE HELP DESK\n'
    'Today: "Have a REAL phone conversation in English"\n\n'
    '[SKELETON SCRIPT]              [RULES]\n'
    'OPENING:                       1. Both must agree\n'
    '"IT Support, this is [name]."  2. Try English FIRST\n'
    '"This is [name] from [dept]."  3. Gestures/Spanish OK if stuck\n'
    '                                4. Goal: understanding\n'
    'DESCRIBING:\n'
    '"My workstation has a problem."\n'
    '"The [X] is [adjective]."     [LIVE ZONE — vacío]\n'
    '"It only has [spec]."          (errores comunes para debrief)\n\n'
    'ASKING:\n'
    '"What about the [X]?"\n'
    '"Is it an SSD or an HDD?"\n\n'
    'CONFIRMING:\n'
    '"Got it. So... new [X]."\n'
    '"I\'ll send the request today."\n\n'
    'CLOSING:\n'
    '"Thanks, you\'re a lifesaver."\n'
    '"No problem. Have a good one."')
p=doc.add_paragraph();r=p.add_run(bt);r.font.name='Consolas';r.font.size=Pt(9)
doc.add_page_break()

h(doc,'MINUTE-BY-MINUTE TIMELINE',1)
tbl(doc,['Tiempo','Dur.','Bloque','Actividad','Agrupación','Notas'],[
    ['0:00-0:08','8','SET-UP','Chunk Recall','Ind→Plenary','Memory'],
    ['0:08-0:15','7','SET-UP','Skeleton Script review','Plenary','Practice chunks'],
    ['0:15-0:20','5','SET-UP','Opening + Rules','Plenary',''],
    ['0:20-0:30','10','WHILE-A','Briefing: roles + cards','Plenary',''],
    ['0:30-0:35','5','WHILE-A','Demo round','Plenary','Volunteer'],
    ['0:35-0:40','5','WHILE-A','✓ Checkpoint','Plenary','Questions'],
    ['0:40-0:45','5','WHILE-A','Form pairs + distribute','Pairs',''],
    ['0:45-1:15','30','WHILE-B','Round 1: FACE-UP','Pairs','2+ rounds'],
    ['1:15-1:20','5','WHILE-B','Debrief R1','Plenary','Hard? Easy?'],
    ['1:20-1:35','15','BREAK','BREAK','—',''],
    ['1:35-1:40','5','WHILE-C','Round 2 intro + new pairs','Pairs','Face-DOWN'],
    ['1:40-2:10','30','WHILE-C','Round 2: FACE-DOWN','Pairs','From memory'],
    ['2:10-2:15','5','WHILE-C','Debrief R2','Plenary','Improvement?'],
    ['2:15-2:20','5','WHILE-D','Round 3: role swap','Pairs','New cards'],
    ['2:20-2:40','20','WHILE-D','Round 3: simulation','Pairs','Final round'],
    ['2:40-2:45','5','WHILE-D','Final debrief','Plenary','Celebration'],
    ['2:45-2:53','8','WRAP-UP','Deliverable self-assess','Individual','4 items'],
    ['2:53-2:58','5','WRAP-UP','Closing + Autónomo','Plenary','Workbook Ch. 6'],
    ['2:58-3:00','2','WRAP-UP','Preview','Plenary',''],
])
p=doc.add_paragraph();r=p.add_run('Total: 180 minutos ✓');r.bold=True;r.font.name='Calibri';r.font.color.rgb=RGBColor(0x2E,0x7D,0x32)
doc.add_page_break()

h(doc,'SET-UP DETALLADO (20 min) — 0:00 a 0:20',2)
h(doc,'Warm-up: "Chunk Recall" (8 min)',3)
bp(doc,'Paso 1 — Write from memory (4 min)')
ip(doc,'"Good morning everyone. Take out a clean piece of paper."')
st(doc,'1. "Session 3 — we extracted Speaking Scripts. From MEMORY write the phrases for 3 functions: Greet, Describe problem, Confirm solution. 3 minutes." — Timer: 3 min.')
bp(doc,'Paso 2 — Check on board (4 min)')
st(doc,'2. Voluntarios escriben en tablero. Instructor confirma.')
fn(doc,'Si no recuerdan: Skeleton Script review las retoma. Recall es activación.')
doc.add_paragraph('')
h(doc,'Skeleton Script Review (7 min)',3)
st(doc,'1. Proyectar slide 28 (Skeleton Script completo).')
st(doc,'2. Instructor lee sección por sección. Clase repite 2 veces cada frase:')
st(doc,'   OPENING: "IT Support, this is [name]."')
st(doc,'   DESCRIBING: "My workstation has a problem."')
st(doc,'   ASKING: "What about the storage?"')
st(doc,'   CONFIRMING: "Got it. So... new SSD, check."')
st(doc,'   CLOSING: "No problem. Have a good one."')
st(doc,'3. "Practice ALL sections with partner. Read aloud. 1 min." — Timer: 1 min.')
doc.add_paragraph('')
h(doc,'Teacher Talk + Rules (5 min)',3)
ip(doc,'"Today is simulation day. Real phone conversation. Role A: Developer — has a problem. Role B: IT Support — has the solution. Problem solved only when BOTH agree. Goal: NOT perfect English. Goal: your partner UNDERSTANDS."')
st(doc,'4 Rules: (1) Both agree (2) English first (3) Gestures/Spanish OK if stuck (4) Understanding, not perfection')
icq(doc,'Is the goal perfect English?','No')
icq(doc,'What is the goal?','Understanding / problem gets solved')
doc.add_page_break()

h(doc,'WHILE DETALLADO (145 min) — 0:20 a 2:55',2)
h(doc,'BLOQUE A — BRIEFING + DEMO (20 min) — 0:20 a 0:40',3)
bp(doc,'Objetivo: ','Los aprendices entienden roles, tarjetas y dinámica antes de empezar.')
doc.add_paragraph('')
h(doc,'Briefing: Roles + Cards (10 min) — 0:20 a 0:30',3)
st(doc,'1. Proyectar slide 29 (Roles).')
bp(doc,'Role A — Developer: ','"Workstation has problems. PROBLEM CARD tells you what\'s wrong. Describe clearly. Request upgrades."')
bp(doc,'Role B — IT Support: ','"STOCK CARD shows what\'s AVAILABLE and what must be ORDERED. Listen, diagnose, confirm, suggest."')
st(doc,'2. Proyectar slide 30 (sample cards).')
st(doc,'3. Instructor lee Problem Card A y Stock Card A en voz alta.')
st(doc,'4. "See? Developer needs CPU — NOT in stock. Must decide: order and wait, or replace SSD/RAM NOW. THIS is where conversation happens."')
fn(doc,'Mismatch INTENCIONAL fuerza negociación real.')
doc.add_paragraph('')
h(doc,'Demo Round (5 min) — 0:30 a 0:35',3)
st(doc,'1. "I\'ll be IT Support. Who wants Developer?" — 1 voluntario.')
st(doc,'2. 2-3 min simulación frente a clase.')
st(doc,'3. "See? Greeting → Problem → Details → Solution → Closing. Now YOU do it."')
doc.add_paragraph('')
ck(doc,'"Questions? Rule 1?" → Both agree | "Rule 2?" → English first | "Rule 4?" → Understanding')
doc.add_page_break()

h(doc,'BLOQUE B — ROUND 1: FACE-UP (40 min) — 0:40 a 1:20',3)
bp(doc,'Objetivo: ','Simulación con soporte de tarjetas visibles — confianza y fluidez básica.')
doc.add_paragraph('')
h(doc,'Form Pairs + Distribute (5 min) — 0:40 a 0:45',3)
st(doc,'1. "Pair up." — 1 min. Instructor distribuye tarjetas (mezclar combinaciones).')
st(doc,'2. "Read your card. DON\'T show partner. 1 min." — Timer: 1 min.')
st(doc,'3. "Round 1: Cards FACE-UP. You CAN look. 2 rounds per pair, ~7-8 min each. Then swap cards with another pair."')
doc.add_paragraph('')
h(doc,'Simulation (30 min) — 0:45 a 1:15',3)
bp(doc,'Mientras circula, el instructor:')
st(doc,'- Escucha conversaciones, NO interrumpe')
st(doc,'- Toma nota de 3-4 errores comunes + 2-3 cosas positivas')
st(doc,'- Usa Observation Criteria: identifica componentes? Usa is/has? Inteligible? Llega a resolución?')
bp(doc,'Timing: ','Ronda 1 (7-8 min) → Intercambio tarjetas (1 min) → Ronda 2 (7-8 min)')
fn(doc,'Primeras conversaciones serán torpes. ESTÁ BIEN. Ronda 2 será mejor. La mejora visible es el objetivo.')
doc.add_paragraph('')
h(doc,'Quick Debrief (5 min) — 1:15 a 1:20',3)
st(doc,'1. "What was HARD?" — 3-4 respuestas. Escribir en LIVE ZONE.')
st(doc,'2. "What was EASY?" — 2-3 respuestas.')
st(doc,'3. Corregir 2-3 errores comunes: "The CPU HAVE → HAS" / "I need buy → I need TO buy" / "Have a one good → Have a good one"')
tr(doc,'Transition → BREAK:','"Take a break. 15 minutes. Come back for Round 2 — NO cards."')
doc.add_page_break()

h(doc,'BREAK — Stretch & Recharge (15 min) — 1:20 a 1:35',2)
doc.add_paragraph('Instructor nota de:')
for item in ['Errores más comunes','Parejas más fluidas (para Round 2 agrupación)','Aprendices con más dificultad']:
    p=doc.add_paragraph(item,style='List Bullet')
    for r in p.runs:r.font.name='Calibri';r.font.size=Pt(11)
doc.add_page_break()

h(doc,'BLOQUE C — ROUND 2: FACE-DOWN (40 min) — 1:35 a 2:15',3)
bp(doc,'Objetivo: ','Simulación sin tarjetas — memoria y fluidez.')
doc.add_paragraph('')
h(doc,'New Pairs + Setup (5 min) — 1:35 a 1:40',3)
tt(doc,'Welcome back. Round 1 was practice. Round 2: NEW partner. Cards FACE-DOWN. You can look ONCE for 30 seconds if needed. But try without.')
st(doc,'1. "New partner — not same as Round 1." — 1 min.')
st(doc,'2. Nuevas tarjetas. "30 seconds to read. Then face-down." — Timer: 30 seg.')
st(doc,'3. "From memory. 2 rounds, 8 min each."')
doc.add_paragraph('')
h(doc,'Simulation (30 min) — 1:45 a 2:15',3)
st(doc,'- Circula y observa. Compara con R1: ¿mejora?')
st(doc,'- Toma nota de errores que persisten + mejoras significativas')
bp(doc,'Timing: ','Ronda 1 (8 min) → Intercambio (1 min) → Ronda 2 (8 min)')
fn(doc,'Mejora DRAMÁTICA entre R1 y R2. Los aprendices confían en memoria y descubren que pueden más de lo que pensaban.')
doc.add_paragraph('')
h(doc,'Quick Debrief (5 min) — 2:10 a 2:15',3)
st(doc,'1. "Round 2 better than Round 1? Thumbs up/sideways/down." — Contar.')
st(doc,'2. "What improved?" — 2-3 respuestas.')
st(doc,'3. Celebrar mejoras: "Much better \'compatible.\' People using \'I need\' instead of \'I want.\'"')
st(doc,'4. "Last round. Role SWAP. Developers → IT Support. IT Support → Developers. New cards."')
doc.add_page_break()

h(doc,'BLOQUE D — ROUND 3: ROLE SWAP (30 min) — 2:15 a 2:45',3)
bp(doc,'Objetivo: ','Ambos roles — empatía y comprensión bidireccional.')
doc.add_paragraph('')
h(doc,'Role Swap + New Cards (5 min) — 2:15 a 2:20',3)
st(doc,'1. "Swap roles. Developer → IT Support. IT Support → Developer."')
st(doc,'2. Nuevas tarjetas. 30 seg para leer. Face-down.')
st(doc,'3. "Final round. Best one yet. 8 minutes. Go."')
doc.add_paragraph('')
h(doc,'Simulation (20 min) — 2:20 a 2:40',3)
st(doc,'- Típicamente la MEJOR ronda de las tres')
st(doc,'- Timing: Ronda 1 (8 min) → Ronda 2 si tiempo (8 min)')
doc.add_paragraph('')
h(doc,'Final Debrief (5 min) — 2:40 a 2:45',3)
st(doc,'1. "Three rounds. Which was your BEST?" — 3-4 voluntarios.')
st(doc,'2. "What surprised you about IT Support role?" — 2-3 respuestas.')
st(doc,'3. "You had real phone conversations in English about hardware. Three times. Different partners. Two roles. That\'s not easy — and you did it. Well done."')
fn(doc,'Feedback final debe ser CELEBRATORIO. Reconocer esfuerzo, no solo resultado.')
doc.add_page_break()

h(doc,'WRAP-UP DETALLADO (15 min) — 2:45 a 3:00',2)
h(doc,'Communication Deliverable Self-Assessment (8 min) — 2:45 a 2:53',3)
bp(doc,'Instrucciones:')
st(doc,'1. "Open PM-2.9 — Communication Deliverable checklist. Self-assess. Did YOUR team achieve these 4 things?"')
bp(doc,'Checklist:')
st(doc,'- [ ] Developer described ALL problems clearly')
st(doc,'- [ ] IT Support asked at least ONE clarifying question')
st(doc,'- [ ] Both agreed on NOW vs. ORDER')
st(doc,'- [ ] Conversation had: Greeting → Problem → Details → Solution → Closing')
st(doc,'2. "Check each one. Honest. 2 minutes." — Timer: 2 min.')
st(doc,'3. "All 4 checked?" — Manos. "3 checked?" — Manos.')
st(doc,'4. "Only 2 or fewer?" — Si hay: "That\'s ok. Next time better. You TRIED."')
doc.add_paragraph('')
h(doc,'Teacher Talk — Closing Script (5 min) — 2:53 a 2:58',3)
ip(doc,'"Excellent work today. Real phone conversations in English. Three rounds. Three partners. Two roles. Real communication.\n\nHomework: Workbook Ch. 6. (1) Reflection paragraph — 5 sentences about what was easy/hard. (2) Review ALL vocab and grammar for the quiz. 45 minutes.\n\nNext session — the quiz. Prove what you know. 50 points."')
icq(doc,'How many sentences in reflection?','5')
icq(doc,'What do you review?','All vocabulary and grammar')
icq(doc,'Why?','For the quiz next session')
doc.add_paragraph('')
h(doc,'Preview — 2:58 a 3:00',3)
ip(doc,'"Next session: Prove What You Know. Cuestionario Técnico. 50 points. 5 sections. Show me what you\'ve learned."')
doc.add_page_break()

h(doc,'EVALUATION CRITERIA',1)
h(doc,'Observation Criteria (para el instructor durante circulación)',2)
tbl(doc,['Criterio','Competente','En desarrollo'],[
    ['Identifica componentes','≥3 términos técnicos correctamente','0-2 términos o confunde'],
    ['Usa is/has','Correcto en ≥70% de oraciones','Confunde frecuentemente'],
    ['Inteligibilidad','Compañero entiende problema + solicitud','Mensaje confuso o incompleto'],
    ['Resolución','Ambos acuerdan NOW vs. ORDER','No llegan a acuerdo'],
],sc='2E75B6')
doc.add_paragraph('')
h(doc,'Correct Phrases del Skeleton Script',2)
tbl(doc,['Function','Correct phrases'],[
    ['Greet','"IT Support, this is [name]." / "This is [name] from DevCore."'],
    ['Describe','"My workstation has a problem." / "The [X] is [adjective]."'],
    ['Ask','"What about the [X]?" / "Is it an SSD or an HDD?"'],
    ['Confirm','"Got it. So... new [X], check." / "I\'ll send the request today."'],
    ['Close','"Thanks, you\'re a lifesaver." / "No problem. Have a good one."'],
],sc='2E75B6')
doc.add_page_break()

h(doc,'DIFFERENTIATION NOTES',1)
h(doc,'Fast Finishers',2)
for t,d in [
    ('Round 1','"Add 1 NEW problem NOT on the card. Invent one."'),
    ('Round 2','"Challenge: no Skeleton Script either. 100% from memory."'),
    ('Round 3','"Be coaches — sit next to struggling pair, whisper suggestions."'),
]:bp(doc,f'{t}: ',d)
doc.add_paragraph('')
h(doc,'More Support Needed',2)
for t,d in [
    ('Ronda 1','Mantener Skeleton Script visible toda la conversación.'),
    ('Ronda 2','Permitir "one look" de 30 seg en tarjeta.'),
    ('Parejas','Emparejar niveles similares.'),
    ('If blocked','Permitir lectura directa de tarjeta. Práctica de lectura oral tiene valor.'),
]:bp(doc,f'{t}: ',d)
doc.add_paragraph('')
h(doc,'Plan B — Fishbowl (>25 aprendices)',2)
st(doc,'1. Grupos de 6-8 en círculo.')
st(doc,'2. 2 voluntarios en centro actúan. 4-6 observan con checklist.')
st(doc,'3. Cada ronda: 5 min. Después: 1 positivo + 1 sugerencia de observadores.')
st(doc,'4. Rotar. En 30 min, 4-6 actúan, resto practica observation + feedback.')
doc.add_page_break()

h(doc,'INSTRUCTOR SELF-CHECK',1)
doc.add_paragraph('')
for i,q in enumerate([
    '¿Todas las parejas completaron al menos 2 rondas? ¿Cuántas completaron las 3?',
    '¿Hubo mejora visible entre Round 1 y Round 2? ¿Y entre Round 2 y Round 3?',
    '¿Los errores más comunes fueron vocabulario, gramática o pronunciación? ¿Cuáles?',
    '¿Los aprendices llegaron a resoluciones? ¿Cuántas parejas NO llegaron a acuerdo?',
    '¿Qué % marcó las 4 casillas del Communication Deliverable?',
],1):
    bp(doc,f'{i}. {q}')
    p=doc.add_paragraph('Respuesta: _______________________________________________')
    for r in p.runs:r.font.name='Calibri';r.font.size=Pt(11)
    doc.add_paragraph('')

doc.add_paragraph('')
for txt in ['SESSION 6: THE HELP DESK — BUILD-OUT','ADSO — GUÍA 1: The Hardware Specialist','Sistema de Prompts Maestros — LG Factory — FPI SENA — Bilingüismo','Instructor Sergio Cortés Perdomo · Marzo 2026']:
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(txt);r.font.size=Pt(9);r.font.name='Calibri';r.font.color.rgb=RGBColor(0x7F,0x8C,0x8D)

doc.save('/Users/Beppo/Projects/fpi-sena-factory/guides/ADSO-G1/ADSO — GUÍA 1 — Session 6 — The Help Desk — Build-Out.docx')
print('Done')
