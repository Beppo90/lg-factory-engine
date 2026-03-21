#!/usr/bin/env python3
"""Generate Session 3 Build-Out as a Word document."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5); s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)
style = doc.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(11)

def shade(cell, color):
    sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), color); sh.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(sh)

def tbl(doc, headers, rows, sc='1F3A5F'):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''
        r = c.paragraphs[0].add_run(str(h)); r.font.size = Pt(10); r.font.name = 'Calibri'; r.bold = True; r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); shade(c, sc)
    for rd in rows:
        row = t.add_row()
        for i, v in enumerate(rd):
            c = row.cells[i]; c.text = ''
            r = c.paragraphs[0].add_run(str(v)); r.font.size = Pt(9); r.font.name = 'Calibri'
    return t

def h(doc, text, level=1):
    hd = doc.add_heading(text, level=level)
    for r in hd.runs: r.font.name = 'Calibri'

def bp(doc, label, value=''):
    p = doc.add_paragraph(); r = p.add_run(label); r.bold = True; r.font.name = 'Calibri'; r.font.size = Pt(11)
    if value:
        r2 = p.add_run(value); r2.font.name = 'Calibri'; r2.font.size = Pt(11)

def tt(doc, text):
    p = doc.add_paragraph(); r = p.add_run('Teacher Talk: '); r.bold = True; r.italic = True; r.font.name = 'Calibri'
    r2 = p.add_run(f'"{text}"'); r2.italic = True; r2.font.name = 'Calibri'

def fn(doc, text):
    p = doc.add_paragraph(); r = p.add_run('💡 '); r.font.name = 'Calibri'
    r2 = p.add_run(text); r2.italic = True; r2.font.name = 'Calibri'; r2.font.color.rgb = RGBColor(0x2E,0x75,0xB6)

def ck(doc, text):
    p = doc.add_paragraph(); r = p.add_run('✓ Checkpoint: '); r.bold = True; r.font.name = 'Calibri'; r.font.color.rgb = RGBColor(0x2E,0x7D,0x32)
    r2 = p.add_run(text); r2.font.name = 'Calibri'

def icq(doc, q, e):
    p = doc.add_paragraph(); r = p.add_run('ICQ: '); r.bold = True; r.font.name = 'Calibri'
    r2 = p.add_run(f'"{q}"'); r2.font.name = 'Calibri'
    r3 = p.add_run(f' → Esperan: "{e}"'); r3.italic = True; r3.font.name = 'Calibri'; r3.font.color.rgb = RGBColor(0x7F,0x8C,0x8D)

def ip(doc, text):
    p = doc.add_paragraph(); r = p.add_run(text); r.italic = True; r.font.name = 'Calibri'; r.font.size = Pt(11)

def st(doc, text):
    p = doc.add_paragraph(text)
    for r in p.runs: r.font.name = 'Calibri'; r.font.size = Pt(11)

def tr(doc, label, text):
    p = doc.add_paragraph(); r = p.add_run(f'{label} '); r.bold = True; r.font.name = 'Calibri'
    r2 = p.add_run(text); r2.italic = True; r2.font.name = 'Calibri'

# ═══ TITLE ═══
for _ in range(4): doc.add_paragraph('')
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('SESSION 3: TUNING IN'); r.bold = True; r.font.size = Pt(20); r.font.name = 'Calibri'; r.font.color.rgb = RGBColor(0x1F,0x3A,0x5F)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('ADSO — GUÍA 1: The Hardware Specialist — Build-Out'); r.font.size = Pt(14); r.font.name = 'Calibri'; r.font.color.rgb = RGBColor(0x44,0x72,0xC4)
doc.add_paragraph('')
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('⚠️ Este documento es SOLO para el instructor. No distribuir a los aprendices.'); r.bold = True; r.font.size = Pt(12); r.font.name = 'Calibri'; r.font.color.rgb = RGBColor(0xC0,0x39,0x2B)
doc.add_paragraph('')
for txt in ['Sistema de Prompts Maestros — LG Factory — FPI SENA — Bilingüismo','Instructor Sergio Cortés Perdomo · Marzo 2026']:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt); r.font.size = Pt(10); r.font.name = 'Calibri'; r.font.color.rgb = RGBColor(0x7F,0x8C,0x8D)
doc.add_page_break()

# ═══ HEADER ═══
h(doc, 'SESSION HEADER', 1)
tbl(doc, ['Campo','Dato'], [
    ['Programa','ADSO — 228118'],['Guía','Guía 1: The Hardware Specialist'],['Session','3: Tuning In'],
    ['Worksheets','PM-2.4 (Listening) + PM-2.5 (Vocabulary)'],['Duración','180 minutos'],
    ['Habilidades foco','L● V●'],['Habilidades soporte','R○'],
    ['Trabajo autónomo','Workbook Ch. 3 — Vocabulary Reinforcement (45 min)'],
    ['Siguiente sesión','Session 4: Say It Right, Build It Right (PM-2.6 + PM-2.7)'],
    ['Plan B','Si audio no disponible: instructor lee script, estudiantes siguen transcripción'],
])
doc.add_paragraph('')

# ═══ MATERIALS ═══
h(doc, 'MATERIALS CHECKLIST', 1)
p = doc.add_paragraph('Marcar ANTES de entrar al aula:'); p.runs[0].italic = True
for item in ['PM-2.4: Listening worksheet (1 por estudiante + 2 extras)','PM-2.5: Vocabulary worksheet (1 por estudiante + 2 extras)',
    'Audio file: Phone Call TTS (0.85x, 2 voices) — REPRODUCIR antes de clase','Speakers o auriculares — VERIFICAR que funcionan',
    'Canva slides 8-14 abiertas y listas','Proyector encendido y funcionando','Tablero preparado según Board Plan',
    'Markers de colores (negro + azul)','Timer visible','Transcripción impresa (2-3 copias — Plan B)']:
    p = doc.add_paragraph(f'☐  {item}')
    for r in p.runs: r.font.name = 'Calibri'; r.font.size = Pt(11)
doc.add_paragraph('')

# ═══ BOARD PLAN ═══
h(doc, 'BOARD PLAN', 1)
bt = ('SESSION 3: TUNING IN\n'
      'Today: "Listen to Carlos on the phone + learn 20 hardware words"\n\n'
      '[AUDIO CUES]                    [FORMULA ZONE]\n'
      '• "gig" = Gigabyte              _____ is a _____\n'
      '• "lemme check" = let me check  _____ has _____\n'
      '• "is" vs "has" sound similar   This is a _____\n\n'
      '[VOCAB WALL — pre-escrito]\n'
      'INPUT: keyboard, mouse, scanner | OUTPUT: monitor, printer\n'
      'STORAGE: SSD, HDD | INTERNAL: CPU, RAM, GPU, motherboard, PSU\n'
      'CONNECTIVITY: USB port, HDMI port, Ethernet cable\n'
      'SPECS: Gigabyte, Terabyte, Gigahertz, Compatible, Portable\n\n'
      '[LIVE ZONE — vacío]')
p = doc.add_paragraph(); r = p.add_run(bt); r.font.name = 'Consolas'; r.font.size = Pt(9)
doc.add_page_break()

# ═══ TIMELINE ═══
h(doc, 'MINUTE-BY-MINUTE TIMELINE', 1)
tbl(doc, ['Tiempo','Dur.','Bloque','Actividad','Agrupación','Notas'], [
    ['0:00-0:05','5','SET-UP','Nombrar 3 componentes','Pairs','Reciclaje PM-2.3'],
    ['0:05-0:12','7','SET-UP','Sentences con is/has','Pairs→Plenary','6 voluntarios'],
    ['0:12-0:17','5','SET-UP','Opening + Objective','Plenary',''],
    ['0:17-0:20','3','SET-UP','ICQ + ✓ Checkpoint','Plenary','Thumbs Carlos'],
    ['0:20-0:30','10','WHILE-A','Listen: Audio Cues','Plenary','3 puntos'],
    ['0:30-0:35','5','WHILE-A','Listen 1: Prediction','Pairs',''],
    ['0:35-0:40','5','WHILE-A','Listen 1: First listen','Individual','Audio 1x'],
    ['0:40-0:45','5','WHILE-A','Listen 1: Gist check','Pairs→Plenary',''],
    ['0:45-0:55','10','WHILE-A','Listen 2: Detail','Individual','Audio 2x'],
    ['0:55-1:10','15','WHILE-A','Activity 1: Triage','Ind→Pairs','Match'],
    ['1:10-1:25','15','WHILE-A','Listen 3 + Activity 2','Ind→Plenary','Immediate Action'],
    ['1:25-1:35','10','WHILE-A','Activity 3: Chunks','Ind→Plenary','Speaking scripts'],
    ['1:35-1:50','15','BREAK','BREAK','—',''],
    ['1:50-1:55','5','WHILE-B','Transición: Vocab','Plenary',''],
    ['1:55-2:10','15','WHILE-B','Toolbelt categorization','Ind→Plenary','8 chunks + 20 terms'],
    ['2:10-2:25','15','WHILE-B','Activity 1: Association','Ind→Pairs→Plenary','Match'],
    ['2:25-2:45','20','WHILE-B','Activity 2: Gap Fill','Ind→Pairs→Plenary','System Log'],
    ['2:45-3:00','15','WHILE-B','Activity 3: Micro-Prod','Ind→Small groups','2 sentences'],
    ['3:00-3:15','15','WHILE-C','Reciclaje: 20 terms','Ind→Pairs','PM-2.3 concordance'],
    ['3:15-3:18','3','WHILE-C','✓ Checkpoint','Plenary','Quick Write'],
    ['3:18-3:26','8','WRAP-UP','Exit Ticket','Individual','3 components'],
    ['3:26-3:31','5','WRAP-UP','Closing + Autónomo','Plenary','Workbook Ch. 3'],
    ['3:31-3:33','2','WRAP-UP','Preview','Plenary',''],
])
p = doc.add_paragraph(); r = p.add_run('Total: 180 minutos ✓'); r.bold = True; r.font.name = 'Calibri'; r.font.color.rgb = RGBColor(0x2E,0x7D,0x32)
doc.add_page_break()

# ═══ SET-UP ═══
h(doc, 'SET-UP DETALLADO (20 min) — 0:00 a 0:20', 2)
h(doc, 'Warm-up: "Quick Recall" (12 min)', 3)
bp(doc, 'Paso 1 — Nombrar 3 componentes (5 min)')
ip(doc, '"Good morning everyone. Turn to the person next to you."')
st(doc, '1. "Last session we read Carlos\'s email. With your partner — name 3 hardware components from the email. Just the names. You have 2 minutes."')
st(doc, '2. Timer: 2 min. Pairs. Instructor circula.')
st(doc, '3. "Quick round — one component from each pair." — 6 parejas dicen 1 componente. Escribir en LIVE ZONE.')
bp(doc, 'Paso 2 — Sentences con is/has (7 min)')
st(doc, '4. "Can you say them in a sentence? Use \'is\' or \'has\'. Example: \'The CPU is old.\' Try with your partner. 2 minutes."')
st(doc, '5. Timer: 2 min. Pairs.')
st(doc, '6. "Who wants to share?" — 6 voluntarios. Instructor escribe 2-3 en LIVE ZONE.')
fn(doc, 'Si confunden "is" y "has": escribir formulas en tablero. "IS = describe. HAS = what it contains."')
doc.add_paragraph('')
h(doc, 'Teacher Talk — Opening Script (5 min)', 3)
ip(doc, '"Good morning everyone. Last session we READ Carlos\'s email. Today we\'re going to LISTEN to Carlos on the phone — he\'s calling IT Support. After the phone call, we learn the 20 key vocabulary words. You\'ll hear real English — fast, informal, with shortcuts. Don\'t worry — we\'ll listen THREE times."')
doc.add_paragraph('')
h(doc, 'Objective (escrito en tablero)', 3)
ip(doc, '"Today you will: (1) listen to a phone call about hardware, and (2) learn 20 key vocabulary words."')
icq(doc, 'Are we reading today or listening?', 'Listening')
icq(doc, 'How many times will we listen?', '3 times')
icq(doc, 'How many vocabulary words will we learn?', '20')
doc.add_paragraph('')
ck(doc, '"Thumbs up if you remember what Carlos\'s problem was in his email."')
fn(doc, 'Si >30% thumbs down: "Carlos\'s CPU is old (i3). His RAM is only 4 GB. His HDD is slow."')
tr(doc, 'Transition → Bloque A:', '"Ok. Before we listen, let me show you 3 things that will help you understand the phone call."')
doc.add_page_break()

# ═══ WHILE — BLOQUE A ═══
h(doc, 'WHILE DETALLADO (145 min) — 0:20 a 3:15', 2)
h(doc, 'BLOQUE A — LISTENING (75 min) — 0:20 a 1:35', 3)
bp(doc, 'Worksheet: ', 'PM-2.4 (The Auditory Anchor)')
bp(doc, 'Agrupación: ', 'Plenary → Individual → Pairs → Plenary')
bp(doc, 'Objetivo: ', 'Los aprendices comprenden audio de llamada telefónica, extraen información y identifican chunks reutilizables.')

doc.add_paragraph('')
h(doc, 'Audio Cues: Introducción (10 min) — 0:20 a 0:30', 3)
bp(doc, 'Instrucciones paso a paso:')
st(doc, '1. Proyectar slide 8 (Audio Cues).')
st(doc, '2. Punto 1 — "gig" = Gigabyte: "In the audio you\'ll hear \'gig.\' Short for Gigabyte. \'4 gigs of RAM\' = 4 Gigabytes."')
st(doc, '3. Punto 2 — "lemme check" = "let me check": "Very fast, very informal. It means \'let me look at the information.\'"')
st(doc, '4. Punto 3 — "is" vs "has": "They sound VERY similar at speed. \'is\' + ADJECTIVE (old, slow). \'has\' + THING (a problem, 4 GB). The grammar helps your ear."')
st(doc, '5. "Write these 3 notes on your PM-2.4 worksheet — Audio Cues section. 2 minutes."')
st(doc, '6. Timer: 2 min. Individual.')
fn(doc, 'Para A1: "If you hear \'is,\' the next word is an ADJECTIVE. If you hear \'has,\' the next word is a THING. The grammar helps your ear."')

doc.add_paragraph('')
h(doc, 'Listen 1: Gist (15 min) — 0:30 a 0:45', 3)
bp(doc, 'Instrucciones paso a paso:')
st(doc, '1. "Before we listen — predict. Who is calling? Who answers? What\'s the problem? Partner. 2 minutes."')
st(doc, '2. Timer: 2 min. Pairs.')
st(doc, '3. "Listen ONCE. Don\'t write — just LISTEN. Focus on WHO and WHAT."')
st(doc, '4. Reproducir audio (≈2 min).')
st(doc, '5. "With your partner — Who is calling? What is the problem? 1 minute."')
st(doc, '6. Timer: 1 min. Pairs.')
st(doc, '7. "Who can answer?" — 2-3 voluntarios. Respuesta: "Carlos is calling IT Support. His workstation is slow."')
st(doc, '8. Instructor escribe: "Caller: Carlos | Problem: Slow workstation"')
fn(doc, 'Si la clase no entiende nada (normal con A1): "Don\'t worry. We listen two more times."')

doc.add_paragraph('')
h(doc, 'Listen 2 + Activity 1 — The Triage (25 min) — 0:45 a 1:10', 3)
bp(doc, 'Instrucciones paso a paso:')
st(doc, '1. "Listen AGAIN. This time: Activity 1 — The Triage. Match each PROBLEM with the COMPONENT. Ready?"')
st(doc, '2. Reproducir audio segunda vez.')
bp(doc, 'Answer Key (in-line):')
tbl(doc, ['Problem','Component'], [
    ['1. "It\'s super slow"','B. CPU (also E)'],
    ['2. "Only has 4 gigs"','D. RAM'],
    ['3. "Only has a VGA port"','A. Monitor'],
    ['4. "It\'s an i3"','B. CPU'],
    ['5. "500 gigs, really slow"','C. Storage (HDD)'],
], sc='2E75B6')
st(doc, '')
st(doc, '3. "3 minutes to finish matching. Alone first."')
st(doc, '4. Timer: 3 min. Individual.')
st(doc, '5. "Compare with partner. If different — discuss. 3 minutes."')
st(doc, '6. Timer: 3 min. Pairs.')
st(doc, '7. "Let\'s check." — ítem por ítem en plenaria.')
fn(doc, 'Ítem 1 y 4 son AMBOS CPU. Algunos aprendices pueden poner E para el ítem 1. Es aceptable pero la respuesta esperada es B para ambos.')
st(doc, '8. "Do you need a third listen? Raise hand if YES."')
st(doc, '   Si >40%: reproducir tercera vez. Si no: avanzar.')

doc.add_page_break()
h(doc, 'Listen 3 + Activity 2 — The Immediate Action (15 min) — 1:10 a 1:25', 3)
bp(doc, 'Instrucciones paso a paso:')
st(doc, '1. "Activity 2. You are Mike (IT Support). What do you do FIRST? Three options. Choose ONE and justify. 3 minutes."')
st(doc, '2. Timer: 3 min. Individual.')
bp(doc, 'Opciones:')
st(doc, 'A. Replace the SSD immediately (it\'s in stock)')
st(doc, 'B. Order a new CPU (takes 5 days)')
st(doc, 'C. Replace the monitor first')
bp(doc, 'Answer Key:')
st(doc, 'Mejor respuesta: A. Replace the SSD — Mike dice: "We have an SSD in stock." Disponible inmediatamente.')
st(doc, 'También aceptable: B. Order CPU — Carlos dice "The CPU is old" como problema principal.')
st(doc, '')
st(doc, '3. "Who chose A? B? C?" — Contar manos.')
st(doc, '4. "2 people explain WHY. Use the sentence from the audio." — 4-5 voluntarios.')
st(doc, '5. Instructor escribe justificación más fuerte en LIVE ZONE.')

doc.add_paragraph('')
h(doc, 'Activity 3 — Chunk Extraction (10 min) — 1:25 a 1:35', 3)
bp(doc, 'Instrucciones paso a paso:')
st(doc, '1. "Activity 3: Noticing & Chunk Extraction. Find the EXACT phrases for 3 functions. 5 minutes."')
st(doc, '2. Timer: 5 min. Individual.')
bp(doc, 'Answer Key (in-line):')
tbl(doc, ['Function','Exact phrase'], [
    ['Greet','"IT Support, this is Mike." / "This is Carlos from DevCore."'],
    ['Describe problem','"My workstation has a problem." / "The computer is super slow."'],
    ['Confirm solution','"Got it. So... new SSD, check the CPU, and HDMI monitor." / "I\'ll send the request today."'],
], sc='2E75B6')
st(doc, '')
st(doc, '3. Revisar en plenaria. Instructor escribe en tablero:')
ip(doc, 'GREET: "IT Support, this is [name]." | DESCRIBE: "My [device] has a problem." | CONFIRM: "Got it. So... new [component]."')
st(doc, '4. "These are YOUR SPEAKING SCRIPTS. Session 6 you use them in a simulation. Save them."')
fn(doc, 'Refuerce: "This is not just listening practice. This is preparation for your phone call in Session 6."')
tr(doc, 'Transition → BREAK:', '"Great work with the listening. Take a break — 15 minutes. Come back ready for vocabulary."')
doc.add_page_break()

# ═══ BREAK ═══
h(doc, 'BREAK — Stretch & Recharge (15 min) — 1:35 a 1:50', 2)
doc.add_paragraph('Instructor toma nota de:')
for item in ['Qué ítems del Triage causaron más confusión','Qué aprendices necesitan reproducción adicional','Verificar Speaking Scripts en tablero']:
    p = doc.add_paragraph(item, style='List Bullet')
    for r in p.runs: r.font.name = 'Calibri'; r.font.size = Pt(11)
doc.add_page_break()

# ═══ BLOQUE B — VOCABULARY ═══
h(doc, 'BLOQUE B — VOCABULARY (60 min) — 1:50 a 3:00', 3)
bp(doc, 'Worksheet: ', 'PM-2.5 (Vocabulary & Language Function)')
bp(doc, 'Agrupación: ', 'Plenary → Individual → Pairs → Small groups → Plenary')
bp(doc, 'Objetivo: ', 'Los aprendices dominan los 20 términos clave a través de categorización, asociación, gap fill y producción.')

doc.add_paragraph('')
h(doc, 'Transición + Toolbelt Categorization (15 min) — 1:50 a 2:05', 3)
tt(doc, 'Welcome back. You heard Carlos and Mike talking about hardware. You understood the problems. Now — let\'s learn the WORDS. The 20 key vocabulary words for hardware. Open your PM-2.5 worksheet.')
bp(doc, 'Instrucciones paso a paso:')
st(doc, '1. Proyectar slide 10 (Toolbelt — 8 chunks).')
st(doc, '2. Instructor presenta cada chunk con uso:')
st(doc, '   CPU → "The CPU is..." (to define)')
st(doc, '   RAM → "It has... GB of RAM" (to describe specs)')
st(doc, '   SSD/HDD → "An SSD is faster than an HDD" (to compare)')
st(doc, '   Monitor → "The monitor has an HDMI port" (to describe features)')
st(doc, '   Keyboard/Mouse → "This is an input device" (to classify)')
st(doc, '   Compatible → "This GPU is compatible with..." (to evaluate)')
st(doc, '   Portable → "A laptop is portable" (to describe)')
st(doc, '   Upgrade → "I need to upgrade the RAM" (to request)')
st(doc, '')
st(doc, '3. "Look at the VOCAB WALL. 20 terms, 6 categories. Create a table in your worksheet. Write each word under its category. 5 minutes."')
st(doc, '4. Timer: 5 min. Individual.')
fn(doc, 'La categorización es clave para la memoria. "INPUT devices SEND information TO the computer. OUTPUT devices SEND information FROM the computer to you."')

doc.add_paragraph('')
h(doc, 'Activity 1 — Logical Association (15 min) — 2:05 a 2:20', 3)
bp(doc, 'Instrucciones paso a paso:')
st(doc, '1. "Match component with FUNCTION. Don\'t translate — think what it DOES. 3 minutes."')
st(doc, '2. Timer: 3 min. Individual.')
bp(doc, 'Answer Key:')
tbl(doc, ['Component','Function'], [
    ['1. CPU','C. Processes all instructions'],
    ['2. RAM','E. Stores data temporarily (fast access)'],
    ['3. GPU','B. Displays images on the screen'],
    ['4. SSD','A. Stores data permanently'],
    ['5. PSU','D. Supplies energy to all parts'],
], sc='2E75B6')
st(doc, '')
st(doc, '3. "Compare with partner. 2 minutes." — Timer: 2 min. Pairs.')
st(doc, '4. Revisar ítem por ítem en plenaria.')
ck(doc, '"Without looking — what does the PSU do?" — 1 voluntario.')

doc.add_paragraph('')
h(doc, 'Activity 2 — Gap Fill: The System Log (20 min) — 2:20 a 2:40', 3)
bp(doc, 'Instrucciones paso a paso:')
st(doc, '1. "Activity 2. System Diagnostic Log — real computer report. Complete with Word Bank. 7 blanks, 7 words. 7 minutes."')
st(doc, '2. Timer: 7 min. Individual.')
bp(doc, 'Answer Key:')
tbl(doc, ['Blank','Answer'], [
    ['1. "The ___ is operating at 95%"','CPU'],
    ['2. "Current ___ : 8 GB"','RAM'],
    ['3. "Storage type: ___ (mechanical)"','HDD'],
    ['4. "Recommendation: ___ the storage"','Upgrade'],
    ['5. "to a 1 TB ___"','SSD'],
    ['6. "The current ___ does not have HDMI"','monitor'],
    ['7. "The port is not ___ with the new GPU"','compatible'],
], sc='2E75B6')
st(doc, '')
st(doc, '3. "Compare with partner. 3 minutes." — Timer: 3 min. Pairs.')
st(doc, '4. Revisar blank por blank en plenaria.')
fn(doc, 'Blanks 4 ("Upgrade") y 7 ("compatible") son los más difíciles para A1. Explique: "Compatible = works together. Upgrade = change for a better one."')
st(doc, '5. "Read the WHOLE log aloud to your partner. Slowly. 2 minutes." — Timer: 2 min. Pairs.')

doc.add_paragraph('')
h(doc, 'Activity 3 — Micro-Production (15 min) — 2:40 a 2:55', 3)
bp(doc, 'Instrucciones paso a paso:')
st(doc, '1. "You are a technician. Write 2 sentences about a computer in YOUR classroom at SENA. Use Toolbelt words."')
st(doc, '2. "Formula: \'The [component] is/has [specification].\' 5 minutes."')
st(doc, '3. Timer: 5 min. Individual.')
st(doc, '4. "Share in groups of 3. Read your sentences. Check: is/has correct? 3 minutes."')
st(doc, '5. Timer: 3 min. Small groups.')
st(doc, '6. "Who wants to share?" — 3-4 voluntarios. Instructor escribe las mejores en LIVE ZONE.')
fn(doc, 'Los aprendices pueden describir computadores reales del salón SENA. "Look at the computer in front of you. What can you see?"')

doc.add_page_break()

# ═══ RECICLAJE ═══
h(doc, 'Reciclaje: Return to PM-2.3 (15 min) — 3:00 a 3:15', 3)
bp(doc, 'Instrucciones paso a paso:')
st(doc, '1. "Open your PM-2.3 — Carlos\'s email. Find the 20 vocabulary words. Underline EVERY word from the VOCAB WALL that appears in the email. 5 minutes."')
st(doc, '2. Timer: 5 min. Individual.')
bp(doc, 'Términos que aparecen en el email (12 de 20):')
st(doc, 'CPU, RAM, SSD, HDD, monitor, keyboard, mouse, compatible, Gigabyte, HDMI port, VGA port, GPU')
st(doc, '')
st(doc, '3. "Compare with partner. How many did you find? 2 minutes." — Timer: 2 min. Pairs.')
st(doc, '4. "How many pairs found more than 10?" — Manos arriba.')
st(doc, '5. Instructor hace mini-concordancia en tablero:')
ip(doc, 'IN EMAIL: ✓ CPU ✓ RAM ✓ SSD ✓ HDD ✓ monitor ✓ keyboard ✓ mouse ✓ compatible ✓ GB ✓ HDMI ✓ VGA ✓ GPU\nNOT IN EMAIL: ✗ Motherboard ✗ PSU ✗ Scanner ✗ Printer ✗ USB ✗ Ethernet ✗ TB ✗ GHz ✗ Portable')
st(doc, '')
st(doc, '6. "12 of 20 in the email. The other 8 you learned TODAY. You now know all 20."')
ck(doc, '"Quick Write — one sentence using \'compatible.\' Show me. 1 minuto."')
tr(doc, 'Transition → WRAP-UP:', '"Good work. Let\'s close."')
doc.add_page_break()

# ═══ WRAP-UP ═══
h(doc, 'WRAP-UP DETALLADO (15 min) — 3:15 a 3:30', 2)
h(doc, 'Exit Ticket (8 min) — 3:15 a 3:23', 3)
bp(doc, 'Instrucciones:')
st(doc, '1. "Write 3 hardware components and 1 sentence about each using \'has\' or \'is.\'"')
st(doc, '2. "Example: \'1. CPU — The CPU is the brain of the computer.\' Now you write 3."')
st(doc, '3. "3 minutes. Go." — Timer: 3 min. Individual.')
st(doc, '4. "Show your partner. 3 sentences? is/has correct? 30 seconds." — Timer: 30 seg. Pairs.')
bp(doc, 'Criterio de éxito: ', '3 componentes con oración que usa "is" o "has" correctamente.')
bp(doc, 'Ejemplos válidos:')
st(doc, '- "1. RAM — The computer has 8 GB of RAM."')
st(doc, '- "2. Monitor — The monitor is big."')
st(doc, '- "3. SSD — An SSD is faster than an HDD."')

doc.add_paragraph('')
h(doc, 'Teacher Talk — Closing Script (5 min) — 3:23 a 3:28', 3)
ip(doc, '"Good work today. You listened to a real phone call — Carlos and Mike. You understood the problems and solutions. And you learned 20 vocabulary words.\n\nYour homework: Workbook Chapter 3 — Vocabulary Reinforcement. Three things: (1) categorize the 20 words into 6 categories. (2) Do the crossword. (3) Write a paragraph — 5 sentences — about \'My Ideal Workstation\' using at least 8 vocabulary words. 45 minutes.\n\nNext session — pronunciation and grammar. How to say the words correctly and how to build sentences. See you next time."')
icq(doc, 'How many categories?', '6')
icq(doc, 'How many sentences in the paragraph?', '5')
icq(doc, 'How many vocabulary words must you use?', '8')

doc.add_paragraph('')
h(doc, 'Preview — 3:28 a 3:30', 3)
ip(doc, '"Next session: Say It Right, Build It Right. Pronunciation — how to say the words. Grammar — how to build sentences with \'is,\' \'has,\' and \'this/that/these/those.\' Come ready to speak."')
doc.add_page_break()

# ═══ ANSWER KEY ═══
h(doc, 'ANSWER KEY CONSOLIDADO', 1)
h(doc, 'PM-2.4 — Listening', 2)
tbl(doc, ['Actividad','Ítem','Respuesta','Alternativas'], [
    ['Activity 1','1. "Super slow"','B. CPU','E (same)'],
    ['Activity 1','2. "4 gigs"','D. RAM','—'],
    ['Activity 1','3. "VGA port"','A. Monitor','—'],
    ['Activity 1','4. "It\'s an i3"','B. CPU','E (same)'],
    ['Activity 1','5. "500 gigs"','C. Storage (HDD)','—'],
    ['Activity 2','Action','A. Replace SSD (in stock)','B. Order CPU (main problem)'],
    ['Activity 3','Greet','"IT Support, this is Mike."','"This is Carlos from DevCore."'],
    ['Activity 3','Describe','"My workstation has a problem."','"The computer is super slow."'],
    ['Activity 3','Confirm','"Got it. So... new SSD..."','"I\'ll send the request today."'],
], sc='2E75B6')
doc.add_paragraph('')
h(doc, 'PM-2.5 — Vocabulary', 2)
tbl(doc, ['Actividad','Ítem','Respuesta','Alternativas'], [
    ['Activity 1','1. CPU','C. Processes instructions','—'],
    ['Activity 1','2. RAM','E. Stores temporarily','—'],
    ['Activity 1','3. GPU','B. Displays images','—'],
    ['Activity 1','4. SSD','A. Stores permanently','—'],
    ['Activity 1','5. PSU','D. Supplies energy','—'],
    ['Activity 2','Blank 1','CPU','—'],
    ['Activity 2','Blank 2','RAM','—'],
    ['Activity 2','Blank 3','HDD','—'],
    ['Activity 2','Blank 4','Upgrade','—'],
    ['Activity 2','Blank 5','SSD','—'],
    ['Activity 2','Blank 6','monitor','—'],
    ['Activity 2','Blank 7','compatible','—'],
], sc='2E75B6')
doc.add_page_break()

# ═══ DIFFERENTIATION ═══
h(doc, 'DIFFERENTIATION NOTES', 1)
h(doc, 'Fast Finishers', 2)
for t,d in [
    ('Activity 1 (Triage)','"Write 1 ADDITIONAL problem from the audio NOT in the 5 items."'),
    ('Activity 2 (Gap Fill)','"Write a SECOND log entry for Workstation #13. Use 5 different words."'),
    ('Activity 3 (Micro-Prod)','"Write 2 MORE sentences about a DIFFERENT device — phone or laptop at home."'),
    ('Reciclaje','"Find the 12 terms in LESS than 5 minutes. Time yourself."'),
    ('Peer tutor','"Help your partner with Gap Fill. Read the log aloud together."'),
]: bp(doc, f'{t}: ', d)
doc.add_paragraph('')
h(doc, 'More Support Needed', 2)
for t,d in [
    ('Listening','Permitir leer transcripción DURANTE segunda escucha.'),
    ('Activity 1','Proporcionar transcripción con problemas SUBRAYADOS.'),
    ('Activity 2','Reducir Word Bank a 5 palabras. Proporcionar primera letra de cada blank.'),
    ('Activity 3','Proporcionar sentence starters: "The _______ is _______."'),
    ('Vocab Wall','Permitir foto del tablero como referencia.'),
]: bp(doc, f'{t}: ', d)
doc.add_paragraph('')
h(doc, 'Plan B — Si audio no disponible', 2)
st(doc, '1. Instructor lee script con dos voces (IT Support + Developer).')
st(doc, '2. Velocidad natural pero clara.')
st(doc, '3. Primera lectura: solo escuchar. Segunda: con transcripción. Tercera: si es necesario.')
doc.add_page_break()

# ═══ SELF-CHECK ═══
h(doc, 'INSTRUCTOR SELF-CHECK', 1)
doc.add_paragraph('')
for i,q in enumerate([
    '¿Todos completaron Activity 1 (Triage)? ¿Cuántos tuvieron 5 ítems correctos?',
    '¿Hubo algún momento donde la mayoría parecía perdida? ¿En qué actividad?',
    '¿Los tiempos se cumplieron o hubo desfases? ¿Dónde?',
    '¿Qué palabras del vocabulario causaron más dificultad? ¿Necesito reforzar?',
    '¿Los aprendices identificaron los chunks del Speaking Script? ¿Están listos para Session 6?',
],1):
    bp(doc, f'{i}. {q}')
    p = doc.add_paragraph('Respuesta: _______________________________________________')
    for r in p.runs: r.font.name = 'Calibri'; r.font.size = Pt(11)
    doc.add_paragraph('')

# ═══ FOOTER ═══
doc.add_paragraph('')
for txt in ['SESSION 3: TUNING IN — BUILD-OUT','ADSO — GUÍA 1: The Hardware Specialist','Sistema de Prompts Maestros — LG Factory — FPI SENA — Bilingüismo','Instructor Sergio Cortés Perdomo · Marzo 2026']:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt); r.font.size = Pt(9); r.font.name = 'Calibri'; r.font.color.rgb = RGBColor(0x7F,0x8C,0x8D)

doc.save('/Users/Beppo/Projects/fpi-sena-factory/guides/ADSO-G1/ADSO — GUÍA 1 — Session 3 — Tuning In — Build-Out.docx')
print('Done')
