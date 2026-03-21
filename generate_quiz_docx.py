#!/usr/bin/env python3
"""Generate Cuestionario Técnico as a Word document."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def sh(cell,color):
    e=OxmlElement('w:shd');e.set(qn('w:fill'),color);e.set(qn('w:val'),'clear')
    cell._tc.get_or_add_tcPr().append(e)

def tbl(doc,headers,rows,sc='1F3A5F'):
    t=doc.add_table(rows=1,cols=len(headers));t.style='Table Grid';t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i];c.text='';r=c.paragraphs[0].add_run(str(h))
        r.font.size=Pt(10);r.font.name='Calibri';r.bold=True;r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF);sh(c,sc)
    for rd in rows:
        row=t.add_row()
        for i,v in enumerate(rd):
            c=row.cells[i];c.text='';r=c.paragraphs[0].add_run(str(v));r.font.size=Pt(9);r.font.name='Calibri'
    return t

def h(d,t,l=1):
    hd=d.add_heading(t,level=l)
    for r in hd.runs:r.font.name='Calibri'

def bp(d,l,v=''):
    p=d.add_paragraph();r=p.add_run(l);r.bold=True;r.font.name='Calibri';r.font.size=Pt(11)
    if v:r2=p.add_run(v);r2.font.name='Calibri';r2.font.size=Pt(11)

def ip(d,t):
    p=d.add_paragraph();r=p.add_run(t);r.italic=True;r.font.name='Calibri';r.font.size=Pt(11)

def st(d,t):
    p=d.add_paragraph(t)
    for r in p.runs:r.font.name='Calibri';r.font.size=Pt(11)

def line(d,n=1):
    for _ in range(n):
        p=d.add_paragraph('____________________________________________________________')
        for r in p.runs:r.font.name='Calibri';r.font.size=Pt(11);r.font.color.rgb=RGBColor(0xBB,0xBB,0xBB)

def cb(d,text):
    p=d.add_paragraph();r=p.add_run(f'☐  {text}');r.font.name='Calibri';r.font.size=Pt(11)

doc = Document()
for s in doc.sections:
    s.top_margin=Cm(2.5);s.bottom_margin=Cm(2.5);s.left_margin=Cm(2.5);s.right_margin=Cm(2.5)
style=doc.styles['Normal'];style.font.name='Calibri';style.font.size=Pt(11)

# TITLE
for _ in range(3):doc.add_paragraph('')
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('ADSO — GUÍA 1: THE HARDWARE SPECIALIST');r.bold=True;r.font.size=Pt(20);r.font.name='Calibri';r.font.color.rgb=RGBColor(0x1F,0x3A,0x5F)
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Cuestionario Técnico (IE-01)');r.font.size=Pt(16);r.font.name='Calibri';r.font.color.rgb=RGBColor(0x44,0x72,0xC4)
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Evidencia de Conocimiento');r.font.size=Pt(14);r.font.name='Calibri';r.italic=True
doc.add_paragraph('')
tbl(doc,['',''],[['Programa','Análisis y Desarrollo de Software (ADSO)'],['Guía','Guía 1: The Hardware Specialist'],['Nivel CEFR','A1.1 — A1.2'],['Total','50 puntos'],['Instructor','Sergio Cortés Perdomo']],sc='4472C4')
doc.add_paragraph('')
ip(doc,'Nombre del Aprendiz: _______________________________________________')
ip(doc,'Ficha: _________________  Fecha: _________________')
doc.add_page_break()

# SECTION 1
h(doc,'SECTION 1: READING COMPREHENSION (10 points)',1)
ip(doc,'Read the following email and answer the questions. / (Lee el siguiente correo y responde las preguntas.)')
doc.add_paragraph('')

# Email box
email_text = (
    "INTERNAL MEMO — DevCore Solutions\n"
    "From: Laura Méndez — QA Tester\n"
    "To: IT Department — Hardware Support\n"
    "Subject: New Workstations for Testing Lab\n\n"
    "Hello,\n\n"
    "My name is Laura Méndez. I am a QA Tester at DevCore Solutions. My department has a big problem. We need new workstations for the testing lab.\n\n"
    "Our current computers are very old. The CPUs are Intel i3 processors. They are slow for running test software. The RAM is only 4 GB on each machine. This is not enough. We need 16 GB of RAM.\n\n"
    "The storage is old HDD drives. They are 500 GB each. We need SSD drives. An SSD is faster than an HDD. This is important for our test results.\n\n"
    "The monitors are small. They only have VGA ports. We need monitors with HDMI ports. Those monitors are compatible with our testing equipment.\n\n"
    "We also need new keyboards and mice. Those keyboards are not compatible with the new software. The mice have broken scroll wheels.\n\n"
    "My request: 5 new workstations. Each with an Intel i5 CPU, 16 GB of RAM, a 1 TB SSD, an HDMI monitor, a keyboard, and a mouse. Total budget: $5,000.\n\n"
    "Thank you,\n"
    "Laura Méndez"
)
p=doc.add_paragraph();r=p.add_run(email_text);r.font.name='Calibri';r.font.size=Pt(10);r.italic=True
doc.add_paragraph('')

h(doc,'TASK A — Main Purpose (2 points)',2)
st(doc,'1. What is the main purpose of this email? (1 point)')
cb(doc,'To ask for a new computer for personal use')
cb(doc,'To request 5 new workstations for the testing lab')
cb(doc,'To complain about the IT department')
cb(doc,'To describe the software testing process')
doc.add_paragraph('')
st(doc,'2. Who is Laura Méndez? (1 point)')
cb(doc,'A Junior Developer');cb(doc,'An IT Support Technician')
cb(doc,'A QA Tester');cb(doc,'A Project Manager')
doc.add_paragraph('')

h(doc,'TASK B — Information Extraction (3 points)',2)
ip(doc,'Complete the table with information from the email.')
tbl(doc,['Component','Current Spec','Requested Spec'],[['CPU','',''],['RAM','',''],['Storage','','']],sc='4472C4')
doc.add_paragraph('')

h(doc,'TASK C — True, False, or Not Given (3 points)',2)
ip(doc,'Write T (True), F (False), or NG (Not Given).')
st(doc,'6. Laura works in the same department as Carlos Ramírez. _____')
st(doc,'7. The current monitors have HDMI ports. _____')
st(doc,'8. Laura needs 5 workstations, not just 1. _____')
doc.add_paragraph('')

h(doc,'TASK D — Detail Analysis (2 points)',2)
st(doc,'9. Why does Laura say SSD is better than HDD? (1 point)')
cb(doc,'Because SSD is cheaper');cb(doc,'Because SSD is faster')
cb(doc,'Because SSD has more storage');cb(doc,'Because SSD is compatible with VGA')
doc.add_paragraph('')
st(doc,'10. What is the TOTAL budget Laura requests? (1 point)')
cb(doc,'$1,200');cb(doc,'$1,000');cb(doc,'$5,000');cb(doc,'$6,000')
doc.add_page_break()

# SECTION 2
h(doc,'SECTION 2: WRITING TASK (10 points)',1)
ip(doc,'You are Miguel Torres, a Junior Developer at DevCore Solutions. Your workstation at SENA has problems. Write a Tech Request email to the IT Department.')
doc.add_paragraph('')
bp(doc,'YOUR SITUATION:')
ip(doc,'You work at workstation #12 in the ADSO classroom. The computer is slow. The CPU is old. The keyboard has a broken Ñ key. The monitor has only VGA (no HDMI). Request upgrades.')
doc.add_paragraph('')
bp(doc,'LANGUAGE BANK: ','has — is — need — compatible — upgrade — old — slow — This is — I need')
doc.add_paragraph('')
ip(doc,'Write your email here:')
doc.add_paragraph('')
bp(doc,'From: _________________________________ — Junior Developer')
bp(doc,'To: IT Department — SENA')
bp(doc,'Subject: _______________________________________________')
ip(doc,'Hello,')
for _ in range(5):line(doc)
ip(doc,'I need ');line(doc)
ip(doc,'Thank you,');doc.add_paragraph('');line(doc)
doc.add_paragraph('')
tbl(doc,['Criterio','Points'],[['Format (From/To/Subject/Greeting/Closing)','/ 2'],['Grammar (is/has/demonstratives)','/ 3'],['Vocabulary (hardware terms)','/ 3'],['Clarity (message understandable)','/ 2'],['Total','/ 10']],sc='4472C4')
doc.add_page_break()

# SECTION 3
h(doc,'SECTION 3: LISTENING COMPREHENSION (10 points)',1)
ip(doc,'You will hear a voicemail from the IT Department. Listen and answer.')
doc.add_paragraph('')
h(doc,'TASK A — Multiple Choice (3 points)',2)
st(doc,'1. What is the main topic of the voicemail?')
cb(doc,'A problem with a workstation');cb(doc,'An office equipment upgrade plan')
cb(doc,'A new employee introduction');cb(doc,'A software update')
doc.add_paragraph('')
st(doc,'2. What CPU replacement does IT offer?')
cb(doc,'i3 to i7');cb(doc,'i5 to i7');cb(doc,'i3 to i5');cb(doc,'i3 to i3')
doc.add_paragraph('')
st(doc,'3. What is NOT compatible with the new GPUs?')
cb(doc,'SSD drives');cb(doc,'VGA monitors');cb(doc,'HDMI monitors');cb(doc,'16 GB RAM')
doc.add_paragraph('')

h(doc,'TASK B — Gap Fill (4 points)',2)
ip(doc,'Complete the summary with words from the Word Bank.')
bp(doc,'Word Bank: ','upgrade — RAM — SSD — compatible — monitor — GPU')
doc.add_paragraph('')
ip(doc,'The IT Department is planning an office _____________ (4) next week. All i3 CPUs will be replaced with i5 processors. If your workstation has less than 8 GB of _____________ (5), send a request. They have 1 TB _____________ (6) drives available. Old VGA monitors are not _____________ (7) with the new GPUs.')
doc.add_paragraph('')
for i in range(4,8):
    bp(doc,f'{i}. ');line(doc)

doc.add_paragraph('')
h(doc,'TASK C — Matching (3 points)',2)
tbl(doc,['IT Action','Detail'],[['8. Replace CPUs','A. One terabyte each'],['9. Provide SSD drives','B. i3 → i5 (free)'],['10. Deliver HDMI monitors','C. Not compatible with VGA']],sc='4472C4')
st(doc,'8 = _____  9 = _____  10 = _____')
doc.add_page_break()

# SECTION 4
h(doc,'SECTION 4: KEY VOCABULARY PRACTICE IN USE — HOTS (10 points)',1)
ip(doc,'This section tests your ABILITY TO USE vocabulary in real situations.')
doc.add_paragraph('')

h(doc,'TASK A — Apply (3 points)',2)
ip(doc,'"My computer takes 5 minutes to open a project file. It is very slow when saving data."')
st(doc,'12. Which component should be upgraded FIRST?')
cb(doc,'Monitor');cb(doc,'Keyboard');cb(doc,'SSD (storage)');cb(doc,'Mouse')
doc.add_paragraph('')
bp(doc,'13. Explain your choice in ONE sentence: (2 points)')
line(doc,2)
doc.add_paragraph('')

h(doc,'TASK B — Analyze (4 points)',2)
ip(doc,'Find the ODD ONE OUT. Write the word AND explain why.')
doc.add_paragraph('')
st(doc,'14. CPU — RAM — Keyboard — GPU — Motherboard')
bp(doc,'Odd one out: ');line(doc)
bp(doc,'Why: ');line(doc)
doc.add_paragraph('')
st(doc,'15. Monitor — Printer — Scanner — SSD — Keyboard')
bp(doc,'Odd one out: ');line(doc)
bp(doc,'Why: ');line(doc)
doc.add_paragraph('')

h(doc,'TASK C — Evaluate (3 points)',2)
ip(doc,'Budget: $500. Buy only TWO items. Which and why?')
tbl(doc,['Item','Price'],[['16 GB RAM stick','$80'],['1 TB SSD','$120'],['HDMI Monitor (24")','$200'],['Intel i5 CPU','$250'],['New Keyboard','$30'],['New Mouse','$20']],sc='4472C4')
doc.add_paragraph('')
bp(doc,'16. I choose: _______________ and _______________')
bp(doc,'17. Justification 1: ');line(doc)
bp(doc,'18. Justification 2: ');line(doc)
doc.add_page_break()

# SECTION 5
h(doc,'SECTION 5: GRAMMAR & STRUCTURE IN REAL LIFE SCENARIO — HOTS (10 points)',1)
ip(doc,'SCENARIO: You are doing an inventory check of the ADSO classroom at SENA.')
doc.add_paragraph('')
tbl(doc,['Computer','CPU','RAM','Storage','Monitor'],[['PC #1','i3','4 GB','HDD 500 GB','VGA'],['PC #2','i5','8 GB','SSD 256 GB','HDMI'],['PC #3','i3','4 GB','HDD 500 GB','VGA']],sc='4472C4')
doc.add_paragraph('')

h(doc,'TASK A — Apply (3 points)',2)
ip(doc,'Write 3 sentences using the data above.')
ip(doc,'Formula: [Component] + IS/HAS + [specification]')
bp(doc,'19. (Use IS): ');line(doc)
bp(doc,'20. (Use HAS): ');line(doc)
bp(doc,'21. (Use a COMPARISON): ');line(doc)
doc.add_paragraph('')

h(doc,'TASK B — Analyze (4 points)',2)
ip(doc,'Find the ERROR. Write the error, correct version, AND the rule.')
doc.add_paragraph('')
st(doc,'22. "Those monitor is not compatible with the GPU."')
bp(doc,'Error: ');line(doc)
bp(doc,'Correct: ');line(doc)
bp(doc,'Rule: ');line(doc)
doc.add_paragraph('')
st(doc,'23. "An HDD is more faster than an SSD."')
bp(doc,'Error: ');line(doc)
bp(doc,'Correct: ');line(doc)
bp(doc,'Rule: ');line(doc)
doc.add_paragraph('')

h(doc,'TASK C — Evaluate (3 points)',2)
ip(doc,'You are the IT Manager. Two departments need upgrades. Approve ONE. Use demonstratives in your justification.')
doc.add_paragraph('')
bp(doc,'Department A — Development: ','10 CPUs are old (i3). Budget: $2,500')
bp(doc,'Department B — Testing Lab: ','5 monitors VGA only. Budget: $1,000')
doc.add_paragraph('')
bp(doc,'24. I approve Department _____ because: (3 points)')
for _ in range(3):line(doc)

doc.add_page_break()

# END
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('— END OF QUIZ —');r.bold=True;r.font.size=Pt(14);r.font.name='Calibri';r.font.color.rgb=RGBColor(0x1F,0x3A,0x5F)
ip(doc,'Review your answers before submitting.')
doc.add_page_break()

# ═══ ANSWER KEY ═══
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('⚠️ ANSWER KEY — SOLO PARA EL INSTRUCTOR');r.bold=True;r.font.size=Pt(16);r.font.name='Calibri';r.font.color.rgb=RGBColor(0xC0,0x39,0x2B)
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('NO distribuir a los aprendices.');r.bold=True;r.font.size=Pt(12);r.font.name='Calibri';r.font.color.rgb=RGBColor(0xC0,0x39,0x2B)
doc.add_page_break()

h(doc,'SECTION 1: READING (10 pts)',1)
tbl(doc,['#','Answer','Pts','Notes'],[
    ['1','B','1','Request 5 workstations'],['2','C','1','QA Tester'],
    ['3','i3 → i5','1','Both columns'],['4','4 GB → 16 GB','1','Both columns'],
    ['5','HDD 500 → SSD 1TB','1','Both columns'],
    ['6','NG','1','No info about Carlos'],['7','F','1','VGA not HDMI'],
    ['8','T','1','"5 new workstations"'],['9','B','1','"SSD is faster"'],
    ['10','C','1','"$5,000"'],
],sc='2E75B6')
doc.add_paragraph('')

h(doc,'SECTION 2: WRITING (10 pts)',1)
tbl(doc,['Criterio','0 pts','1 pt','2 pts','3 pts'],[
    ['Format','Sin formato','Parcial','Completo','—'],
    ['Grammar','0-1','2','3-4','5+'],
    ['Vocabulary','0-1','2-3','4-5','6+'],
    ['Clarity','Incomprensible','Parcial','Claro','—'],
],sc='2E75B6')
doc.add_paragraph('')

h(doc,'SECTION 3: LISTENING (10 pts)',1)
tbl(doc,['#','Answer','Pts'],[
    ['1','B','1'],['2','C','1'],['3','B','1'],
    ['4','upgrade','1'],['5','RAM','1'],['6','SSD','1'],
    ['7','HDMI','1'],['8','compatible','1'],
    ['9','B','1'],['10','A','1'],['11','C','1'],
],sc='2E75B6')
doc.add_page_break()

h(doc,'SECTION 4: VOCABULARY HOTS (10 pts)',1)
tbl(doc,['#','Answer','Pts','Notes'],[
    ['12','C (SSD)','1','Fixes slow data access'],
    ['13','Open','2','1 vocab + 1 logic'],
    ['14','Keyboard','1+1','Input device / external'],
    ['15','SSD','1+1','Storage / not I-O'],
    ['16','Valid ≤$500','1','Any 2 items'],
    ['17','Justification 1','1','Logical reason'],
    ['18','Justification 2','1','Logical reason'],
],sc='2E75B6')
doc.add_paragraph('')

h(doc,'SECTION 5: GRAMMAR HOTS (10 pts)',1)
tbl(doc,['#','Answer','Pts','Notes'],[
    ['19','IS sentence','1','Correct + from table'],
    ['20','HAS sentence','1','Correct + from table'],
    ['21','Comparison','1','Correct structure'],
    ['22','Those monitor is','1+1','Correct: Those monitors ARE / That monitor IS. Rule: plural/singular.'],
    ['23','more faster','1+1','Correct: faster/slower. Rule: no double comparative.'],
    ['24','Justification','1+1+1','1 choice + 1 logic + 1 demonstrative'],
],sc='2E75B6')
doc.add_page_break()

h(doc,'ASSESSMENT SUMMARY',1)
tbl(doc,['Section','Topic','Pts','Type'],[
    ['1. Reading','Laura Méndez — Workstations','10','Receptiva'],
    ['2. Writing','Miguel Torres — Tech Request','10','Productiva'],
    ['3. Listening','IT Voicemail — Upgrade Plan','10','Receptiva'],
    ['4. Vocabulary HOTS','Diagnosis + Odd One Out + Budget','10','Productiva'],
    ['5. Grammar HOTS','Inventory + Error Log + Recommendation','10','Productiva'],
    ['Total','','50',''],
],sc='2E75B6')
doc.add_paragraph('')

h(doc,'BLOOM DISTRIBUTION',1)
tbl(doc,['Section','Remember','Understand','Apply','Analyze','Evaluate'],[
    ['1. Reading','1,2','3-5,9,10','—','6-8','—'],
    ['2. Writing','—','—','All','—','—'],
    ['3. Listening','1-3','9-11','4-8','—','—'],
    ['4. Vocabulary','—','—','12-13','14-15','16-18'],
    ['5. Grammar','—','—','19-21','22-23','24'],
],sc='2E75B6')

doc.add_page_break()
for txt in ['ADSO — GUÍA 1: The Hardware Specialist — Cuestionario Técnico (IE-01)','Sistema de Prompts Maestros — LG Factory — FPI SENA — Bilingüismo','Instructor Sergio Cortés Perdomo · Marzo 2026']:
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(txt);r.font.size=Pt(9);r.font.name='Calibri';r.font.color.rgb=RGBColor(0x7F,0x8C,0x8D)

doc.save('/Users/Beppo/Projects/fpi-sena-factory/guides/ADSO-G1/ADSO — GUÍA 1 — The Hardware Specialist — Cuestionario Técnico (IE-01).docx')
print('Done')
