#!/usr/bin/env python3
"""Generate Session 2 Build-Out as a Word document."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)

style = doc.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(11)

def set_cell_shading(cell, color):
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), color); s.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(s)

def make_table(doc, headers, rows, shade='1F3A5F'):
    table = doc.add_table(rows=1, cols=len(headers)); table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, t in enumerate(headers):
        c = table.rows[0].cells[i]; c.text = ''
        r = c.paragraphs[0].add_run(str(t)); r.font.size = Pt(10); r.font.name = 'Calibri'
        r.bold = True; r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); set_cell_shading(c, shade)
    for rd in rows:
        row = table.add_row()
        for i, t in enumerate(rd):
            c = row.cells[i]; c.text = ''
            r = c.paragraphs[0].add_run(str(t)); r.font.size = Pt(9); r.font.name = 'Calibri'
    return table

def h(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs: r.font.name = 'Calibri'

def bp(doc, label, value=''):
    p = doc.add_paragraph(); r = p.add_run(label); r.bold = True; r.font.name = 'Calibri'; r.font.size = Pt(11)
    if value:
        r2 = p.add_run(value); r2.font.name = 'Calibri'; r2.font.size = Pt(11)

def tt(doc, text):
    p = doc.add_paragraph(); r = p.add_run('Teacher Talk: '); r.bold = True; r.italic = True; r.font.name = 'Calibri'
    r2 = p.add_run(f'"{text}"'); r2.italic = True; r2.font.name = 'Calibri'

def fn(doc, text):
    p = doc.add_paragraph(); r = p.add_run('💡 '); r.font.name = 'Calibri'
    r2 = p.add_run(text); r2.italic = True; r2.font.name = 'Calibri'; r2.font.color.rgb = RGBColor(0x2E,0x75,0xB6)

def ck(doc, text):
    p = doc.add_paragraph(); r = p.add_run('✓ Checkpoint: '); r.bold = True; r.font.name = 'Calibri'; r.font.color.rgb = RGBColor(0x2E,0x7D,0x32)
    r2 = p.add_run(text); r2.font.name = 'Calibri'

def icq(doc, q, e):
    p = doc.add_paragraph(); r = p.add_run('ICQ: '); r.bold = True; r.font.name = 'Calibri'
    r2 = p.add_run(f'"{q}"'); r2.font.name = 'Calibri'
    r3 = p.add_run(f' → Esperan: "{e}"'); r3.italic = True; r3.font.name = 'Calibri'; r3.font.color.rgb = RGBColor(0x7F,0x8C,0x8D)

def italic_para(doc, text):
    p = doc.add_paragraph(); r = p.add_run(text); r.italic = True; r.font.name = 'Calibri'; r.font.size = Pt(11)

def step(doc, text):
    p = doc.add_paragraph(text)
    for r in p.runs: r.font.name = 'Calibri'; r.font.size = Pt(11)

# ═══ TITLE ═══
for _ in range(4): doc.add_paragraph('')
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('SESSION 2: READ THE REQUEST'); r.bold = True; r.font.size = Pt(20); r.font.name = 'Calibri'; r.font.color.rgb = RGBColor(0x1F,0x3A,0x5F)
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run('ADSO — GUÍA 1: The Hardware Specialist — Build-Out'); r.font.size = Pt(14); r.font.name = 'Calibri'; r.font.color.rgb = RGBColor(0x44,0x72,0xC4)
doc.add_paragraph('')
w = doc.add_paragraph(); w.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = w.add_run('⚠️ Este documento es SOLO para el instructor. No distribuir a los aprendices.'); r.bold = True; r.font.size = Pt(12); r.font.name = 'Calibri'; r.font.color.rgb = RGBColor(0xC0,0x39,0x2B)
doc.add_paragraph('')
for txt in ['Sistema de Prompts Maestros — LG Factory — FPI SENA — Bilingüismo','Instructor Sergio Cortés Perdomo · Marzo 2026']:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt); r.font.size = Pt(10); r.font.name = 'Calibri'; r.font.color.rgb = RGBColor(0x7F,0x8C,0x8D)
doc.add_page_break()

# ═══ SESSION HEADER ═══
h(doc, 'SESSION HEADER', 1)
make_table(doc, ['Campo','Dato'], [
    ['Programa','Análisis y Desarrollo de Software (ADSO) — 228118'],
    ['Guía','Guía 1: The Hardware Specialist'],
    ['Session','2: Read the Request'],
    ['Worksheets','PM-2.3 (The Master Anchor — Reading)'],
    ['Duración','180 minutos'],
    ['Habilidades foco','R●'],
    ['Habilidades soporte','V○'],
    ['Trabajo autónomo','Workbook Ch. 2 — Reading Extension (45 min)'],
    ['Siguiente sesión','Session 3: Tuning In (PM-2.4 + PM-2.5)'],
])
doc.add_paragraph('')

# ═══ MATERIALS CHECKLIST ═══
h(doc, 'MATERIALS CHECKLIST', 1)
p = doc.add_paragraph('Marcar ANTES de entrar al aula:'); p.runs[0].italic = True
for item in ['PM-2.3: The Master Anchor — impresos (1 por estudiante + 2 extras)','Canva slides 5-7 abiertas y listas','Proyector encendido y funcionando','Tablero preparado según Board Plan','Markers de colores (negro + 1 de color para resaltar)','Timer visible en pantalla o celular']:
    p = doc.add_paragraph(f'☐  {item}')
    for r in p.runs: r.font.name = 'Calibri'; r.font.size = Pt(11)
doc.add_paragraph('')

# ═══ BOARD PLAN ═══
h(doc, 'BOARD PLAN', 1)
p = doc.add_paragraph('Preparar ANTES de que lleguen los aprendices:'); p.runs[0].italic = True
bt = ('SESSION 2: READ THE REQUEST\n\n'
      'Today: "Read a real tech request email from Carlos"\n\n'
      '[TOOLBELT — pre-escrito]      [LIVE ZONE — vacío]\n'
      '1. Spec Sheet                 (Se llena con respuestas\n'
      '2. Upgrade                     del Information Transfer\n'
      '3. Compatible                  y decisiones de grupos)\n'
      '4. Performance\n'
      '5. Budget\n\n'
      '[FORMULA ZONE — permanente]\n'
      '"The [component] is [problem]."\n'
      '"I need [a new component]."\n'
      '"An SSD is faster than an HDD."')
p = doc.add_paragraph(); r = p.add_run(bt); r.font.name = 'Consolas'; r.font.size = Pt(9)
doc.add_page_break()

# ═══ TIMELINE ═══
h(doc, 'MINUTE-BY-MINUTE TIMELINE', 1)
make_table(doc, ['Tiempo','Dur.','Bloque','Actividad','Agrupación','Notas'], [
    ['0:00-0:05','5 min','SET-UP','Warm-up: Quick Recall','Pairs','Learning Contract'],
    ['0:05-0:08','3 min','SET-UP','Compartir en plenaria','Plenary','3 voluntarios'],
    ['0:08-0:13','5 min','SET-UP','Opening Script + Objective','Plenary',''],
    ['0:13-0:15','2 min','SET-UP','✓ Checkpoint','Plenary','Thumbs up/down'],
    ['0:15-0:25','10 min','WHILE-A','Pre-Read: Toolbelt intro','Plenary','5 palabras'],
    ['0:25-0:35','10 min','WHILE-A','Pre-Read: Toolbelt practice','Individual','Ejemplos'],
    ['0:35-0:40','5 min','WHILE-A','Pre-Read: Quick check','Plenary','Choral response'],
    ['0:40-0:45','5 min','WHILE-B','Read 1: lectura silenciosa','Individual','Gist question'],
    ['0:45-0:50','5 min','WHILE-B','Read 1: gist check','Pairs','What does Carlos want?'],
    ['0:50-0:52','2 min','WHILE-B','Read 1: plenaria','Plenary',''],
    ['0:52-1:10','18 min','WHILE-B','Read 2: Info Transfer table','Individual','Current vs Requested'],
    ['1:10-1:25','15 min','WHILE-B','Read 2: pair check + plenary','Pairs→Plenary','Revisar tabla'],
    ['1:25-1:30','5 min','WHILE-B','✓ Checkpoint','Pairs','5 filas completas?'],
    ['1:30-1:45','15 min','BREAK','BREAK — Stretch & Recharge','—',''],
    ['1:45-1:50','5 min','WHILE-C','Transición: Post-Read','Plenary','$800 budget'],
    ['1:50-2:05','15 min','WHILE-C','Post-Read: Tech Decision','Individual','Elegir 2 + justificar'],
    ['2:05-2:25','20 min','WHILE-C','Post-Read: debate','Groups of 4','Comparar decisiones'],
    ['2:25-2:35','10 min','WHILE-C','Post-Read: presentaciones','Plenary','2-3 grupos'],
    ['2:35-2:45','10 min','WHILE-D','Deconstruct: find formulas','Individual','3 funciones'],
    ['2:45-2:50','5 min','WHILE-D','Write formulas on board','Plenary','Formula Zone'],
    ['2:55-3:03','8 min','WRAP-UP','Exit Ticket','Individual','2 sentences'],
    ['3:03-3:08','5 min','WRAP-UP','Closing + Autónomo','Plenary','Workbook Ch. 2'],
    ['3:08-3:10','2 min','WRAP-UP','Preview','Plenary',''],
])
p = doc.add_paragraph(); r = p.add_run('Total: 180 minutos ✓'); r.bold = True; r.font.name = 'Calibri'; r.font.color.rgb = RGBColor(0x2E,0x7D,0x32)
doc.add_page_break()

# ═══ SET-UP ═══
h(doc, 'SET-UP DETALLADO (15 min) — 0:00 a 0:15', 2)
h(doc, 'Warm-up: "Quick Recall" (8 min)', 3)
bp(doc, 'Paso 1 — Compartir con partner (5 min)')
italic_para(doc, '"Good morning everyone. Turn to the person next to you."')
step(doc, '1. "Last session, you wrote a Learning Contract. Open your PM-2.2 worksheet — look at your contract. With your partner, share: What are the 2 most important things you need to learn? You have 3 minutes."')
step(doc, '2. Timer: 3 minutos. Pairs. Instructor circula.')
step(doc, '3. "Quick — also tell your partner: how many of the 5 Blind Spots from the board did YOU have? You have 1 more minute."')
step(doc, '4. Timer: 1 minuto.')
bp(doc, 'Paso 2 — Compartir en plenaria (3 min)')
step(doc, '5. "Who wants to share one thing from their Learning Contract? Just one thing." — 3 voluntarios.')
step(doc, '6. Instructor confirma: "Good. These are exactly the things we start working on today."')
fn(doc, 'Si nadie quiere compartir: "Ok, I heard one pair talking about \'CPU\' — is that something you need to learn? Raise your hand if CPU is in your contract."')

doc.add_paragraph('')
h(doc, 'Teacher Talk — Opening Script (3 min)', 3)
italic_para(doc, '"Good morning everyone. Last session we discovered our blind spots — what we don\'t know yet. Today we start FILLING those gaps. We\'re going to read a REAL email from a developer named Carlos. He works at DevCore Solutions. He has a problem with his workstation — and he needs to write a request to IT. Your job: understand what Carlos needs, decide if you would approve his request, and learn the structures he uses so YOU can write your own email later."')

doc.add_paragraph('')
h(doc, 'Objective (escrito en tablero)', 3)
italic_para(doc, '"Today you will: read a real tech request email, extract information, make a decision, and learn writing structures."')
icq(doc, 'Are we writing an email today or reading one?', 'Reading')
icq(doc, 'Who is the email from?', 'Carlos')

doc.add_paragraph('')
ck(doc, '"Quick check — thumbs up if you remember who Carlos is from last session."')
fn(doc, 'Si >30% thumbs down: "Carlos is a Junior Developer at DevCore Solutions. Last session we talked about a scenario where a computer crashes — Carlos is the developer."')
doc.add_page_break()

# ═══ WHILE DETALLADO ═══
h(doc, 'WHILE DETALLADO (150 min) — 0:15 a 2:55', 2)

# BLOQUE A
h(doc, 'BLOQUE A — PRE-READING: TOOLBELT (20 min) — 0:15 a 0:35', 3)
bp(doc, 'Worksheet: ', 'PM-2.3, The Toolbelt section')
bp(doc, 'Agrupación: ', 'Plenary → Individual → Plenary')
bp(doc, 'Objetivo: ', 'Los aprendices conocen las 5 palabras clave antes de leer.')

doc.add_paragraph('')
h(doc, 'Toolbelt: Introducción (10 min) — 0:15 a 0:25', 3)
bp(doc, 'Instrucciones paso a paso:')
step(doc, '1. Proyectar slide 5 (Toolbelt words con definiciones).')
step(doc, '2. Instructor presenta cada palabra, una por una:')
step(doc, '   - Dice la palabra: "The first word is SPEC SHEET."')
step(doc, '   - Da la definición: "A Spec Sheet is a document with technical information about a product."')
step(doc, '   - Da un ejemplo: "Example: \'I need the spec sheet for this computer.\'"')
step(doc, '   - Los aprendices repiten 2 veces en coro.')
step(doc, '3. Las 5 palabras en orden: Spec Sheet, Upgrade, Compatible, Performance, Budget.')
step(doc, '4. "Open your PM-2.3 worksheet. Look at The Toolbelt. Write a simple example sentence for EACH word. You have 5 minutes."')
step(doc, '5. Timer: 5 minutos. Individual.')
fn(doc, 'Los aprendices A1 pueden tener dificultad con "compatible" y "performance". Para "compatible", use gesto de encajar piezas. Para "performance": "Fast = good. Slow = bad."')

doc.add_paragraph('')
h(doc, 'Toolbelt: Quick Check (10 min) — 0:25 a 0:35', 3)
step(doc, '1. "I say the DEFINITION — you say the WORD." — Choral response para las 5 palabras.')
step(doc, '2. "Now I say the WORD — you give me a quick example." — 2-3 voluntarios por palabra.')
ck(doc, '"Write ONE sentence using the word \'upgrade\'. Show me. 1 minuto."')
step(doc, '3. Timer: 1 minuto. Instructor circula rápido y verifica >70%.')
p = doc.add_paragraph(); r = p.add_run('Transition → Bloque B: '); r.bold = True; r.font.name = 'Calibri'
r2 = p.add_run('"Good. You know the 5 key words. Now let\'s read Carlos\'s email. These words are ALL in the email."'); r2.italic = True; r2.font.name = 'Calibri'

doc.add_page_break()

# BLOQUE B
h(doc, 'BLOQUE B — WHILE-READING (50 min) — 0:35 a 1:25', 3)
bp(doc, 'Worksheet: ', 'PM-2.3, Activities 1 + 2')
bp(doc, 'Agrupación: ', 'Individual → Pairs → Plenary')
bp(doc, 'Objetivo: ', 'Los aprendices comprenden el email, extraen datos en tabla, verifican comprensión.')

doc.add_paragraph('')
h(doc, 'First Read: Gist (15 min) — 0:35 a 0:50', 3)
bp(doc, 'Instrucciones paso a paso:')
step(doc, '1. Proyectar slide 6 (email de Carlos).')
step(doc, '2. "Read this email silently. Don\'t answer any questions yet — just READ. You have 5 minutes."')
step(doc, '3. Timer: 5 minutos. Individual. Silencio.')
step(doc, '4. "Done? With your partner — answer: What does Carlos want? One sentence. 2 minutes."')
step(doc, '5. Timer: 2 minutos. Pairs.')
step(doc, '6. "Who can answer?" — 2-3 voluntarios. Respuesta: "He wants to upgrade his workstation."')
fn(doc, 'Es normal que A1 no comprendan todo en la primera lectura. La segunda lectura con la tabla construye la comprensión detallada.')

doc.add_paragraph('')
h(doc, 'Second Read: Information Transfer (35 min) — 0:50 a 1:25', 3)
bp(doc, 'Instrucciones paso a paso:')
step(doc, '1. "Read the email AGAIN. This time you have a TASK. Activity 1: fill in Current Spec and Requested Spec for each component. Use the EXACT words from the email. You have 10 minutes."')
step(doc, '2. Timer: 10 minutos. Individual. Silencio.')
doc.add_paragraph('')
bp(doc, 'Answer Key (in-line):')
make_table(doc, ['Component','Current Spec','Requested Spec'], [
    ['CPU','Intel i3 (old, slow for compiling)','Intel i5 or higher'],
    ['RAM','4 GB','16 GB'],
    ['Storage','500 GB HDD (mechanical, slow)','1 TB SSD'],
    ['Monitor','Small, only VGA port','HDMI monitor'],
    ['Mouse','Broken scroll wheel','New mouse'],
], shade='2E75B6')
step(doc, '')
step(doc, '3. "Turn to your partner. Compare. If different — check the email together. 5 minutes."')
step(doc, '4. Timer: 5 minutos. Pairs.')
step(doc, '5. "Let\'s check. Eyes on me." — Revisar componente por componente. Escribir en LIVE ZONE.')
step(doc, '6. "Total budget?" → "$1,200"')
fn(doc, 'Error común: escribir "4" en lugar de "4 GB". Insista en las unidades — es ESP.')

doc.add_paragraph('')
ck(doc, '"Do you BOTH have 5 components filled in? Raise your hand if yes."')
step(doc, 'Si <80%: 2 más minutos para completar. Si >80%: continuar.')
p = doc.add_paragraph(); r = p.add_run('Transition → BREAK: '); r.bold = True; r.font.name = 'Calibri'
r2 = p.add_run('"Take a break. 15 minutes. Come back ready to make a tough decision."'); r2.italic = True; r2.font.name = 'Calibri'

doc.add_page_break()

# BREAK
h(doc, 'BREAK — Stretch & Recharge (15 min) — 1:30 a 1:45', 2)
p = doc.add_paragraph('Instructor toma nota de:')
for item in ['Qué componentes causaron más confusión (para aclarar después)','Qué aprendices necesitan más apoyo con la lectura']:
    p = doc.add_paragraph(item, style='List Bullet')
    for r in p.runs: r.font.name = 'Calibri'; r.font.size = Pt(11)
doc.add_page_break()

# BLOQUE C
h(doc, 'BLOQUE C — POST-READING (40 min) — 1:45 a 2:25', 3)
bp(doc, 'Worksheet: ', 'PM-2.3, Activity 2')
bp(doc, 'Agrupación: ', 'Plenary → Individual → Groups of 4 → Plenary')
bp(doc, 'Objetivo: ', 'Los aprendices priorizan upgrades con presupuesto limitado, justificando con evidencia del texto.')

doc.add_paragraph('')
h(doc, 'Transición (5 min) — 1:45 a 1:50', 3)
tt(doc, 'Welcome back. You know what Carlos needs. Five upgrades, $1,200 total. But here\'s the problem: you are the IT Manager. Your budget is $800. Not $1,200. You can\'t approve everything. You have to choose.')
step(doc, '1. Proyectar slide 7 (escenario del presupuesto).')
step(doc, '2. Escribir en tablero: Budget: $800 (not $1,200)')

doc.add_paragraph('')
h(doc, 'Technical Decision (15 min) — 1:50 a 2:05', 3)
bp(doc, 'Instrucciones paso a paso:')
step(doc, '1. "Look at Activity 2. 5 options. Budget is $800. Which TWO are MOST important? Approve ✓ or Reject ✗ for EACH. You have 5 minutes. ALONE."')
step(doc, '2. Timer: 5 minutos. Individual.')
step(doc, '3. "Now write your justification: \'I approve _______ because the text says: ________________.\' You must cite evidence. 5 minutes."')
step(doc, '4. Timer: 5 minutos. Individual.')
fn(doc, 'No hay respuesta "correcta" única. Las más defendibles: CPU ("very slow for compiling"), RAM ("only 4 GB, not enough"), SSD ("HDD is slow, SSD is faster"). El mouse y monitor son menos urgentes.')
fn(doc, 'NO diga cuál es la "correcta". Cualquier elección es válida SI se justifica con texto.')

doc.add_paragraph('')
h(doc, 'Debate en grupos de 4 (20 min) — 2:05 a 2:25', 3)
bp(doc, 'Instrucciones paso a paso:')
step(doc, '1. "Join the pair behind you. Group of 4." — 1 min transición.')
step(doc, '2. "Share your TWO choices. Do you all agree? Discuss. Each person explains WHY. Use the sentence from the email. 8 minutes."')
step(doc, '3. Timer: 8 minutos. Groups of 4. Instructor circula.')
step(doc, '4. "Your group must agree on TWO upgrades. Write the group decision on ONE paper. 2 minutes."')
step(doc, '5. Timer: 2 minutos.')
step(doc, '6. "Group 2 — what did you choose and why?" — 2-3 grupos comparten (2 min cada uno).')
step(doc, '7. Después de cada presentación: "Do you agree? Why or why not?" — 1 min intercambio.')
fn(doc, 'Los debates más productivos ocurren cuando hay opiniones diferentes. Si todos eligieron lo mismo: "Did anyone want to approve the monitor? Why?"')
p = doc.add_paragraph(); r = p.add_run('Transition → Bloque D: '); r.bold = True; r.font.name = 'Calibri'
r2 = p.add_run('"Good decisions. Now — Carlos used some very useful sentences in his email. Let\'s find them."'); r2.italic = True; r2.font.name = 'Calibri'

doc.add_page_break()

# BLOQUE D
h(doc, 'BLOQUE D — DESTRUCTURE FOR FUTURE OUTPUT (15 min) — 2:35 a 2:50', 3)
bp(doc, 'Worksheet: ', 'PM-2.3, Activity 3')
bp(doc, 'Agrupación: ', 'Individual → Plenary')
bp(doc, 'Objetivo: ', 'Los aprendices identifican 3 estructuras funcionales del email como "formulas" reutilizables.')

doc.add_paragraph('')
h(doc, 'Encontrar las funciones (10 min) — 2:35 a 2:45', 3)
bp(doc, 'Instrucciones paso a paso:')
step(doc, '1. "Activity 3. 3 functions: describe a problem, request something, compare two things. Find the EXACT sentence from the email for EACH. Copy word for word. 5 minutes."')
step(doc, '2. Timer: 5 minutos. Individual.')
doc.add_paragraph('')
bp(doc, 'Answer Key (in-line):')
make_table(doc, ['Function','Exact sentence from the text'], [
    ['Describe a problem','"The CPU is old." / "It is very slow for compiling code." / "The RAM is only 4 GB."'],
    ['Request something','"I need a new CPU (Intel i5 or higher), 16 GB of RAM, a 1 TB SSD, an HDMI monitor, and a new mouse."'],
    ['Compare two things','"An SSD is faster than an HDD."'],
], shade='2E75B6')
step(doc, '')
step(doc, '3. "What sentence for \'describe a problem\'?" — 2-3 voluntarios. Instructor confirma.')

doc.add_paragraph('')
h(doc, 'Escribir formulas en tablero (5 min) — 2:45 a 2:50', 3)
step(doc, '1. Instructor escribe en Formula Zone:')
p = doc.add_paragraph(); r = p.add_run('DESCRIBE: "The [component] is [problem]."\nREQUEST: "I need [a new component]."\nCOMPARE: "An [X] is [adjective-er] than a [Y]."')
r.font.name = 'Consolas'; r.font.size = Pt(10)
step(doc, '')
step(doc, '2. "These are your WRITING FORMULAS. You use them in Session 5. Memorize them. Screenshot if you want."')
fn(doc, 'Si los aprendices tienen celular, déjelos tomar foto. Estas formulas son la base de la escritura en Session 5.')
p = doc.add_paragraph(); r = p.add_run('Transition → WRAP-UP: '); r.bold = True; r.font.name = 'Calibri'
r2 = p.add_run('"Good work. Let\'s close."'); r2.italic = True; r2.font.name = 'Calibri'

doc.add_page_break()

# ═══ WRAP-UP ═══
h(doc, 'WRAP-UP DETALLADO (15 min) — 2:55 a 3:10', 2)

h(doc, 'Exit Ticket (8 min) — 2:55 a 3:03', 3)
bp(doc, 'Instrucciones:')
step(doc, '1. "Write Carlos\'s problem in YOUR words. 2 sentences. Use \'is\' and \'has\'."')
step(doc, '2. "Example: \'The CPU is old. The computer has only 4 GB of RAM.\' — But YOUR words."')
step(doc, '3. "3 minutes. Go."')
step(doc, '4. Timer: 3 minutos. Individual.')
step(doc, '5. "Show your partner. Does your partner have 2 sentences? 30 seconds."')
step(doc, '6. Timer: 30 segundos. Pairs.')
bp(doc, 'Criterio de éxito: ', 'El aprendiz escribe 2 oraciones con "is" (problema) y "has" (característica) sobre el email de Carlos.')
bp(doc, 'Ejemplos válidos:')
step(doc, '- "The CPU is old. It has 4 GB of RAM."')
step(doc, '- "The workstation is slow. It has an HDD."')
step(doc, '- "The monitor is small. It has a VGA port."')

doc.add_paragraph('')
h(doc, 'Teacher Talk — Closing Script (5 min) — 3:03 a 3:08', 3)
italic_para(doc, '"Good work today. You read a real email, extracted the information, made a tough decision with a limited budget, and found the writing formulas Carlos used. These formulas — \'The CPU is old,\' \'I need a new CPU,\' \'An SSD is faster than an HDD\' — these are YOUR tools now.\n\nYour homework: Workbook Chapter 2 — Reading Extension. Go back to Carlos\'s email. Underline ALL examples of \'is\' and \'has\' — count them. Then write 3 new sentences about YOUR real computer using the same structure. You have 45 minutes.\n\nNext session — we LISTEN. Carlos on the phone, calling IT Support. And we learn the 20 key vocabulary words. See you next time."')
icq(doc, 'What do you do with Carlos\'s email?', 'Underline is and has')
icq(doc, 'How many new sentences do you write?', '3')
icq(doc, 'About whose computer?', 'My real computer')

doc.add_paragraph('')
h(doc, 'Preview — 3:08 a 3:10', 3)
italic_para(doc, '"Next session: Tuning In. We listen to Carlos on the phone — and we learn 20 hardware vocabulary words. Come ready to listen."')

doc.add_page_break()

# ═══ ANSWER KEY ═══
h(doc, 'ANSWER KEY CONSOLIDADO', 1)
h(doc, 'PM-2.3 — The Master Anchor (Reading)', 2)

make_table(doc, ['Actividad','Ítem','Respuesta Correcta','Alternativas'], [
    ['Toolbelt','Spec Sheet','Document with technical info about a product','—'],
    ['Toolbelt','Upgrade','To change a part for a better one','—'],
    ['Toolbelt','Compatible','When two parts can work together','—'],
    ['Toolbelt','Performance','How fast or how well something works','—'],
    ['Toolbelt','Budget','The money available to buy something','—'],
    ['Activity 1','CPU — Current','Intel i3 (old, slow)','"An i3"'],
    ['Activity 1','CPU — Requested','Intel i5 or higher','"A new CPU, i5+"'],
    ['Activity 1','RAM — Current','4 GB','"Only 4 GB"'],
    ['Activity 1','RAM — Requested','16 GB','"16 GB of RAM"'],
    ['Activity 1','Storage — Current','500 GB HDD (mechanical, slow)','"HDD 500 GB"'],
    ['Activity 1','Storage — Requested','1 TB SSD','"An SSD, 1 TB"'],
    ['Activity 1','Monitor — Current','Small, only VGA port','"VGA monitor"'],
    ['Activity 1','Monitor — Requested','HDMI monitor','—'],
    ['Activity 1','Mouse — Current','Broken scroll wheel','"Broken mouse"'],
    ['Activity 1','Mouse — Requested','New mouse','—'],
    ['Activity 1','Total Budget','$1,200','—'],
    ['Activity 2','Decision','Abierta — justificado con texto','Cualquier par defendible'],
    ['Activity 3','Describe','"The CPU is old." / "It is very slow..."','Cualquier frase de problema'],
    ['Activity 3','Request','"I need a new CPU..."','Frase parcial con "I need"'],
    ['Activity 3','Compare','"An SSD is faster than an HDD."','—'],
    ['Exit Ticket','Sent 1','Frase con "is" sobre problema','"The CPU is old."'],
    ['Exit Ticket','Sent 2','Frase con "has" sobre característica','"It has 4 GB of RAM."'],
], shade='2E75B6')

doc.add_paragraph('')
p = doc.add_paragraph(); r = p.add_run('Nota: '); r.bold = True; r.font.name = 'Calibri'
r2 = p.add_run('Activity 2 no tiene respuesta "correcta" única. Se evalúa la calidad de la justificación con evidencia textual.'); r2.font.name = 'Calibri'

doc.add_page_break()

# ═══ DIFFERENTIATION ═══
h(doc, 'DIFFERENTIATION NOTES', 1)
h(doc, 'Fast Finishers', 2)
for t,d in [('Activity 1','"Write 1 additional row: What would YOU request for YOUR workstation?"'),('Activity 2','"What if the budget is $1,000? Which THREE would you approve?"'),('Activity 3','"Write 3 NEW sentences about a phone (not a computer) using the same formulas."'),('Peer tutor','"Help your partner find the sentences in Activity 3. Point to the email."')]:
    bp(doc, f'{t}: ', d)

doc.add_paragraph('')
h(doc, 'More Support Needed', 2)
for t,d in [('Toolbelt','Permitir que mantenga el Toolbelt visible toda la sesión como andaje.'),
    ('Activity 1','Proporcionar versión del email con respuestas SUBRAYADAS (2-3 copias preparadas).'),
    ('Activity 2','Reducir a 3 opciones (CPU, RAM, SSD). Simplificar justificación.'),
    ('Activity 3','Proporcionar funciones con primera mitad: "The _______ is _______."'),
    ('Exit Ticket','Permitir copiar una frase del email como práctica guiada antes de escribir la segunda.')]:
    bp(doc, f'{t}: ', d)

doc.add_page_break()

# ═══ SELF-CHECK ═══
h(doc, 'INSTRUCTOR SELF-CHECK', 1)
p = doc.add_paragraph('Después de la sesión, responda:'); p.runs[0].italic = True
doc.add_paragraph('')
for i,q in enumerate(['¿Todos completaron la tabla de Information Transfer? ¿Cuántos tuvieron 5 filas correctas?','¿Hubo algún momento donde la mayoría parecía perdida? ¿En qué actividad?','¿Los tiempos se cumplieron o hubo desfases? ¿Dónde?','¿Qué errores comunes observé en las justificaciones de Activity 2?','¿Los aprendices identificaron las 3 writing formulas en Activity 3?'],1):
    bp(doc, f'{i}. {q}')
    p = doc.add_paragraph('Respuesta: _______________________________________________')
    for r in p.runs: r.font.name = 'Calibri'; r.font.size = Pt(11)
    doc.add_paragraph('')

# ═══ FOOTER ═══
doc.add_paragraph('')
for txt in ['SESSION 2: READ THE REQUEST — BUILD-OUT','ADSO — GUÍA 1: The Hardware Specialist','Sistema de Prompts Maestros — LG Factory — FPI SENA — Bilingüismo','Instructor Sergio Cortés Perdomo · Marzo 2026']:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt); r.font.size = Pt(9); r.font.name = 'Calibri'; r.font.color.rgb = RGBColor(0x7F,0x8C,0x8D)

doc.save('/Users/Beppo/Projects/fpi-sena-factory/guides/ADSO-G1/ADSO — GUÍA 1 — Session 2 — Read the Request — Build-Out.docx')
print('Done')
