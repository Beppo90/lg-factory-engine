#!/usr/bin/env python3
"""Generate Evaluation Instruments as a Word document."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def sh(cell,color):
    e=OxmlElement('w:shd');e.set(qn('w:fill'),color);e.set(qn('w:val'),'clear')
    cell._tc.get_or_add_tcPr().append(e)

def tbl(doc,headers,rows,sc='1F3A5F'):
    t=doc.add_table(rows=1,cols=len(headers));t.style='Table Grid';t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i];c.text='';r=c.paragraphs[0].add_run(str(h))
        r.font.size=Pt(10);r.font.name='Calibri';r.bold=True;r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF);sh(c,sc)
    for rd in rows:
        row=t.add_row()
        for i,v in enumerate(rd):
            c=row.cells[i];c.text='';r=c.paragraphs[0].add_run(str(v));r.font.size=Pt(9);r.font.name='Calibri'
    return t

def h(d,t,l=1):
    hd=d.add_heading(t,level=l)
    for r in hd.runs:r.font.name='Calibri'

def bp(d,l,v=''):
    p=d.add_paragraph();r=p.add_run(l);r.bold=True;r.font.name='Calibri';r.font.size=Pt(11)
    if v:r2=p.add_run(v);r2.font.name='Calibri';r2.font.size=Pt(11)

def ip(d,t):
    p=d.add_paragraph();r=p.add_run(t);r.italic=True;r.font.name='Calibri';r.font.size=Pt(11)

def st(d,t):
    p=d.add_paragraph(t)
    for r in p.runs:r.font.name='Calibri';r.font.size=Pt(11)

def line(d,n=1):
    for _ in range(n):
        p=d.add_paragraph('____________________________________________________________')
        for r in p.runs:r.font.name='Calibri';r.font.size=Pt(11);r.font.color.rgb=RGBColor(0xBB,0xBB,0xBB)

doc = Document()
for s in doc.sections:
    s.top_margin=Cm(2.5);s.bottom_margin=Cm(2.5);s.left_margin=Cm(2.5);s.right_margin=Cm(2.5)
style=doc.styles['Normal'];style.font.name='Calibri';style.font.size=Pt(11)

# TITLE
for _ in range(4):doc.add_paragraph('')
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('ADSO — GUÍA 1: THE HARDWARE SPECIALIST');r.bold=True;r.font.size=Pt(20);r.font.name='Calibri';r.font.color.rgb=RGBColor(0x1F,0x3A,0x5F)
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Instrumentos de Evaluación Formativa');r.font.size=Pt(16);r.font.name='Calibri';r.font.color.rgb=RGBColor(0x44,0x72,0xC4)
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Checklist de Observación + Feedback Loop');r.font.size=Pt(14);r.font.name='Calibri';r.italic=True
doc.add_paragraph('')
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('⚠️ SOLO PARA EL INSTRUCTOR. No distribuir a los aprendices.');r.bold=True;r.font.size=Pt(12);r.font.name='Calibri';r.font.color.rgb=RGBColor(0xC0,0x39,0x2B)
doc.add_paragraph('')
for txt in ['Sistema de Prompts Maestros — LG Factory — FPI SENA — Bilingüismo','Instructor Sergio Cortés Perdomo · Marzo 2026']:
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(txt);r.font.size=Pt(10);r.font.name='Calibri';r.font.color.rgb=RGBColor(0x7F,0x8C,0x8D)
doc.add_page_break()

# ═══ §3: CHECKLIST ═══
h(doc,'§3: CHECKLIST DE OBSERVACIÓN — EVIDENCIA DE DESEMPEÑO',1)
bp(doc,'Uso: ','Durante Session 6 (The Help Desk) — Speaking Simulation')
bp(doc,'Evaluador: ','Instructor')
bp(doc,'Puntuación: ','10 puntos por aprendiz (5 criterios × 2 pts)')
doc.add_paragraph('')

h(doc,'Criterios de Observación',2)
tbl(doc,['#','Criterio','Competente (2 pts)','En desarrollo (1 pt)','No evidenciado (0 pts)'],[
    ['1','Identifica componentes','≥3 términos técnicos correctamente','1-2 términos o con errores','No usa términos o incorrectamente'],
    ['2','Usa estructuras gram.','is/has correctos en ≥70%','is/has con errores 30-50%','Confunde is/has en >50%'],
    ['3','Mensaje inteligible','Compañero entiende problema + solicitud','Entiende parcialmente, 1-2 repeticiones','No entiende o traducción constante'],
    ['4','Usa frases Skeleton','≥3 frases de forma natural','1-2 frases o lee mecánicamente','No usa frases del Skeleton'],
    ['5','Llega a resolución','Acuerdo claro: NOW vs ORDER','Avanza pero sin acuerdo claro','No intercambio significativo'],
],sc='2E75B6')
doc.add_page_break()

h(doc,'Formato de Registro (1 por aprendiz — fotocopiar)',2)
bp(doc,'ADSO — GUÍA 1: The Hardware Specialist — Session 6: The Help Desk')
doc.add_paragraph('')
bp(doc,'Aprendiz: ');line(doc)
bp(doc,'Fecha: _________________   Ronda: ☐ 1  ☐ 2  ☐ 3   Rol: ☐ Developer  ☐ IT Support')
bp(doc,'Compañero(a): ');line(doc)
doc.add_paragraph('')

tbl(doc,['#','Criterio','2 pts','1 pt','0 pts','Notas'],[
    ['1','Identifica componentes','☐','☐','☐',''],
    ['2','Usa estructuras gram.','☐','☐','☐',''],
    ['3','Mensaje inteligible','☐','☐','☐',''],
    ['4','Usa frases Skeleton','☐','☐','☐',''],
    ['5','Llega a resolución','☐','☐','☐',''],
],sc='4472C4')
doc.add_paragraph('')
bp(doc,'TOTAL: _____ / 10 puntos')
doc.add_paragraph('')
bp(doc,'Frase destacada del aprendiz (literal):')
st(doc,'"____________________________________________________________"')
doc.add_paragraph('')
bp(doc,'Observaciones:')
line(doc,2)
doc.add_paragraph('')
bp(doc,'Instructor: ________________________  Firma: _____________')
doc.add_page_break()

h(doc,'Registro Rápido de Clase',2)
ip(doc,'Usar durante circulación — anotar 3-5 aprendices que necesitan atención especial:')
tbl(doc,['Aprendiz','Ronda','Observación rápida'],[
    ['','',''],['','',''],['','',''],['','',''],['','',''],
],sc='4472C4')
doc.add_page_break()

# ═══ §5: FEEDBACK INDIVIDUAL ═══
h(doc,'§5: FEEDBACK INDIVIDUAL — POST-CUESTIONARIO',1)
bp(doc,'Uso: ','Después de calificar PM-4.2 — entregar en Session 7 o después de clase')
bp(doc,'Formato: ','1 por aprendiz (fotocopiar)')
doc.add_paragraph('')

h(doc,'Formato de Registro',2)
bp(doc,'ADSO — GUÍA 1: The Hardware Specialist — Feedback Individual')
doc.add_paragraph('')
bp(doc,'Aprendiz: ');line(doc)
bp(doc,'Ficha: _________________  Fecha: _________________')
doc.add_paragraph('')

h(doc,'Resultados del Cuestionario (PM-4.2 — 50 pts)',3)
tbl(doc,['Sección','Pts obtenidos','Pts totales','Nivel'],[
    ['1. Reading','____','10',''],['2. Writing','____','10',''],
    ['3. Listening','____','10',''],['4. Vocabulary HOTS','____','10',''],
    ['5. Grammar HOTS','____','10',''],['TOTAL','____','50',''],
],sc='4472C4')
doc.add_paragraph('')
bp(doc,'NIVEL GENERAL:')
st(doc,'☐ Avanzado (45-50)  ☐ Adecuado (35-44)  ☐ En desarrollo (25-34)  ☐ Requiere refuerzo (<25)')
doc.add_paragraph('')

bp(doc,'FORTALEZA PRINCIPAL:');line(doc)
bp(doc,'ÁREA DE MEJORA PRINCIPAL:');line(doc)
bp(doc,'RECOMENDACIÓN DEL INSTRUCTOR:');line(doc,2)
doc.add_paragraph('')

bp(doc,'OBSERVACIÓN DE DESEMPEÑO ORAL (Session 6):')
bp(doc,'Puntuación Checklist: _____ / 10')
bp(doc,'Comentario: ');line(doc)
doc.add_paragraph('')
ip(doc,'Este feedback es FORMATIVO — no es una calificación oficial. Úsalo para saber qué practicar más.')
doc.add_paragraph('')
bp(doc,'Instructor: ________________________  Firma: _____________')
doc.add_page_break()

h(doc,'Guía de Completado para el Instructor',2)
bp(doc,'Fortaleza Principal — Qué buscar:')
st(doc,'- La sección con MÁS puntos obtenidos')
st(doc,'- O un logro específico (ej: "Usó \'compatible\' correctamente en la simulación oral")')
doc.add_paragraph('')
bp(doc,'Área de Mejora — Qué buscar:')
st(doc,'- La sección con MENOS puntos obtenidos')
st(doc,'- O un patrón de error recurrente (ej: "Confunde \'is\' y \'has\'")')
doc.add_paragraph('')
bp(doc,'Recomendación — Ejemplos:')
st(doc,'- "Practica el vocabulario: haz flashcards con los 20 términos."')
st(doc,'- "Revisa la diferencia entre \'is\' (describe) y \'has\' (specs). Workbook Ch. 4."')
st(doc,'- "Practica leer emails en voz alta para mejorar fluidez oral."')
st(doc,'- "Reescribe tu Tech Request prestando atención al formato."')
doc.add_page_break()

# ═══ §5: FEEDBACK COLECTIVO ═══
h(doc,'§5: FEEDBACK COLECTIVO — GUÍA PARA EL INSTRUCTOR',1)
bp(doc,'Uso: ','Session 7 (Prove What You Know) — 25 minutos de feedback colectivo')
doc.add_paragraph('')

h(doc,'Paso a Paso',2)
tbl(doc,['Paso','Tiempo','Acción','Qué hacer'],[
    ['1','2 min','Set the tone','"I\'m NOT giving grades today. Let\'s learn from common mistakes."'],
    ['2','5 min','Section 1 (Reading)','Errores de scanning. "The answer was IN THE TEXT." NO revelar respuestas.'],
    ['3','5 min','Section 2 (Writing)','Errores de formato (Subject line) y grammar (is/has). Ejemplos en tablero.'],
    ['4','5 min','Section 3 (Listening)','"compatible" no "compitable." Escribir correcto en verde.'],
    ['5','5 min','Section 4-5 (V/G)','Odd One Out: "JUSTIFY with a sentence." Error Log: "Those cables ARE."'],
    ['6','3 min','Questions','"Any questions?" 2-3 preguntas. Guiar sin revelar.'],
],sc='2E75B6')
doc.add_paragraph('')

h(doc,'Errores Comunes Esperados',2)
bp(doc,'Section 1 (Reading):')
st(doc,'- No encontrar datos en Information Extraction / Confundir NG con False')
doc.add_paragraph('')
bp(doc,'Section 2 (Writing):')
st(doc,'- Olvidar Subject line / "I want" vs "I need" / No cerrar con "Thank you"')
doc.add_paragraph('')
bp(doc,'Section 3 (Listening):')
st(doc,'"compitable" vs "compatible" / No distinguir is/has en audio')
doc.add_paragraph('')
bp(doc,'Section 4 (Vocabulary):')
st(doc,'- Justificar mal el Odd One Out / Exceder presupuesto')
doc.add_paragraph('')
bp(doc,'Section 5 (Grammar):')
st(doc,'"Those cable is" → "Those cables ARE" / "more faster" → "faster"')
doc.add_page_break()

h(doc,'Registro de Errores Observados',2)
ip(doc,'Durante la calificación, anotar los 3-5 errores más comunes por sección:')
tbl(doc,['Sección','Error más común','# de aprendices'],[
    ['1. Reading','',''],['2. Writing','',''],['3. Listening','',''],
    ['4. Vocabulary','',''],['5. Grammar','',''],
],sc='4472C4')

doc.add_page_break()
for txt in ['ADSO — GUÍA 1: The Hardware Specialist — Instrumentos de Evaluación Formativa','Sistema de Prompts Maestros — LG Factory — FPI SENA — Bilingüismo','Instructor Sergio Cortés Perdomo · Marzo 2026']:
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(txt);r.font.size=Pt(9);r.font.name='Calibri';r.font.color.rgb=RGBColor(0x7F,0x8C,0x8D)

doc.save('/Users/Beppo/Projects/fpi-sena-factory/guides/ADSO-G1/ADSO — GUÍA 1 — Instrumentos de Evaluación Formativa.docx')
print('Done')
